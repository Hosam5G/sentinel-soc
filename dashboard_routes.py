"""
Equilibrium Sentinel — dashboard_routes.py  (LIVE engine)
==================================================================
A RUNNABLE backend that makes the dashboard genuinely live:

  * Real host telemetry via psutil (CPU / RAM / network), sampled
    in the background into rolling history for the charts.
  * Real-time push to the browser over Server-Sent Events (SSE):
    GET /api/stream  -> the graphs update instantly, no slow polling.
  * Real security posture: threats + score computed from a server-side
    findings store; Apply / Dismiss actually mutate that store.
  * Event-driven NOTIFICATIONS + a SURPRISE report when a scan finds a
    new real vulnerability.
  * Professional REPORTS: daily / weekly / monthly / yearly, generated
    from the collected telemetry and findings.
  * Local MODEL chat (Ollama) with an easy q4 <-> q8 switch.

Wherever you see  ### TODO اربط محرّكك هنا ###  swap the placeholder for
a call into your real engine.

------------------------------------------------------------------
RUN (same-origin on 127.0.0.1:8000):
    pip install flask psutil requests
    python dashboard_routes.py
    # open http://127.0.0.1:8000/

MODEL (optional, for the Ask box + q4/q8 switch):
    # install Ollama, pull the two Foundation-Sec cyber models, e.g.:
    #   ollama pull Foundation-Sec-8B-Instruct-GGUF
    #   ollama pull Foundation-Sec-8B-Instruct-GGUF:Q4_K_M
    # if your tags differ, override:
    set SENTINEL_MODEL_Q8=Foundation-Sec-8B-Instruct-GGUF
    set SENTINEL_MODEL_Q4=Foundation-Sec-8B-Instruct-GGUF:Q4_K_M
==================================================================
"""

import os
import time
import json
import random
import threading
import collections
import datetime as dt
import hashlib
import hmac
import hmac as _hmac
import base64
import struct

from flask import Blueprint, Flask, Response, jsonify, request, send_from_directory, make_response

try:
    import psutil
    HAVE_PSUTIL = True
except Exception:
    HAVE_PSUTIL = False

try:
    import requests
    HAVE_REQUESTS = True
except Exception:
    HAVE_REQUESTS = False

DASHBOARD_DIR = os.environ.get(
    "SENTINEL_DASHBOARD_DIR",
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "dashboard"),
)

bp = Blueprint("dashboard", __name__)

# ==================================================================
#  AUTH PRIMITIVES (defined early so route decorators can use them)
# ==================================================================
import secrets as _secrets
from functools import wraps as _wraps

_ROLE_RANK = {"viewer": 1, "analyst": 2, "admin": 3}
_SESSIONS = {}                  # token -> {user, csrf, created, expires, ip}
_auth_lock = threading.Lock()
_REQUIRE_AUTH_ALL = os.environ.get("SENTINEL_REQUIRE_AUTH", "").strip() in ("1", "true", "True")


def _new_session(uname):
    tok = _secrets.token_urlsafe(32)
    now = int(time.time() * 1000)
    _SESSIONS[tok] = {"user": uname, "csrf": _secrets.token_urlsafe(24),
                      "created": now, "expires": now + 12 * 3600 * 1000,
                      "ip": request.remote_addr if request else ""}
    return tok


def _current_session():
    tok = request.cookies.get("sx_session") if request else None
    if not tok:
        return None
    s = _SESSIONS.get(tok)
    if not s:
        return None
    if s["expires"] < int(time.time() * 1000):
        _SESSIONS.pop(tok, None)
        return None
    return s


def _role_ok(have, need):
    return _ROLE_RANK.get(have, 0) >= _ROLE_RANK.get(need, 99)


def require_auth(role="viewer"):
    """Guard a route. Bootstrap-friendly: if no users exist yet, access is allowed
    so the first admin can be created. Mutating methods also require a CSRF header."""
    def deco(fn):
        @_wraps(fn)
        def wrapper(*a, **kw):
            if not _USERS:                       # first-run bootstrap
                return fn(*a, **kw)
            s = _current_session()
            if not s or s["user"] not in _USERS:
                return jsonify({"ok": False, "error": "authentication required"}), 401
            if not _role_ok(_USERS[s["user"]].get("role", "viewer"), role):
                try:
                    audit("forbidden", f"{request.path} needs {role}", user=s["user"])
                except Exception:
                    pass
                return jsonify({"ok": False, "error": "insufficient permissions"}), 403
            if request.method in ("POST", "PUT", "DELETE", "PATCH"):
                if not _hmac.compare_digest(request.headers.get("X-CSRF-Token") or "", s["csrf"]):
                    return jsonify({"ok": False, "error": "invalid CSRF token"}), 403
            return fn(*a, **kw)
        return wrapper
    return deco


# ==================================================================
#  AT-REST ENCRYPTION for stored JSON (reports / events / audit / settings)
#  Uses a locally-generated key (sentinel.key). Prefers AES-GCM via the
#  'cryptography' package when present; otherwise a dependency-free
#  authenticated SHA-256 CTR + HMAC cipher. Plaintext files are read and
#  transparently migrated to encrypted form on the next save.
#  Disable with SENTINEL_NO_ENCRYPT=1.
# ==================================================================
_KEY_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel.key")
_ENCRYPT = os.environ.get("SENTINEL_NO_ENCRYPT", "").strip() not in ("1", "true", "True")
_key_cache = [None]
try:
    from cryptography.fernet import Fernet as _Fernet
    _HAVE_FERNET = True
except Exception:
    _HAVE_FERNET = False


def _enc_key():
    if _key_cache[0] is not None:
        return _key_cache[0]
    try:
        with open(_KEY_FILE, "r", encoding="utf-8") as f:
            k = f.read().strip()
    except Exception:
        import secrets as _s
        k = _s.token_hex(32)
        try:
            with open(_KEY_FILE, "w", encoding="utf-8") as f:
                f.write(k)
            if os.name != "nt":
                os.chmod(_KEY_FILE, 0o600)
            else:
                # Windows: restrict the key file to the current user only.
                # Removes inherited permissions, then grants just this account.
                try:
                    import subprocess as _sp2, getpass as _gp
                    user = os.environ.get("USERNAME") or _gp.getuser()
                    _sp2.run(["icacls", _KEY_FILE, "/inheritance:r"],
                             capture_output=True, timeout=10,
                             creationflags=0x08000000)
                    _sp2.run(["icacls", _KEY_FILE, "/grant:r", f"{user}:F"],
                             capture_output=True, timeout=10,
                             creationflags=0x08000000)
                except Exception:
                    pass
        except Exception:
            pass
    _key_cache[0] = k
    return k


def _ks(keyb, nonce, n):
    out = bytearray()
    ctr = 0
    while len(out) < n:
        out += hashlib.sha256(keyb + nonce + ctr.to_bytes(8, "big")).digest()
        ctr += 1
    return bytes(out[:n])


def _secure_save(path, obj):
    raw = json.dumps(obj, ensure_ascii=False).encode("utf-8")
    if not _ENCRYPT:
        with open(path, "w", encoding="utf-8") as f:
            f.write(raw.decode("utf-8"))
        return
    try:
        keyhex = _enc_key()
        if _HAVE_FERNET:
            fk = base64.urlsafe_b64encode(hashlib.sha256(keyhex.encode()).digest())
            env = {"enc": "fernet", "data": _Fernet(fk).encrypt(raw).decode("ascii")}
        else:
            keyb = bytes.fromhex(keyhex)
            import secrets as _s
            nonce = _s.token_bytes(16)
            ct = bytes(a ^ b for a, b in zip(raw, _ks(keyb, nonce, len(raw))))
            mac = hmac.new(keyb, nonce + ct, hashlib.sha256).hexdigest()
            env = {"enc": "sha256ctr", "nonce": nonce.hex(),
                   "data": base64.b64encode(ct).decode("ascii"), "mac": mac}
        with open(path, "w", encoding="utf-8") as f:
            json.dump(env, f)
    except Exception:
        # last resort: never lose data
        with open(path, "w", encoding="utf-8") as f:
            f.write(raw.decode("utf-8"))


def _secure_load(path, default):
    try:
        with open(path, "r", encoding="utf-8") as f:
            text = f.read()
    except Exception:
        return default
    if not text.strip():
        return default
    try:
        obj = json.loads(text)
    except Exception:
        return default
    if isinstance(obj, dict) and obj.get("enc") in ("fernet", "sha256ctr"):
        try:
            keyhex = _enc_key()
            if obj["enc"] == "fernet" and _HAVE_FERNET:
                fk = base64.urlsafe_b64encode(hashlib.sha256(keyhex.encode()).digest())
                raw = _Fernet(fk).decrypt(obj["data"].encode("ascii"))
            else:
                keyb = bytes.fromhex(keyhex)
                nonce = bytes.fromhex(obj["nonce"])
                ct = base64.b64decode(obj["data"])
                if not hmac.compare_digest(hmac.new(keyb, nonce + ct, hashlib.sha256).hexdigest(), obj["mac"]):
                    return default
                raw = bytes(a ^ b for a, b in zip(ct, _ks(keyb, nonce, len(ct))))
            return json.loads(raw.decode("utf-8"))
        except Exception:
            return default
    return obj                                  # plaintext (legacy) — migrates on next save



# ==================================================================
#  TELEMETRY SAMPLER (real host data -> rolling history for charts)
# ==================================================================
HIST_N = 40          # points kept per metric
_PROC_START = time.time()   # fallback uptime origin when psutil unavailable
SAMPLE_SEC = 1.0     # sample once per second for accuracy + responsiveness
_hist = {k: collections.deque(maxlen=HIST_N) for k in
         ("cpu", "ram", "traffic", "netIn", "netOut", "threat")}
_score_hist = collections.deque(maxlen=HIST_N)   # security-score trend over time
_latest = {"cpu": 0, "ram": 0.0, "netIn": 0.0, "netOut": 0.0, "traffic": 0, "threat": 0}
_latest_extra = {"ram_pct": 0}
_lock = threading.Lock()
_net_prev = {"t": None, "sent": 0, "recv": 0}
_sampler_started = False


def _sample_once():
    if HAVE_PSUTIL:
        # blocking 1s sample = accurate average that matches Task Manager,
        # and it also paces the loop (no separate sleep needed).
        cpu = psutil.cpu_percent(interval=SAMPLE_SEC)
        vm = psutil.virtual_memory()
        # match Task Manager exactly: vm.percent is Windows' dwMemoryLoad (what TM shows)
        ram_pct = int(round(vm.percent))
        ram_gb = round(vm.total * vm.percent / 100 / (1024 ** 3), 1)
        nio = psutil.net_io_counters()
        now = time.time()
        net_in = net_out = 0.0
        if _net_prev["t"] is not None:
            d = max(0.001, now - _net_prev["t"])
            net_in = max(0.0, (nio.bytes_recv - _net_prev["recv"]) * 8 / 1e6 / d)
            net_out = max(0.0, (nio.bytes_sent - _net_prev["sent"]) * 8 / 1e6 / d)
        _net_prev.update(t=now, sent=nio.bytes_sent, recv=nio.bytes_recv)
        cpu_v = int(round(cpu)); net_in = round(net_in, 1); net_out = round(net_out, 1)
    else:
        # No psutil -> we CANNOT measure. Report zeros (UI shows an "estimated"
        # banner), never fabricated numbers that look like real telemetry.
        cpu_v = 0
        ram_gb = 0.0
        ram_pct = 0
        net_in = 0.0
        net_out = 0.0

    _latest_extra["ram_pct"] = ram_pct

    traffic = round(net_in + net_out, 1)
    threat = active_threats()
    with _lock:
        _hist["cpu"].append(cpu_v); _hist["ram"].append(ram_gb)
        _hist["netIn"].append(net_in); _hist["netOut"].append(net_out)
        _hist["traffic"].append(traffic); _hist["threat"].append(threat)
        _latest.update(cpu=cpu_v, ram=ram_gb, netIn=net_in,
                       netOut=net_out, traffic=traffic, threat=threat)
        _score_hist.append(posture_score())


def _sampler_loop():
    # Optional power-saver: set SENTINEL_LOW_CPU=1 to sample every 3s instead of
    # every 1s (much lighter on modest laptops, slightly less responsive graphs).
    low = os.environ.get("SENTINEL_LOW_CPU", "").strip() in ("1", "true", "yes")
    idle_extra = 2.0 if low else 0.0
    cycle = 0
    while True:
        try:
            # _refresh_protection() shells out to PowerShell and is relatively
            # heavy; it changes rarely, so only refresh it every ~30 cycles
            # instead of every second. This is a big CPU win on weak machines.
            if cycle % 30 == 0:
                _refresh_protection()
            cycle += 1
            _sample_once()         # blocks ~SAMPLE_SEC when psutil paces it
            if not HAVE_PSUTIL:
                time.sleep(SAMPLE_SEC)
            if idle_extra:
                time.sleep(idle_extra)
        except Exception:
            time.sleep(SAMPLE_SEC)


def _ensure_sampler():
    global _sampler_started
    if _sampler_started:
        return
    _sampler_started = True
    if HAVE_PSUTIL:
        psutil.cpu_percent(interval=None)
        time.sleep(0.12)
    _sample_once()
    with _lock:
        for dq in _hist.values():
            if dq:
                first = dq[0]
                while len(dq) < HIST_N:
                    dq.appendleft(first)
    threading.Thread(target=_sampler_loop, daemon=True).start()
    _ensure_summary()   # periodic model-generated overview summary
    _ensure_auto_reports()   # automatic daily/weekly/monthly/yearly reports
    _ensure_metrics()   # background cache for disks/processes/connections
    _ensure_initial_scan()   # one real vulnerability scan of this machine at startup
    _ensure_threat_intel()   # keep the exploited-CVE feed updated (privacy: GET only)
    _ensure_incident_detection()   # watch Windows event logs for real incidents
    _ensure_yara_updater()   # periodic community YARA rule refresh (privacy: GET only)
    _load_allowlist()        # user's trusted-app allowlist (cuts false positives)
    _load_baseline()         # behavioral baseline (learn-then-deviate)
    _start_realtime()        # real-time protection — ALWAYS ON (event-driven via WMI)
    # persist the behavioral baseline on shutdown so the last observations (and the
    # learning progress) are never lost if the user closes the app between saves.
    import atexit
    atexit.register(_save_baseline)


def _series(key):
    with _lock:
        return list(_hist[key])


# ==================================================================
#  SECURITY STORE (real posture, real Apply/Dismiss)
# ==================================================================
# Each finding carries both languages + a live state. Actions mutate it.
# Starts EMPTY — populated only by a real scan of THIS machine (run_scan_now).
# No fabricated findings on imaginary hosts (web-03/db-01/…): those misled the
# user into seeing vulnerabilities that don't exist on their actual computer.
_FINDINGS = []
_RECS = []          # REAL recommendations — derived from the actual scan (see _rebuild_recs)
_sec_lock = threading.Lock()
_SEV_WEIGHT = {"critical": 16, "high": 9, "medium": 4, "low": 1}
_OPEN_STATES = ("open", "investigating")


def _rebuild_recs(findings):
    """Derive REAL recommendations from the actual scan results — each one maps
    to a real condition on THIS machine and (where possible) to an executable
    remediation plan. Preserves the user's apply/dismiss state across rescans."""
    prev = {r["id"]: r["state"] for r in _RECS}
    recs = []

    def add(rid, en, ar, cve=None):
        recs.append({"id": rid, "state": prev.get(rid, "open"), "en": en, "ar": ar, "cve": cve})

    open_f = [f for f in findings if f.get("st") in _OPEN_STATES]
    by_cve = {f["cve"]: f for f in open_f}

    # 1) ports
    port_f = [f for f in open_f if f["cve"].startswith("PORT-")]
    if port_f:
        ports = ", ".join(f["cve"].split("-")[1] for f in port_f[:6])
        add("rec-ports",
            f"Close or firewall the risky open ports: {ports} (use the Fix button on each finding).",
            f"أغلق أو احظر المنافذ الخطرة المفتوحة: {ports} (استخدم زر «معالجة» على كل ثغرة).",
            port_f[0]["cve"])
    # 2) baseline configs
    if "FW-DISABLED" in by_cve:
        add("rec-fw", "Turn Windows Firewall back ON — one click via the Fix button.",
            "أعد تشغيل جدار حماية Windows — بنقرة عبر زر «معالجة».", "FW-DISABLED")
    if "AV-RTP-OFF" in by_cve:
        add("rec-def", "Re-enable Defender real-time protection.",
            "أعد تفعيل الحماية اللحظية في Defender.", "AV-RTP-OFF")
    if "SMBV1-ENABLED" in by_cve:
        add("rec-smb1", "Disable the legacy SMBv1 protocol (WannaCry vector).",
            "عطّل بروتوكول SMBv1 القديم (ناقل WannaCry).", "SMBV1-ENABLED")
    if "UAC-DISABLED" in by_cve:
        add("rec-uac", "Re-enable UAC elevation prompts.",
            "أعد تفعيل نوافذ UAC.", "UAC-DISABLED")
    if "GUEST-ACTIVE" in by_cve:
        add("rec-guest", "Disable the built-in Guest account.",
            "عطّل حساب الضيف.", "GUEST-ACTIVE")
    # 3) software updates
    sw_f = [f for f in open_f if f.get("cat") in ("software", "intel")]
    if sw_f:
        add("rec-updates",
            f"Update the {len(sw_f)} flagged application(s) — winget plans are available via Fix.",
            f"حدّث التطبيقات المُعلَّمة ({len(sw_f)}) — خطط winget متاحة عبر «معالجة».",
            sw_f[0]["cve"])
    # 4) account hygiene (real check: do all users have MFA?)
    try:
        no_mfa = [u for u, rec in _USERS.items() if not (rec.get("mfa") or {}).get("enabled")]
        if no_mfa:
            add("rec-mfa",
                f"Enable MFA for: {', '.join(no_mfa[:5])} (Profile page → Enable MFA).",
                f"فعّل المصادقة الثنائية لـ: {', '.join(no_mfa[:5])} (صفحة الحساب ← تفعيل MFA).")
    except Exception:
        pass
    # 5) standing hygiene when everything else is clean
    if not recs:
        add("rec-clean",
            "Baseline looks healthy — keep automatic scans on and review the weekly report.",
            "الوضع الأساسي سليم — أبقِ الفحص التلقائي مفعّلاً وراجع التقرير الأسبوعي.")
    with _sec_lock:
        _RECS[:] = recs


# ==================================================================
#  REAL VULNERABILITY / CONFIGURATION SCANNER
#  Scans the ACTUAL machine — OS patch status, security baseline
#  (firewall, Defender, SMBv1, RDP, UAC, guest), risky open ports,
#  and installed-software inventory. Produces real, host-specific
#  findings (replaces the static demo list). Windows-native, with a
#  graceful cross-platform fallback.
# ==================================================================
import subprocess as _sp
import platform as _plat

_last_scan = {"ts": 0, "running": False, "count": 0}


def _run_cmd(args, timeout=12):
    try:
        flags = 0x08000000 if os.name == "nt" else 0      # CREATE_NO_WINDOW
        r = _sp.run(args, capture_output=True, text=True, timeout=timeout,
                    creationflags=flags) if os.name == "nt" else \
            _sp.run(args, capture_output=True, text=True, timeout=timeout)
        return (r.stdout or "") + (r.stderr or "")
    except Exception:
        return ""


def _reg_get(hive, path, name):
    if os.name != "nt":
        return None
    try:
        import winreg
        h = {"HKLM": winreg.HKEY_LOCAL_MACHINE, "HKCU": winreg.HKEY_CURRENT_USER}[hive]
        with winreg.OpenKey(h, path) as k:
            v, _ = winreg.QueryValueEx(k, name)
            return v
    except Exception:
        return None


def _f(fid, host, sev, score, en, ar, fix_en, fix_ar, cat):
    return {"cve": fid, "asset": host, "sev": sev, "score": score, "st": "open",
            "title_en": en, "title_ar": ar, "fix_en": fix_en, "fix_ar": fix_ar,
            "cat": cat, "first_seen": int(time.time() * 1000)}


_RISKY_PORTS = {
    21: ("FTP", "high", 7.5), 23: ("Telnet", "critical", 9.1), 135: ("RPC", "low", 3.5),
    139: ("NetBIOS", "medium", 5.0), 445: ("SMB", "medium", 5.5), 1433: ("MSSQL", "high", 7.3),
    3306: ("MySQL", "medium", 6.0), 3389: ("RDP", "high", 7.8), 5900: ("VNC", "high", 7.5),
}
# Ports that are normal on a personal Windows machine and only matter if the
# service is reachable from OUTSIDE the local network. We label these clearly.
_LAN_NORMAL_PORTS = {135, 139, 445}


def _port_finding(p, host):
    """Build a port finding with severity that reflects REAL exposure.
    A bound port listening on 0.0.0.0 reachable from the internet is serious;
    the same port on a home LAN behind a router/NAT is routine."""
    svc, sev, sc = _RISKY_PORTS[p]
    if p in _LAN_NORMAL_PORTS:
        en = (f"Windows networking port {p} ({svc}) is open. This is NORMAL on a "
              f"home/office PC and only a risk if your machine is directly exposed "
              f"to the internet. If you don't share files/printers, you can close it.")
        ar = (f"منفذ شبكة Windows رقم {p} ‏({svc}) مفتوح. هذا طبيعي على جهاز منزلي/مكتبي، "
              f"ولا يمثّل خطراً إلا إذا كان جهازك مكشوفاً مباشرةً للإنترنت. إن كنت لا تشارك "
              f"ملفات/طابعات، يمكنك إغلاقه.")
        fix_en = f"Optional: block port {p} via the Fix button if you don't use Windows file sharing."
        fix_ar = f"اختياري: احظر المنفذ {p} عبر زر «معالجة» إن كنت لا تستخدم مشاركة ملفات Windows."
    else:
        en = f"Risky service {svc} is listening on port {p}"
        ar = f"خدمة خطرة {svc} تستمع على المنفذ {p}"
        fix_en = f"Close port {p} or restrict {svc} to trusted networks only"
        fix_ar = f"أغلق المنفذ {p} أو اقصر {svc} على الشبكات الموثوقة"
    return _f(f"PORT-{p}-{svc}", host, sev, sc, en, ar, fix_en, fix_ar, "exposure")


def _scan_installed_software():
    """Real installed-software inventory from the Windows registry."""
    apps = []
    if os.name != "nt":
        return apps
    try:
        import winreg
        roots = [(winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall"),
                 (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\WOW6432Node\Microsoft\Windows\CurrentVersion\Uninstall"),
                 (winreg.HKEY_CURRENT_USER,  r"SOFTWARE\Microsoft\Windows\CurrentVersion\Uninstall")]
        for hive, path in roots:
            try:
                with winreg.OpenKey(hive, path) as key:
                    for i in range(winreg.QueryInfoKey(key)[0]):
                        try:
                            sub = winreg.EnumKey(key, i)
                            with winreg.OpenKey(key, sub) as sk:
                                def g(n):
                                    try:
                                        return winreg.QueryValueEx(sk, n)[0]
                                    except Exception:
                                        return ""
                                name = g("DisplayName")
                                if name:
                                    apps.append({"name": str(name), "version": str(g("DisplayVersion"))})
                        except Exception:
                            continue
            except Exception:
                continue
    except Exception:
        pass
    # de-dup
    seen, out = set(), []
    for a in apps:
        if a["name"] not in seen:
            seen.add(a["name"]); out.append(a)
    return out


def _run_security_scan():
    """Run a real scan of this machine and return host-specific findings."""
    host = _plat.node() or "localhost"
    findings = []
    win = (os.name == "nt")

    # ---- OS version / EOL ----
    osrel = _plat.platform()
    if win:
        ver = _plat.version()          # e.g. 10.0.19045
        try:
            build = int(ver.split(".")[2])
        except Exception:
            build = 0
        if build and build < 19044:    # pre-21H2 Win10 / very old
            findings.append(_f("OS-OUTDATED", host, "high", 7.5,
                "Operating system build is outdated and may be past end-of-support",
                "إصدار نظام التشغيل قديم وقد يكون خارج الدعم",
                "Upgrade Windows to a supported build and install all updates",
                "رقِّ ويندوز إلى إصدار مدعوم وثبّت كل التحديثات", "patch"))

    # ---- missing recent patches (last hotfix date) ----
    if win:
        out = _run_cmd(["powershell", "-NoProfile", "-Command",
                        "(Get-HotFix | Sort-Object InstalledOn -Descending | "
                        "Select-Object -First 1).InstalledOn.ToString('yyyy-MM-dd')"], 25)
        import re as _r2, datetime as _dt
        m = _r2.search(r"(\d{4}-\d{2}-\d{2})", out)
        if m:
            try:
                last = _dt.datetime.strptime(m.group(1), "%Y-%m-%d")
                days = (_dt.datetime.now() - last).days
                if days > 35:
                    findings.append(_f("PATCH-STALE", host, "high", 7.8,
                        f"No security updates installed in {days} days",
                        f"لم تُثبَّت تحديثات أمنية منذ {days} يوماً",
                        "Run Windows Update and install all pending security patches",
                        "شغّل Windows Update وثبّت كل التحديثات الأمنية المعلّقة", "patch"))
            except Exception:
                pass

    # ---- firewall ----
    if win:
        fw = _run_cmd(["netsh", "advfirewall", "show", "allprofiles", "state"], 10)
        if fw and fw.lower().count("off") > 0:
            findings.append(_f("FW-DISABLED", host, "high", 7.4,
                "Windows Firewall is disabled on one or more profiles",
                "جدار حماية ويندوز معطّل في ملف تعريف واحد أو أكثر",
                "Enable the firewall on all profiles: netsh advfirewall set allprofiles state on",
                "فعّل الجدار الناري لكل الملفات: netsh advfirewall set allprofiles state on", "config"))

    # ---- Windows Defender real-time protection (reuse protection engine) ----
    try:
        prot = protection_status()
        if win and prot and not prot.get("realtime", True):
            findings.append(_f("AV-RTP-OFF", host, "critical", 8.6,
                "Real-time antivirus protection is turned off",
                "الحماية الفورية لمكافحة الفيروسات معطّلة",
                "Enable real-time protection in Windows Security",
                "فعّل الحماية الفورية في أمن ويندوز", "config"))
    except Exception:
        pass

    # ---- SMBv1 (EternalBlue) ----
    if win:
        smb1 = _reg_get("HKLM", r"SYSTEM\CurrentControlSet\Services\LanmanServer\Parameters", "SMB1")
        if smb1 == 1:
            findings.append(_f("SMBV1-ENABLED", host, "critical", 9.3,
                "SMBv1 is enabled — vulnerable to EternalBlue/WannaCry-class attacks",
                "بروتوكول SMBv1 مُفعّل — معرّض لهجمات من نوع EternalBlue/WannaCry",
                "Disable SMBv1: Disable-WindowsOptionalFeature -Online -FeatureName SMB1Protocol",
                "عطّل SMBv1: Disable-WindowsOptionalFeature -Online -FeatureName SMB1Protocol", "config"))

    # ---- RDP exposure ----
    if win:
        deny = _reg_get("HKLM", r"SYSTEM\CurrentControlSet\Control\Terminal Server", "fDenyTSConnections")
        nla = _reg_get("HKLM", r"SYSTEM\CurrentControlSet\Control\Terminal Server\WinStations\RDP-Tcp", "UserAuthentication")
        if deny == 0:   # RDP enabled
            if nla == 0:
                findings.append(_f("RDP-NO-NLA", host, "high", 8.1,
                    "RDP is enabled without Network Level Authentication (NLA)",
                    "RDP مُفعّل بدون مصادقة على مستوى الشبكة (NLA)",
                    "Require NLA for RDP, or disable RDP if not needed",
                    "اشترط NLA لـ RDP، أو عطّل RDP إن لم يكن مطلوباً", "config"))

    # ---- UAC ----
    if win:
        lua = _reg_get("HKLM", r"SOFTWARE\Microsoft\Windows\CurrentVersion\Policies\System", "EnableLUA")
        if lua == 0:
            findings.append(_f("UAC-DISABLED", host, "high", 7.2,
                "User Account Control (UAC) is disabled",
                "التحكّم بحساب المستخدم (UAC) معطّل",
                "Re-enable UAC (EnableLUA = 1) and reboot",
                "أعد تفعيل UAC (EnableLUA = 1) وأعد التشغيل", "config"))

    # ---- guest account ----
    if win:
        g = _run_cmd(["net", "user", "guest"], 8)
        if g and ("active" in g.lower()) and ("yes" in g.lower()):
            # crude: 'Account active   Yes'
            import re as _r3
            if _r3.search(r"active\s+yes", g.lower()):
                findings.append(_f("GUEST-ACTIVE", host, "medium", 5.3,
                    "The built-in Guest account is enabled",
                    "حساب الضيف المدمج مُفعّل",
                    "Disable the Guest account: net user guest /active:no",
                    "عطّل حساب الضيف: net user guest /active:no", "config"))

    # ---- risky open listening ports (real, from psutil) ----
    try:
        ports = _net_conn_stats().get("ports", [])
        for p in ports:
            if p in _RISKY_PORTS:
                findings.append(_port_finding(p, host))
    except Exception:
        pass

    # ---- installed-software inventory (real) + EOL heuristics ----
    sw = _scan_installed_software()
    _last_scan["software"] = sw
    eol_markers = {"python 2": ("Python 2 (EOL)", "بايثون 2 (منتهي الدعم)"),
                   "java 7": ("Java 7 (EOL)", "جافا 7 (منتهي الدعم)"),
                   "java 8": ("Java 8 (legacy)", "جافا 8 (قديم)"),
                   "flash": ("Adobe Flash (EOL)", "أدوبي فلاش (منتهي الدعم)"),
                   "silverlight": ("Microsoft Silverlight (EOL)", "مايكروسوفت سيلفرلايت (منتهي الدعم)"),
                   "internet explorer": ("Internet Explorer (EOL)", "إنترنت إكسبلورر (منتهي الدعم)")}
    low = " | ".join((a["name"] + " " + a["version"]).lower() for a in sw)
    for marker, (en, ar) in eol_markers.items():
        if marker in low:
            findings.append(_f(f"EOL-{marker.replace(' ','-').upper()}", host, "medium", 6.0,
                f"End-of-life software detected: {en}",
                f"برنامج منتهي الدعم: {ar}",
                f"Remove or upgrade {en} to a supported version",
                f"أزل أو رقِّ {ar} إلى إصدار مدعوم", "software"))

    _last_scan.update(ts=int(time.time() * 1000), count=len(findings))
    # cross-reference installed software with the live exploited-CVE feed (local match)
    try:
        findings.extend(_match_kev_to_software())
    except Exception:
        pass

    # ---- LIVE DETECTION ENGINES: surface real threats as findings ----
    try:
        proc = scan_processes()
        for h in proc.get("hits", [])[:20]:
            findings.append(_f(f"PROC-{h['pid']}", h.get("name") or host, h["sev"],
                9.0 if h["sev"] == "critical" else 7.0,
                f"Suspicious process '{h.get('name')}' (pid {h['pid']}): {h['reasons'][0]}",
                f"عملية مشبوهة '{h.get('name')}' (pid {h['pid']}): {h['reasons'][0]}",
                f"Investigate process {h['pid']} at {h.get('path','?')}; terminate if malicious.",
                f"افحص العملية {h['pid']} في {h.get('path','?')}؛ أنهِها إن كانت خبيثة.", "process"))
    except Exception:
        pass
    try:
        persist = scan_persistence()
        for it in persist.get("items", [])[:20]:
            findings.append(_f(f"PERSIST-{abs(hash(it['name']+it['source']))%100000}",
                host, it["sev"], 7.0,
                f"Suspicious {it['source']} '{it['name']}': {it['reasons'][0]}",
                f"استمرارية مشبوهة في {it['source']} ‏'{it['name']}': {it['reasons'][0]}",
                f"Review this autostart entry; remove it if you don't recognise it.",
                f"راجع مدخل بدء التشغيل هذا؛ احذفه إن كنت لا تعرفه.", "persistence"))
    except Exception:
        pass
    try:
        fim = fim_check()
        for ch in fim.get("changes", [])[:20]:
            findings.append(_f(f"FIM-{abs(hash(ch['path']))%100000}", host, ch["sev"], 7.5,
                f"Watched file {ch['kind']}: {ch['path']}",
                f"ملف مراقَب {('تغيّر' if ch['kind']=='modified' else ch['kind'])}: {ch['path']}",
                "Verify this change was intentional; restore from backup if not.",
                "تأكّد أن هذا التغيير مقصود؛ استعد من نسخة احتياطية إن لم يكن كذلك.", "fim"))
    except Exception:
        pass
    try:
        net = scan_network_threats()
        for h in net.get("hits", [])[:20]:
            findings.append(_f(f"NET-{abs(hash(h['detail']))%100000}", host, h["sev"],
                9.0 if h["sev"] == "critical" else 6.5,
                f"Network: {h['detail']}", f"الشبكة: {h['detail']}",
                "Investigate this connection; if malicious, terminate the process and block the IP.",
                "افحص هذا الاتصال؛ إن كان خبيثاً فأنهِ العملية واحظر العنوان.", "network"))
    except Exception:
        pass
    try:
        acct = scan_accounts()
        for it in acct.get("items", [])[:20]:
            findings.append(_f(f"ACCT-{abs(hash(it['detail']))%100000}", host, it["sev"], 6.5,
                f"Account: {it['detail']}", f"الحسابات: {it['detail']}",
                "Review this account; disable or remove it if unauthorised.",
                "راجع هذا الحساب؛ عطّله أو احذفه إن كان غير مصرّح.", "account"))
    except Exception:
        pass
    try:
        rk = scan_rootkit()
        for h in rk.get("hidden_processes", [])[:10]:
            findings.append(_f(f"RKPROC-{h['pid']}", host, "critical", 9.5,
                f"Hidden process detected: pid {h['pid']} ({h.get('name','?')}) visible in {','.join(h['seen_in'])} but hidden from {','.join(h['hidden_from'])}",
                f"عملية مخفية: pid {h['pid']} ({h.get('name','?')}) ظاهرة في {','.join(h['seen_in'])} لكنها مخفية عن {','.join(h['hidden_from'])}",
                "Strong rootkit indicator. Run an offline antivirus scan and consider the machine compromised until cleared.",
                "مؤشّر Rootkit قوي. شغّل فحصاً دون اتصال واعتبر الجهاز مخترقاً حتى التأكّد.", "rootkit"))
        for d in rk.get("unsigned_drivers", [])[:10]:
            findings.append(_f(f"RKDRV-{abs(hash(d['name']))%100000}", host, d["sev"], 9.0,
                f"Suspicious kernel driver '{d['name']}' ({d['signature']}): {d['path']}",
                f"درايفر نواة مشبوه '{d['name']}' ({d['signature']}): {d['path']}",
                "Unsigned/abnormal kernel drivers can be rootkits. Verify the driver's source; remove if untrusted.",
                "درايفرات النواة غير الموقّعة قد تكون rootkits. تحقّق من مصدره؛ أزله إن لم يكن موثوقاً.", "rootkit"))
    except Exception:
        pass

    # ---- FINAL DEDUP: one row per (cve,asset); keep the HIGHEST severity ----
    sev_rank = {"critical": 4, "high": 3, "medium": 2, "low": 1, "info": 0}
    best = {}
    for f in findings:
        k = (f["cve"], f["asset"])
        if k not in best or sev_rank.get(f["sev"], 0) > sev_rank.get(best[k]["sev"], 0):
            best[k] = f
    findings = list(best.values())
    _last_scan["count"] = len(findings)
    return findings


# ==================================================================
#  REMEDIATION ENGINE — real fixes, executed ONLY with explicit consent
#  For each finding category we define one or more remediation OPTIONS.
#  The UI shows the exact commands; the user picks one and approves; only
#  then is it executed (admin role + CSRF required). Every action is:
#    - reversible where possible (we also provide the undo command)
#    - logged to the audit trail
#    - verified afterwards by re-checking the condition
#  PRIVACY: all commands run locally; nothing leaves the machine.
# ==================================================================

def _remediation_plans(f):
    """Return a list of remediation options for a finding. Each option:
    {id, title_en, title_ar, risk, commands:[...], undo:[...], note_en, note_ar}
    risk: 'safe' | 'moderate' | 'careful'"""
    cve = f.get("cve", "")
    cat = f.get("cat", "")
    plans = []

    # ---- risky listening ports ----
    if cve.startswith("PORT-"):
        try:
            port = cve.split("-")[1]
        except Exception:
            port = ""
        if port:
            plans.append({
                "id": "fw-block-port", "risk": "safe",
                "title_en": f"Block inbound port {port} with Windows Firewall",
                "title_ar": f"حظر المنفذ الوارد {port} عبر جدار حماية Windows",
                "commands": [f'netsh advfirewall firewall add rule name="Sentinel-Port-{port}" dir=in action=block protocol=TCP localport={port}'],
                "undo": [f'netsh advfirewall firewall delete rule name="Sentinel-Port-{port}"'],
                "note_en": "Blocks new inbound connections to this port. Local programs keep working; remote access to this service stops.",
                "note_ar": "يحظر الاتصالات الواردة الجديدة لهذا المنفذ. البرامج المحلية تستمر؛ يتوقف الوصول البعيد للخدمة.",
            })
            if port == "445" or port == "139":
                plans.append({
                    "id": "disable-smb1", "risk": "moderate",
                    "title_en": "Disable the legacy SMBv1 protocol",
                    "title_ar": "تعطيل بروتوكول SMBv1 القديم",
                    "commands": ['powershell -NoProfile -Command "Disable-WindowsOptionalFeature -Online -FeatureName SMB1Protocol -NoRestart"'],
                    "undo": ['powershell -NoProfile -Command "Enable-WindowsOptionalFeature -Online -FeatureName SMB1Protocol -NoRestart"'],
                    "note_en": "SMBv1 is the protocol WannaCry exploited. Modern Windows shares use SMBv2/3 and keep working.",
                    "note_ar": "SMBv1 هو ما استغلّته WannaCry. مشاركات Windows الحديثة تستخدم SMBv2/3 وتستمر بالعمل.",
                })
            if port == "3389":
                plans.append({
                    "id": "rdp-nla", "risk": "safe",
                    "title_en": "Require Network Level Authentication for RDP",
                    "title_ar": "فرض مصادقة NLA لسطح المكتب البعيد",
                    "commands": ['reg add "HKLM\\SYSTEM\\CurrentControlSet\\Control\\Terminal Server\\WinStations\\RDP-Tcp" /v UserAuthentication /t REG_DWORD /d 1 /f'],
                    "undo": ['reg add "HKLM\\SYSTEM\\CurrentControlSet\\Control\\Terminal Server\\WinStations\\RDP-Tcp" /v UserAuthentication /t REG_DWORD /d 0 /f'],
                    "note_en": "Attackers must authenticate BEFORE a session is created. No impact on normal RDP use.",
                    "note_ar": "يجبر المهاجم على المصادقة قبل إنشاء الجلسة. لا تأثير على الاستخدام الطبيعي.",
                })

    # ---- live-detection findings ----
    if cve.startswith("PROC-"):
        try:
            pid = cve.split("-")[1]
        except Exception:
            pid = ""
        if pid:
            plans.append({
                "id": "kill-process", "risk": "careful",
                "title_en": f"Terminate the suspicious process (PID {pid})",
                "title_ar": f"إنهاء العملية المشبوهة (PID {pid})",
                "commands": [f"taskkill /F /PID {pid}"],
                "undo": [],
                "note_en": "Force-kills the process. Make sure it isn't a program you need — note its path first.",
                "note_ar": "ينهي العملية إجبارياً. تأكّد أنها ليست برنامجاً تحتاجه — دوّن مسارها أولاً.",
            })
    if cve.startswith("RKPROC-") or cve.startswith("RKDRV-"):
        plans.append({
            "id": "defender-offline", "risk": "careful",
            "title_en": "Run a Windows Defender OFFLINE scan (reboots to scan before Windows loads)",
            "title_ar": "تشغيل فحص Windows Defender دون اتصال (يعيد التشغيل ويفحص قبل تحميل ويندوز)",
            "commands": ['powershell -NoProfile -Command "Start-MpWDOScan"'],
            "undo": [],
            "note_en": "Defender Offline restarts the PC and scans from a trusted environment where rootkits can't hide. SAVE YOUR WORK FIRST — the machine reboots. Note: Defender's cloud submission settings still apply; review them if privacy is critical.",
            "note_ar": "يعيد التشغيل ويفحص من بيئة موثوقة لا تستطيع الـ rootkits الاختباء فيها. احفظ عملك أولاً — الجهاز سيعيد التشغيل. ملاحظة: إعدادات رفع العيّنات لدى Defender تبقى سارية؛ راجعها إن كانت الخصوصية حرجة.",
        })
        plans.append({
            "id": "list-drivers", "risk": "safe",
            "title_en": "Export the full list of loaded drivers for manual review",
            "title_ar": "تصدير قائمة الدرايفرات المحمّلة كاملةً للمراجعة اليدوية",
            "commands": ["driverquery /v /fo table"],
            "undo": [],
            "note_en": "Lists every loaded driver with its path and state so you can investigate the flagged one.",
            "note_ar": "يسرد كل درايفر محمّل مع مساره وحالته لتفحص المُعلَّم منها.",
        })

    if cve.startswith("FIM-"):
        plans.append({
            "id": "fim-rebaseline", "risk": "moderate",
            "title_en": "Accept current state as the new integrity baseline",
            "title_ar": "اعتماد الحالة الحالية كخط أساس جديد للسلامة",
            "commands": [],
            "undo": [],
            "note_en": "Only do this if you KNOW the change was legitimate (e.g. a Windows update). Otherwise investigate first.",
            "note_ar": "افعل هذا فقط إن كنت متأكّداً أن التغيير شرعي (مثل تحديث ويندوز). وإلا فافحص أولاً.",
        })
    if cve.startswith("PERSIST-"):
        plans.append({
            "id": "open-startup-apps", "risk": "safe",
            "title_en": "Open Startup Apps settings to review this entry",
            "title_ar": "فتح إعدادات تطبيقات بدء التشغيل لمراجعة هذا المدخل",
            "commands": ["start ms-settings:startupapps"],
            "undo": [],
            "note_en": "Shows programs that launch with Windows so you can disable anything you don't recognise. Safe — only opens Settings.",
            "note_ar": "يعرض البرامج التي تبدأ مع ويندوز لتعطّل ما لا تعرفه. آمن — يفتح الإعدادات فقط.",
        })
        plans.append({
            "id": "open-task-manager-startup", "risk": "safe",
            "title_en": "Open Task Manager to review startup items",
            "title_ar": "فتح مدير المهام لمراجعة عناصر بدء التشغيل",
            "commands": ["taskmgr"],
            "undo": [],
            "note_en": "Go to the Startup tab to see and disable suspicious auto-start programs.",
            "note_ar": "اذهب لتبويب بدء التشغيل لرؤية وتعطيل البرامج المشبوهة.",
        })

    if cve.startswith("ACCT-"):
        plans.append({
            "id": "open-users", "risk": "safe",
            "title_en": "Open Local Users & Groups to review accounts",
            "title_ar": "فتح المستخدمين والمجموعات المحليين لمراجعة الحسابات",
            "commands": ["start lusrmgr.msc"],
            "undo": [],
            "note_en": "Review the flagged account; disable or delete it if you don't recognise it.",
            "note_ar": "راجع الحساب المُعلَّم؛ عطّله أو احذفه إن كنت لا تعرفه.",
        })

    # ---- config findings (use the SCANNER's real finding ids) ----
    if cve == "FW-DISABLED":
        plans.append({
            "id": "fw-on", "risk": "safe",
            "title_en": "Turn ON Windows Firewall (all profiles)",
            "title_ar": "تشغيل جدار حماية Windows (كل الأوضاع)",
            "commands": ["netsh advfirewall set allprofiles state on"],
            "undo": ["netsh advfirewall set allprofiles state off"],
            "note_en": "Standard protection; existing allowed apps keep working.",
            "note_ar": "حماية قياسية؛ التطبيقات المسموحة تبقى كما هي.",
        })
    if cve == "AV-RTP-OFF":
        plans.append({
            "id": "defender-on", "risk": "safe",
            "title_en": "Enable Windows Defender real-time protection",
            "title_ar": "تفعيل الحماية اللحظية في Windows Defender",
            "commands": ['powershell -NoProfile -Command "Set-MpPreference -DisableRealtimeMonitoring $false"'],
            "undo": ['powershell -NoProfile -Command "Set-MpPreference -DisableRealtimeMonitoring $true"'],
            "note_en": "Re-enables live malware scanning.",
            "note_ar": "يعيد تفعيل فحص البرمجيات الخبيثة اللحظي.",
        })
    if cve == "SMBV1-ENABLED":
        plans.append({
            "id": "disable-smb1", "risk": "moderate",
            "title_en": "Disable the legacy SMBv1 protocol",
            "title_ar": "تعطيل بروتوكول SMBv1 القديم",
            "commands": ['powershell -NoProfile -Command "Disable-WindowsOptionalFeature -Online -FeatureName SMB1Protocol -NoRestart"'],
            "undo": ['powershell -NoProfile -Command "Enable-WindowsOptionalFeature -Online -FeatureName SMB1Protocol -NoRestart"'],
            "note_en": "Only ancient devices (Windows XP-era NAS) still need SMBv1.",
            "note_ar": "فقط الأجهزة القديمة جداً (عصر XP) ما تزال تحتاج SMBv1.",
        })
    if cve == "UAC-DISABLED":
        plans.append({
            "id": "uac-on", "risk": "safe",
            "title_en": "Re-enable User Account Control (UAC)",
            "title_ar": "إعادة تفعيل التحكم بحساب المستخدم (UAC)",
            "commands": ['reg add "HKLM\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Policies\\System" /v EnableLUA /t REG_DWORD /d 1 /f'],
            "undo": ['reg add "HKLM\\SOFTWARE\\Microsoft\\Windows\\CurrentVersion\\Policies\\System" /v EnableLUA /t REG_DWORD /d 0 /f'],
            "note_en": "Restores elevation prompts. Takes effect after a reboot.",
            "note_ar": "يعيد نوافذ طلب الصلاحيات. يسري بعد إعادة التشغيل.",
        })
    if cve == "GUEST-ACTIVE":
        plans.append({
            "id": "guest-off", "risk": "safe",
            "title_en": "Disable the built-in Guest account",
            "title_ar": "تعطيل حساب الضيف",
            "commands": ["net user guest /active:no"],
            "undo": ["net user guest /active:yes"],
            "note_en": "The Guest account is a classic foothold; disabling it is standard hardening.",
            "note_ar": "حساب الضيف موطئ قدم كلاسيكي للمهاجمين؛ تعطيله إجراء قياسي.",
        })
    if cve == "RDP-NO-NLA":
        plans.append({
            "id": "rdp-nla", "risk": "safe",
            "title_en": "Require Network Level Authentication for RDP",
            "title_ar": "فرض مصادقة NLA لسطح المكتب البعيد",
            "commands": ['reg add "HKLM\\SYSTEM\\CurrentControlSet\\Control\\Terminal Server\\WinStations\\RDP-Tcp" /v UserAuthentication /t REG_DWORD /d 1 /f'],
            "undo": ['reg add "HKLM\\SYSTEM\\CurrentControlSet\\Control\\Terminal Server\\WinStations\\RDP-Tcp" /v UserAuthentication /t REG_DWORD /d 0 /f'],
            "note_en": "Attackers must authenticate BEFORE a session is created.",
            "note_ar": "يجبر المهاجم على المصادقة قبل إنشاء الجلسة.",
        })
    if cve in ("PATCH-STALE", "OS-OUTDATED"):
        plans.append({
            "id": "windows-update", "risk": "safe",
            "title_en": "Open Windows Update to install pending patches",
            "title_ar": "فتح Windows Update لتثبيت التحديثات المعلّقة",
            "commands": ["start ms-settings:windowsupdate"],
            "undo": [],
            "note_en": "Opens the Update panel so you can review and install. (Auto-install needs the PSWindowsUpdate module.)",
            "note_ar": "يفتح لوحة التحديثات لتراجع وتثبّت. (التثبيت التلقائي يحتاج وحدة PSWindowsUpdate.)",
        })

    # ---- software / intel findings (updates) ----
    if cat in ("software", "intel"):
        # try to derive a winget id for common apps
        name_l = (f.get("title_en", "") + " " + f.get("cve", "")).lower()
        winget_map = [("firefox", "Mozilla.Firefox"), ("chrome", "Google.Chrome"),
                      ("7-zip", "7zip.7zip"), ("7zip", "7zip.7zip"), ("git", "Git.Git"),
                      ("vlc", "VideoLAN.VLC"), ("notepad++", "Notepad++.Notepad++"),
                      ("python", "Python.Python.3.12"), ("node", "OpenJS.NodeJS.LTS"),
                      ("edge", "Microsoft.Edge"), ("zoom", "Zoom.Zoom")]
        wid = next((w for k, w in winget_map if k in name_l), None)
        if wid:
            plans.append({
                "id": "winget-update", "risk": "safe",
                "title_en": f"Update via winget ({wid})",
                "title_ar": f"تحديث عبر winget ‏({wid})",
                "commands": [f"winget upgrade --id {wid} --silent --accept-package-agreements --accept-source-agreements"],
                "undo": [],
                "note_en": "Official package from the Windows Package Manager. The app may restart.",
                "note_ar": "حزمة رسمية من مدير حزم Windows. قد يُعاد تشغيل التطبيق.",
            })
        plans.append({
            "id": "winget-update-all", "risk": "moderate",
            "title_en": "Update ALL outdated apps via winget",
            "title_ar": "تحديث كل التطبيقات القديمة عبر winget",
            "commands": ["winget upgrade --all --silent --accept-package-agreements --accept-source-agreements"],
            "undo": [],
            "note_en": "Updates everything winget manages. Takes several minutes; running apps may restart.",
            "note_ar": "يحدّث كل ما يديره winget. يستغرق دقائق؛ قد تُعاد تشغيل التطبيقات المفتوحة.",
        })
        if cat == "software" and "eol" in cve.lower():
            plans.append({
                "id": "manual-remove", "risk": "careful",
                "title_en": "Uninstall the end-of-life software (manual)",
                "title_ar": "إزالة البرنامج منتهي الدعم (يدوي)",
                "commands": ["start ms-settings:appsfeatures"],
                "undo": [],
                "note_en": "Opens Settings > Apps so you can review and uninstall it yourself — safest for EOL software you might still need.",
                "note_ar": "يفتح الإعدادات > التطبيقات لتراجع وتزيل بنفسك — الأسلم للبرامج التي قد تحتاجها.",
            })
        # system/runtime products (DirectX, VC++, .NET) update via Windows Update
        name_lc = (f.get("title_en", "") + " " + f.get("cve", "")).lower()
        if cat == "intel" and any(v in name_lc for v in ("directx", "visual-studio", "visual studio", ".net", "dotnet", "redistributable", "webview")):
            plans.append({
                "id": "open-windows-update", "risk": "safe",
                "title_en": "Open Windows Update (recommended for system components)",
                "title_ar": "فتح Windows Update (المُوصى به لمكوّنات النظام)",
                "commands": ["start ms-settings:windowsupdate"],
                "undo": [],
                "note_en": "System runtimes like DirectX/.NET/VC++ update through Windows Update, not winget. This opens it so you can install pending updates.",
                "note_ar": "مكوّنات النظام مثل DirectX/.NET تُحدّث عبر Windows Update لا winget. هذا يفتحه لتثبّت التحديثات المعلّقة.",
            })

    return plans


def _remediation_verify(f):
    """Re-check the condition behind a finding after a fix. Returns True if resolved."""
    cve = f.get("cve", "")
    try:
        if cve.startswith("PORT-"):
            port = int(cve.split("-")[1])
            ports = _net_conn_stats().get("ports", [])
            return port not in ports
        if cve == "FW-DISABLED":
            out = _run_cmd(["netsh", "advfirewall", "show", "allprofiles", "state"], 10).lower()
            return "off" not in out.replace("turn off", "")
        if cve == "GUEST-ACTIVE":
            out = _run_cmd(["net", "user", "guest"], 10).lower()
            return "account active               no" in out or "no" in out.split("account active")[-1][:20]
        if cve == "AV-RTP-OFF":
            out = _run_cmd(["powershell", "-NoProfile", "-Command",
                            "(Get-MpPreference).DisableRealtimeMonitoring"], 15).strip().lower()
            return out == "false"
    except Exception:
        pass
    return False   # unknown -> let the next full scan decide


_remediation_log = collections.deque(maxlen=200)


def execute_remediation(cve, plan_id, user):
    """Run an APPROVED remediation plan. Logs everything; verifies afterwards."""
    f = next((x for x in _FINDINGS if x["cve"] == cve), None)
    if not f:
        # synthesize from the id (same as the options endpoint) so a plan still
        # resolves after a rescan rebuilt the findings store — prevents HTTP 500.
        cat = ""
        up = cve.upper()
        if up.startswith("KEV-"):
            cat = "intel"
        elif up.startswith(("PATCH", "OS-")):
            cat = "config"
        elif up.startswith("PORT-"):
            cat = "exposure"
        f = {"cve": cve, "asset": _plat.node() or "host", "sev": "medium",
             "score": 5.0, "fix_en": "", "fix_ar": "", "cat": cat,
             "title_en": cve, "title_ar": cve}
    plan = next((p for p in _remediation_plans(f) if p["id"] == plan_id), None)
    if not plan:
        return {"ok": False, "error": "plan not found"}
    results = []
    try:
        for cmd in plan["commands"]:
            if os.name == "nt":
                args = cmd if isinstance(cmd, list) else (
                    cmd.split() if cmd.startswith(("netsh", "net ", "reg ")) else ["cmd", "/c", cmd])
                out = _run_cmd(args, 180)
            else:
                out = "(not on Windows — dry run)"
            results.append({"cmd": cmd if isinstance(cmd, str) else " ".join(cmd),
                            "output": (out or "")[:500]})
    except Exception as e:
        return {"ok": False, "error": f"command failed: {e}", "results": results}
    try:
        verified = _remediation_verify(f) if os.name == "nt" else False
    except Exception:
        verified = False
    if verified:
        with _sec_lock:
            f["st"] = "patched"
    rec = {"ts": int(time.time() * 1000), "cve": cve, "plan": plan_id, "user": user,
           "verified": verified, "results": results, "undo": plan.get("undo", [])}
    _remediation_log.append(rec)
    audit("remediation_executed", f"{cve} via {plan_id} verified={verified}")
    log_event("remediation", sev="info", cve=cve, asset=f.get("asset", ""),
              text_en=f"Remediation applied: {plan['title_en']}" + (" (verified)" if verified else " (pending verification)"),
              text_ar=f"طُبّقت المعالجة: {plan['title_ar']}" + (" (تم التحقق)" if verified else " (بانتظار التحقق)"))
    return {"ok": True, "verified": verified, "results": results, "undo": plan.get("undo", [])}


def _apply_scan_results(findings):
    """Replace live findings with scan results, preserving prior remediated state."""
    with _sec_lock:
        prev = {f["cve"]: f["st"] for f in _FINDINGS}
        for f in findings:
            if prev.get(f["cve"]) in ("patched", "dismissed", "resolved"):
                f["st"] = prev[f["cve"]]      # keep what the analyst already actioned
        _FINDINGS[:] = findings
    _rebuild_recs(findings)   # recommendations now derive from REAL scan state
    for f in findings:
        if f["st"] == "open":
            log_event("finding_open", sev=f["sev"], cve=f["cve"], asset=f["asset"],
                      text_en=f.get("title_en", f["cve"]), text_ar=f.get("title_ar", f["cve"]))


def run_scan_now():
    if _last_scan.get("running"):
        return _last_scan
    _last_scan["running"] = True
    try:
        findings = _run_security_scan()
        _apply_scan_results(findings)
        log_event("scan", sev="info", asset=_plat.node() or "host",
                  text_en=f"System vulnerability scan completed: {len(findings)} findings",
                  text_ar=f"اكتمل فحص ثغرات النظام: {len(findings)} نتيجة")
    finally:
        _last_scan["running"] = False
    return _last_scan


_scan_thread = [None]


def _periodic_scan_loop():
    """Run the full vulnerability + live-detection scan automatically on a timer
    so the user never has to press "Scan". First pass runs at startup, then every
    SENTINEL_SCAN_MINUTES (default 60). Lightweight: skips if one is in flight."""
    try:
        mins = int(os.environ.get("SENTINEL_SCAN_MINUTES", "60") or "60")
    except Exception:
        mins = 60
    interval = max(300, mins * 60)   # never faster than every 5 min
    first = True
    while True:
        try:
            if not first:
                time.sleep(interval)
            first = False
            run_scan_now()
            # also refresh the live-detection engines so threats surface on their
            # own and get logged as events, without the user pressing anything.
            try:
                scan_processes(); scan_persistence(); fim_check()
                scan_network_threats(); scan_accounts(); scan_rootkit()
                # behavioral heuristics + optional auto-isolation of clear threats
                heur = _heuristic_scan()
                if heur.get("available"):
                    _auto_isolate_from_alerts(heur.get("alerts", []))
            except Exception:
                pass
        except Exception:
            time.sleep(interval)


def _ensure_initial_scan():
    if _scan_thread[0] is None:
        t = threading.Thread(target=_periodic_scan_loop, daemon=True)
        _scan_thread[0] = t
        t.start()


# ==================================================================
#  THREAT INTELLIGENCE  (privacy-first, always-updated)
#  Pulls the PUBLIC CISA "Known Exploited Vulnerabilities" feed — the
#  authoritative list of CVEs being actively exploited in the wild —
#  and matches it LOCALLY against the installed-software inventory.
#  PRIVACY: the ONLY network traffic is a one-way GET of a public feed.
#  Nothing about this machine is ever transmitted. Set SENTINEL_OFFLINE=1
#  to disable all network access entirely.
# ==================================================================
import urllib.request as _ur

_TI_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_threatintel.json")
def _load_json_file(path, default):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default


_TI = _load_json_file(_TI_FILE, {"kev": [], "ts": 0, "count": 0})
_TI_URL = os.environ.get(
    "SENTINEL_KEV_URL",
    "https://www.cisa.gov/sites/default/files/feeds/known_exploited_vulnerabilities.json")
_OFFLINE = os.environ.get("SENTINEL_OFFLINE", "").strip() not in ("", "0", "false", "False")
_ti_thread = [None]
_ti_lock = threading.Lock()


def _ti_stale():
    return (int(time.time() * 1000) - _TI.get("ts", 0)) > 12 * 3600 * 1000


def update_threat_intel(force=False):
    """Download the latest public exploited-vulnerability feed. Privacy: GET only,
    no data about this host is ever sent. No-op in offline mode."""
    if _OFFLINE:
        _TI["status"] = "offline"
        return _TI
    if not force and _TI.get("kev") and not _ti_stale():
        return _TI
    try:
        req = _ur.Request(_TI_URL, headers={"User-Agent": "Sentinel-SOC"})  # generic UA, no host data
        with _ur.urlopen(req, timeout=25) as r:
            data = json.loads(r.read().decode("utf-8", "replace"))
        vulns = data.get("vulnerabilities", []) if isinstance(data, dict) else []
        kev = []
        for v in vulns:
            kev.append({"cve": v.get("cveID", ""), "vendor": v.get("vendorProject", ""),
                        "product": v.get("product", ""), "name": v.get("vulnerabilityName", ""),
                        "added": v.get("dateAdded", ""), "due": v.get("dueDate", ""),
                        "ransomware": v.get("knownRansomwareCampaignUse", "")})
        with _ti_lock:
            _TI["kev"] = kev
            _TI["ts"] = int(time.time() * 1000)
            _TI["count"] = len(kev)
            _TI["status"] = "ok"
            _TI.pop("error", None)
        _save_json_file(_TI_FILE, _TI)
    except Exception as e:
        _TI["status"] = "error"
        _TI["error"] = str(e)[:200]
    return _TI


def _save_json_file(path, obj):
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(obj, f, ensure_ascii=False)
    except Exception:
        pass


def _ti_loop():
    update_threat_intel(force=not _TI.get("kev"))
    while True:
        time.sleep(6 * 3600)          # refresh every 6 hours
        update_threat_intel()


def _ensure_threat_intel():
    if _ti_thread[0] is None and not _OFFLINE:
        t = threading.Thread(target=_ti_loop, daemon=True)
        _ti_thread[0] = t
        t.start()


# Two-tier filtering against false positives:
#   UMBRELLA: if ANY token of the KEV product is one of these, skip the CVE entirely
#             (e.g., KEV "DirectX End-User Runtime" → contains "directx" → skip).
#   NOISE:    tokens that don't count as "specific" enough to identify a product
_KEV_UMBRELLA = {"windows", "office", "java", "python", "directx", "framework",
                 "linux", "android", "macos", "ios"}
_KEV_NOISE = {"server", "client", "system", "core", "service", "driver", "tools",
              "runtime", "mode", "edition", "package", "library", "redistributable",
              "language", "pack", "environment", "manager", "helper", "plugin",
              "end", "user", "net", "browser", "engine", "shared", "support",
              "professional", "standard", "express", "enterprise", "starter",
              "x86", "x64", "win", "win32", "win64", "amd64", "arm", "arm64"}

# Products with documented release years — used to filter physically-impossible matches
# (e.g. WebView2 didn't exist in 2018, so CVE-2018-xxxx against WebView2 is bogus).
_PRODUCT_FIRST_YEAR = {"webview2": 2020, "edge": 2015, "teams": 2017,
                       "asp net core": 2016, "windows 11": 2021}


def _kev_normalize(s):
    s = s.lower()
    for ch in "-_.,()[]/\\":
        s = s.replace(ch, " ")
    return " ".join(s.split())


_KEV_SKIP_EOL = ("silverlight", "internet explorer", "flash", "java 7", "java 8", "python 2")

# System/runtime products whose installed VERSION can't be reliably read from the
# registry display name (DirectX, VC++ redistributables, .NET, VS setup bits).
# We can't prove they're vulnerable, so we DON'T raise an alarming "vulnerability"
# — at most an informational "verify" note, and only when not clearly current.
_KEV_VERIFY_ONLY = ("directx", "visual studio", "visual c++", "vc redist",
                    "redistributable", ".net", "dotnet", "runtime", "webview")


def _version_year_estimate(name_norm, ver):
    """Best-effort estimate of the RELEASE YEAR of the installed version, for
    fast-moving apps. Lets us drop KEV CVEs that predate the installed build
    (i.e. the user already updated). Returns None when unknown."""
    if not ver:
        return None
    try:
        parts = [int(x) for x in str(ver).replace(",", ".").split(".")[:3] if x.strip().isdigit()]
    except Exception:
        return None
    if not parts:
        return None
    major = parts[0]
    minor = parts[1] if len(parts) > 1 else 0
    try:
        if "firefox" in name_norm:
            # Firefox 100 = May 2022, ~12 majors/year
            return 2022 + max(0, (major - 100)) // 12 if major >= 100 else 2004 + major // 8
        if "chrome" in name_norm or "chromium" in name_norm:
            # Chrome 100 = Mar 2022, ~13 majors/year (releases slowed to ~8/yr later — be generous)
            return 2022 + max(0, (major - 100)) // 8 if major >= 100 else 2008 + major // 8
        if "edge" in name_norm and major >= 100:
            return 2022 + max(0, (major - 100)) // 8
        if "7 zip" in name_norm or "7zip" in name_norm:
            # 7-Zip major = year-2000 (21.x=2021, 24.x=2024)
            return 2000 + major if 15 <= major <= 40 else None
        if "git" in name_norm and major == 2:
            # Git ~4 minors/year. 2.39=Dec2022, 2.43=2023, 2.47=2024, 2.49=2025.
            # Map generously so a current install is recognised as up-to-date.
            if minor >= 30:
                return 2022 + max(0, (minor - 39)) // 4 + (1 if minor >= 39 else 0)
            return None
        if "vlc" in name_norm and major == 3:
            return 2018 + minor // 4
        if "python" in name_norm and major == 3:
            # Python 3.11 = Oct 2022, one minor/year
            return 2022 + (minor - 11) if minor >= 8 else None
        if "node" in name_norm:
            # Node even majors yearly: 18=2022, 20=2023, 22=2024
            return 2022 + (major - 18) // 2 if major >= 14 else None
    except Exception:
        return None
    return None


def _match_kev_to_software():
    """LOCAL match between installed software and CISA's exploited-CVE feed.
    Strict matching to avoid false positives:
      - umbrella/noise token filters (Windows, Office, DirectX, …)
      - full-vendor check
      - temporal sanity vs product first-release year
      - VERSION-AWARE: skip CVEs that predate the installed build's estimated
        release year (an updated Firefox/Chrome/7-Zip/Git is NOT re-flagged)
      - ONE aggregated finding per software product (no duplicate rows with
        mixed severities for the same app)"""
    out = []
    kev = _TI.get("kev") or []
    if not kev:
        return out
    sw = _last_scan.get("software", [])
    if not sw:
        return out
    host = _plat.node() or "localhost"
    sw_norm = [(_kev_normalize(a["name"]), a.get("version", ""), a["name"]) for a in sw]
    # group matches per installed product: name -> {"cves": [...], "ransom": bool}
    per_product = {}
    for entry in kev:
        prod = _kev_normalize(entry.get("product") or "")
        vend = _kev_normalize(entry.get("vendor") or "")
        if not prod or not vend:
            continue
        prod_tokens = prod.split()
        # rule 1: ANY umbrella token poisons the whole product (DirectX, Office, …)
        if any(t in _KEV_UMBRELLA for t in prod_tokens):
            continue
        # rule 2: specific tokens = non-noise AND length >= 3 (so "git", "zip", "asp" count)
        specific = [t for t in prod_tokens if t not in _KEV_NOISE and len(t) >= 3]
        if not specific:
            continue
        cve_year = None
        try:
            cve_year = int(entry.get("cve", "").split("-")[1])
        except Exception:
            pass
        for name_norm, ver, name in sw_norm:
            # skip products already flagged as End-Of-Life (one EOL finding is enough)
            if any(eol in name_norm for eol in _KEV_SKIP_EOL):
                continue
            # require FULL vendor string in installed name (short vendors like "Git", "7-Zip")
            if vend not in name_norm:
                continue
            if not all(t in name_norm for t in specific):
                continue
            # rule 3: temporal sanity — skip CVE older than the product line itself
            if cve_year:
                skip = False
                for prod_key, first_year in _PRODUCT_FIRST_YEAR.items():
                    if prod_key in name_norm and cve_year < first_year:
                        skip = True
                        break
                if skip:
                    break
                # rule 4 (NEW): version-aware — installed build released AFTER the
                # CVE year means the user already updated past it; drop the match.
                vy = _version_year_estimate(name_norm, ver)
                if vy is not None and vy > cve_year:
                    break
            ransom = str(entry.get("ransomware", "")).lower().startswith("known")
            slot = per_product.setdefault(name, {"cves": [], "ransom": False, "names": []})
            if entry["cve"] not in slot["cves"]:
                slot["cves"].append(entry["cve"])
                slot["names"].append(entry.get("name", "")[:60])
                slot["ransom"] = slot["ransom"] or ransom
            break
    # one aggregated finding per product
    for name, info in per_product.items():
        cves = info["cves"]
        name_l = name.lower()
        verify_only = any(v in name_l for v in _KEV_VERIFY_ONLY)
        cve_list = ", ".join(cves[:6]) + (" …" if len(cves) > 6 else "")
        if verify_only:
            # System runtimes (DirectX, Visual Studio, .NET, VC++) update silently
            # through Windows Update and we CAN'T read their exact version. Showing
            # them as permanent "open findings" is pure noise — they can never be
            # cleared by the user and aren't confirmed vulnerabilities. So we SKIP
            # them entirely (do not add to findings). They remain covered by the
            # global CISA KEV reference list shown separately on the Security page.
            continue
        else:
            sev, score = ("high", 7.5) if info["ransom"] else ("medium", 5.0)
            out.append(_f(f"KEV-{_kev_normalize(name)[:30].replace(' ', '-')}", host, sev, score,
                          f"Possible exposure (verify version): '{name}' may be affected by {len(cves)} exploited CVE(s): {cve_list}",
                          f"تعرّض محتمل (تحقّق من الإصدار): قد يكون '{name}' عرضة لـ {len(cves)} ثغرة مستغلّة: {cve_list}",
                          f"Verify the installed version of {name} is current; if outdated, update it. CVEs: {cve_list} (CISA KEV).",
                          f"تأكّد أن إصدار {name} هو الأحدث؛ إن كان قديماً فحدّثه. الثغرات: {cve_list} (قائمة CISA).",
                          "intel"))
    return out


# ==================================================================
#  REAL INCIDENT DETECTION  (Windows Event Log)
#  Reads the actual Windows event logs to detect security incidents:
#  brute-force logon attempts, Defender malware detections, suspicious
#  PowerShell, audit-log clearing, and new service installs.
#  PRIVACY: reads LOCAL logs only; nothing leaves the machine.
# ==================================================================
_incident_thread = [None]
_incident_seen = set()          # de-dup by event RecordId
_BRUTE_THRESHOLD = 10           # failed logons within window => brute-force incident


def _ps_json(ps_script, timeout=25):
    """Run a PowerShell snippet that emits JSON; return parsed object or None."""
    if os.name != "nt":
        return None
    out = _run_cmd(["powershell", "-NoProfile", "-NonInteractive", "-Command", ps_script], timeout)
    out = out.strip()
    if not out:
        return None
    try:
        return json.loads(out)
    except Exception:
        return None


def _detect_incidents():
    """One pass over the Windows event logs. Logs real incidents via log_event."""
    if os.name != "nt":
        return []
    host = _plat.node() or "localhost"
    incidents = []

    # ---- 1) Failed logons in the last 30 min -> brute-force ----
    n = _ps_json(
        "$s=(Get-Date).AddMinutes(-30);"
        "$e=Get-WinEvent -FilterHashtable @{LogName='Security';Id=4625;StartTime=$s} "
        "-ErrorAction SilentlyContinue;"
        "@{count=($e|Measure-Object).Count} | ConvertTo-Json -Compress")
    if isinstance(n, dict) and (n.get("count") or 0) >= _BRUTE_THRESHOLD:
        c = int(n["count"])
        incidents.append(("brute", "danger", host,
                          f"Possible brute-force: {c} failed logon attempts in 30 min",
                          f"محاولة تخمين محتملة: {c} محاولة دخول فاشلة خلال ٣٠ دقيقة"))

    # ---- 2) Windows Defender malware detections (last 24h) ----
    det = _ps_json(
        "$s=(Get-Date).AddHours(-24);"
        "$e=Get-WinEvent -FilterHashtable @{LogName='Microsoft-Windows-Windows Defender/Operational';"
        "Id=1116,1117;StartTime=$s} -ErrorAction SilentlyContinue |"
        "Select-Object -First 5;"
        "$e | ForEach-Object {{ @{{id=$_.Id; msg=$_.Message}} }} | ConvertTo-Json -Compress")
    if det:
        items = det if isinstance(det, list) else [det]
        for it in items:
            msg = (it.get("msg") or "Malware detected")[:120]
            incidents.append(("malware", "danger", host,
                              f"Windows Defender threat detected: {msg}",
                              f"اكتشف Defender تهديداً: {msg}"))

    # ---- 3) Security audit log cleared (anti-forensics) ----
    cl = _ps_json(
        "$s=(Get-Date).AddHours(-24);"
        "$e=Get-WinEvent -FilterHashtable @{LogName='Security';Id=1102;StartTime=$s} "
        "-ErrorAction SilentlyContinue;"
        "@{count=($e|Measure-Object).Count} | ConvertTo-Json -Compress")
    if isinstance(cl, dict) and (cl.get("count") or 0) > 0:
        incidents.append(("logclear", "danger", host,
                          "Security audit log was cleared (possible anti-forensics)",
                          "تم مسح سجلّ تدقيق الأمان (تغطية محتملة على هجوم)"))

    # ---- 4) New service installed (common persistence, Id 7045) ----
    sv = _ps_json(
        "$s=(Get-Date).AddHours(-24);"
        "$e=Get-WinEvent -FilterHashtable @{LogName='System';Id=7045;StartTime=$s} "
        "-ErrorAction SilentlyContinue | Select-Object -First 3;"
        "$e | ForEach-Object {{ @{{msg=$_.Message}} }} | ConvertTo-Json -Compress")
    if sv:
        items = sv if isinstance(sv, list) else [sv]
        for it in items:
            msg = (it.get("msg") or "New service")[:100]
            incidents.append(("service", "warning", host,
                              f"New service installed: {msg}",
                              f"تم تثبيت خدمة جديدة: {msg}"))

    # ---- 5) Suspicious PowerShell: handled by the real-time risk-scoring engine
    # (_rt_inspect_process), which respects the allowlist, trusted parents, and
    # requires multiple signals. The old log-scraper here produced duplicate,
    # context-free alerts every cycle (a bare "-enc" is common in legit software),
    # so it has been removed in favour of the smarter engine. ----

    # log + alert each new incident
    for kind, sev, asset, en, ar in incidents:
        log_event("incident", sev=sev, asset=asset, text_en=en, text_ar=ar)
        try:
            add_notification("alert", sev, en, ar, "", "", route="alerts")
        except Exception:
            pass
        try:
            _dispatch_alert(en, f"{asset}: {en}", "critical" if sev == "danger" else "high")
        except Exception:
            pass
    return incidents


def _sysmon_available():
    """Check if Sysmon is installed (its operational log exists). Sysmon gives
    rich ETW-sourced events (process create w/ cmdline+hashes, network, file)
    via the Windows event log — far better than polling, no driver to write."""
    if os.name != "nt":
        return False
    out = _ps_json(
        "$l=Get-WinEvent -ListLog 'Microsoft-Windows-Sysmon/Operational' "
        "-ErrorAction SilentlyContinue; @{ok=[bool]$l} | ConvertTo-Json -Compress")
    return bool(isinstance(out, dict) and out.get("ok"))


_sysmon_last_record = [0]


def _sysmon_scan_once():
    """Read NEW Sysmon events since last pass and turn the high-signal ones into
    real-time detections. Sysmon event IDs: 1=process create, 8=CreateRemoteThread
    (injection), 25=process tampering. ETW-grade telemetry from the event log."""
    if os.name != "nt":
        return
    js = _ps_json(
        "$e=Get-WinEvent -FilterHashtable @{LogName='Microsoft-Windows-Sysmon/Operational';"
        "Id=1,8,25} -MaxEvents 40 -ErrorAction SilentlyContinue;"
        "$e | ForEach-Object { @{id=$_.Id; rec=$_.RecordId; msg=$_.Message} } | ConvertTo-Json -Compress")
    if not js:
        return
    items = js if isinstance(js, list) else [js]
    new_max = _sysmon_last_record[0]
    for it in items:
        try:
            rec = int(it.get("rec") or 0)
        except Exception:
            rec = 0
        if rec <= _sysmon_last_record[0]:
            continue
        new_max = max(new_max, rec)
        eid = it.get("id")
        low = (it.get("msg") or "").lower()
        # require BOTH an encode/hide primitive AND a download/exec primitive so
        # we don't flag the many legit apps that use a bare -enc / hidden window.
        _dl = any(k in low for k in ("frombase64string", "downloadstring", "downloadfile",
                                     "iex ", "invoke-expression", "certutil -urlcache",
                                     "bitsadmin /transfer", "webclient", "invoke-webrequest"))
        _enc = any(k in low for k in ("-enc", "-encodedcommand", "frombase64string"))
        _hid = "-windowstyle hidden" in low or "-w hidden" in low
        if eid == 1 and _dl and (_enc or _hid):
            _rt_log("critical",
                    "Sysmon: process created with encoded download/exec command",
                    "Sysmon: عملية أُنشئت بأمر تنزيل/تنفيذ مشفّر",
                    reason="sysmon encoded download process-create")
        elif eid == 8:
            _rt_log("high",
                    "Sysmon: remote thread injection detected (CreateRemoteThread)",
                    "Sysmon: حقن خيط عن بُعد (CreateRemoteThread)",
                    reason="sysmon remote-thread injection")
        elif eid == 25:
            _rt_log("critical",
                    "Sysmon: process tampering detected (hollowing/herpaderping)",
                    "Sysmon: تلاعب بعملية (process hollowing)",
                    reason="sysmon process tampering")
    _sysmon_last_record[0] = new_max


def _sysmon_loop():
    """Drain the Sysmon log every 5s. Sysmon itself is the efficient ETW consumer
    (kernel-sourced); we just read its already-collected events quickly."""
    try:
        js = _ps_json(
            "$e=Get-WinEvent -FilterHashtable @{LogName='Microsoft-Windows-Sysmon/Operational'} "
            "-MaxEvents 1 -ErrorAction SilentlyContinue;"
            "@{rec=($e.RecordId)} | ConvertTo-Json -Compress")
        if isinstance(js, dict) and js.get("rec"):
            _sysmon_last_record[0] = int(js["rec"])
    except Exception:
        pass
    while _realtime_enabled[0]:
        try:
            _sysmon_scan_once()
        except Exception:
            pass
        time.sleep(5)


def _incident_loop():
    while True:
        try:
            _detect_incidents()
        except Exception:
            pass
        time.sleep(180)               # check the event logs every 3 minutes


def _ensure_incident_detection():
    if _incident_thread[0] is None and os.name == "nt":
        t = threading.Thread(target=_incident_loop, daemon=True)
        _incident_thread[0] = t
        t.start()


# ==================================================================
#  EXTERNAL ALERTING  (Email / Telegram / Webhook)
#  Sends alerts ONLY to channels the user configures (their own SMTP,
#  their own Telegram bot, their own webhook URL). The only thing sent
#  is the alert text the user opted into — never system data otherwise.
# ==================================================================
_SETTINGS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_settings.json")
_SETTINGS = _secure_load(_SETTINGS_FILE, {
    "alerts": {"enabled": False, "min_sev": "high",
               "email": {"enabled": False, "host": "", "port": 587, "tls": True,
                         "user": "", "password": "", "to": ""},
               "telegram": {"enabled": False, "token": "", "chat_id": ""},
               "webhook": {"enabled": False, "url": ""}}})
_settings_lock = threading.Lock()
_SEV_ORDER = {"low": 1, "medium": 2, "high": 3, "critical": 4, "danger": 4, "info": 0}


def _save_settings():
    with _settings_lock:
        _secure_save(_SETTINGS_FILE, _SETTINGS)


def _alerts_cfg():
    return _SETTINGS.get("alerts", {})


def _send_email(cfg, subject, body):
    import smtplib
    import ssl as _ssl
    from email.mime.text import MIMEText
    host = cfg.get("host", "").strip()
    user = cfg.get("user", "").strip()
    to = cfg.get("to", "").strip()
    # clear, actionable validation BEFORE attempting a connection (the raw
    # smtplib error "please run connect() first" is confusing to users).
    missing = []
    if not host:
        missing.append("SMTP host (e.g. smtp.gmail.com)")
    if not to:
        missing.append("recipient address (To)")
    if missing:
        raise ValueError("email not configured — fill in: " + ", ".join(missing) +
                         ". Open Settings > Alerts > Email.")
    msg = MIMEText(body, "plain", "utf-8")
    msg["Subject"] = subject
    msg["From"] = user or "sentinel@localhost"
    msg["To"] = to
    port = int(cfg.get("port", 587) or 587)
    ctx = _ssl.create_default_context()
    if port == 465:
        # SMTP_SSL — wraps the connection in SSL from the start
        with smtplib.SMTP_SSL(host, port, timeout=15, context=ctx) as s:
            if cfg.get("user"):
                s.login(cfg["user"], cfg.get("password", ""))
            s.sendmail(msg["From"], [a.strip() for a in cfg.get("to", "").split(",") if a.strip()], msg.as_string())
    else:
        # Plain SMTP then STARTTLS (ports 25, 587…)
        with smtplib.SMTP(host, port, timeout=15) as s:
            s.ehlo()
            if cfg.get("tls", True):
                s.starttls(context=ctx)
                s.ehlo()
            if cfg.get("user"):
                s.login(cfg["user"], cfg.get("password", ""))
            s.sendmail(msg["From"], [a.strip() for a in cfg.get("to", "").split(",") if a.strip()], msg.as_string())


def _send_telegram(cfg, text):
    import ssl as _ssl
    token, chat = cfg.get("token", ""), cfg.get("chat_id", "")
    if not token or not chat:
        raise ValueError("telegram not configured")
    url = f"https://api.telegram.org/bot{token}/sendMessage"
    data = json.dumps({"chat_id": chat, "text": text}).encode("utf-8")
    req = _ur.Request(url, data=data, headers={"Content-Type": "application/json"})
    # use a permissive SSL context to handle self-signed / corporate proxy certs
    ctx = _ssl.create_default_context()
    try:
        _ur.urlopen(req, timeout=15, context=ctx).read()
    except Exception:
        # last resort: skip certificate verification (corporate proxies / old Windows cert store)
        ctx2 = _ssl.create_default_context()
        ctx2.check_hostname = False
        ctx2.verify_mode = _ssl.CERT_NONE
        _ur.urlopen(req, timeout=15, context=ctx2).read()


def _send_webhook(cfg, title, text, sev):
    import ssl as _ssl
    url = cfg.get("url", "")
    if not url:
        raise ValueError("webhook not configured")
    payload = json.dumps({"text": f"*{title}*\n{text}", "title": title,
                          "message": text, "severity": sev,
                          "source": "Sentinel SOC"}).encode("utf-8")
    req = _ur.Request(url, data=payload, headers={"Content-Type": "application/json"})
    ctx = _ssl.create_default_context()
    try:
        _ur.urlopen(req, timeout=15, context=ctx).read()
    except Exception:
        ctx2 = _ssl.create_default_context()
        ctx2.check_hostname = False
        ctx2.verify_mode = _ssl.CERT_NONE
        _ur.urlopen(req, timeout=15, context=ctx2).read()


def _dispatch_alert(title, body, sev="high"):
    """Fan out an alert to all enabled channels (non-blocking, best-effort)."""
    cfg = _alerts_cfg()
    if not cfg.get("enabled"):
        return {"sent": [], "skipped": "alerts disabled"}
    if _SEV_ORDER.get(sev, 0) < _SEV_ORDER.get(cfg.get("min_sev", "high"), 3):
        return {"sent": [], "skipped": "below min severity"}

    def worker():
        results = {}
        e = cfg.get("email", {})
        if e.get("enabled"):
            try:
                _send_email(e, f"[Sentinel] {title}", body); results["email"] = "ok"
            except Exception as ex:
                results["email"] = f"error: {str(ex)[:80]}"
        tg = cfg.get("telegram", {})
        if tg.get("enabled"):
            try:
                _send_telegram(tg, f"🛡 {title}\n{body}"); results["telegram"] = "ok"
            except Exception as ex:
                results["telegram"] = f"error: {str(ex)[:80]}"
        wh = cfg.get("webhook", {})
        if wh.get("enabled"):
            try:
                _send_webhook(wh, title, body, sev); results["webhook"] = "ok"
            except Exception as ex:
                results["webhook"] = f"error: {str(ex)[:80]}"
        try:
            audit("alert_dispatched", f"{title} -> {results}", user="system")
        except Exception:
            pass
    threading.Thread(target=worker, daemon=True).start()
    return {"sent": "dispatching"}


def _dispatch_alert_sync(title, body, sev="high"):
    """Synchronous test dispatch — returns per-channel results."""
    cfg = _alerts_cfg()
    results = {}
    e = cfg.get("email", {})
    if e.get("enabled"):
        try:
            _send_email(e, f"[Sentinel] {title}", body); results["email"] = "ok"
        except Exception as ex:
            results["email"] = f"error: {str(ex)[:120]}"
    tg = cfg.get("telegram", {})
    if tg.get("enabled"):
        try:
            _send_telegram(tg, f"🛡 {title}\n{body}"); results["telegram"] = "ok"
        except Exception as ex:
            results["telegram"] = f"error: {str(ex)[:120]}"
    wh = cfg.get("webhook", {})
    if wh.get("enabled"):
        try:
            _send_webhook(wh, title, body, sev); results["webhook"] = "ok"
        except Exception as ex:
            results["webhook"] = f"error: {str(ex)[:120]}"
    return results


def active_threats():
    # Active threats = currently OPEN critical/high vulnerability findings.
    # (Windows Defender HISTORICAL detections are kept separate as info — they
    #  were inflating this and contradicting a 100 score with 0 open findings.)
    with _sec_lock:
        return sum(1 for f in _FINDINGS if f["sev"] in ("critical", "high") and f["st"] in _OPEN_STATES)


def posture_score():
    with _sec_lock:
        penalty = sum(_SEV_WEIGHT.get(f["sev"], 0) for f in _FINDINGS if f["st"] in _OPEN_STATES)
    return max(0, min(100, 100 - penalty))


def _counts():
    with _sec_lock:
        c = {"critical": 0, "high": 0, "medium": 0, "low": 0}
        for f in _FINDINGS:
            if f["st"] in _OPEN_STATES and f["sev"] in c:
                c[f["sev"]] += 1
    return [{"k": k, "n": v} for k, v in c.items()]


def _findings_out(lang):
    with _sec_lock:
        return [{"cve": f["cve"], "asset": f["asset"], "sev": f["sev"], "score": f["score"],
                 "st": f["st"], "fix": f["fix_ar" if lang == "ar" else "fix_en"],
                 "title": f.get("title_ar" if lang == "ar" else "title_en", ""),
                 "first_seen": f.get("first_seen")} for f in _FINDINGS]


def _recs_out(lang):
    with _sec_lock:
        return [{"id": r["id"], "state": r["state"], "cve": r.get("cve"),
                 "text": r["ar" if lang == "ar" else "en"]} for r in _RECS]


# ==================================================================
#  EVENT BUS + NOTIFICATIONS  (drives SSE + the bell dropdown)
# ==================================================================
_subscribers = []          # list of per-connection Queue
_sub_lock = threading.Lock()
_NOTIFS = collections.deque(maxlen=40)
_notif_seq = [0]


def _push_event(payload):
    """Fan out an event to every connected SSE client."""
    with _sub_lock:
        for q in list(_subscribers):
            try:
                q.append(payload)
            except Exception:
                pass


def add_notification(ic, tone, en, ar, sub_en="", sub_ar="", route="alerts", report_id=None, dedupe_sec=600):
    # collapse repeated identical notifications: if the same English text was
    # pushed within dedupe_sec, bump the existing one instead of stacking copies.
    now = int(time.time() * 1000)
    if dedupe_sec:
        cutoff = now - dedupe_sec * 1000
        for existing in _NOTIFS:
            if existing.get("en") == en and existing.get("ts", 0) >= cutoff:
                existing["ts"] = now
                existing["count"] = existing.get("count", 1) + 1
                _push_event({"type": "notification", "notification": existing})
                return existing
    _notif_seq[0] += 1
    n = {"id": _notif_seq[0], "ic": ic, "tone": tone,
         "en": en, "ar": ar, "sub_en": sub_en, "sub_ar": sub_ar,
         "route": route, "report_id": report_id,
         "ts": now}
    _NOTIFS.appendleft(n)
    _push_event({"type": "notification", "notification": n})
    return n


def _notifs_out(lang):
    ar = lang == "ar"
    return [{"id": n["id"], "ic": n["ic"], "tone": n["tone"],
             "t": n["ar"] if ar else n["en"], "s": n["sub_ar"] if ar else n["sub_en"],
             "route": n.get("route", "alerts"), "report_id": n.get("report_id"),
             "ts": n["ts"]} for n in _NOTIFS]


# Notifications start empty — real ones are added as actual events happen
# (scans, incidents, blocks). No seeded fake alerts about non-existent hosts.


# ==================================================================
#  REPORTS  (real, persisted, period-aware; daily/weekly/monthly/yearly)
# ==================================================================
_REPORTS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_reports.json")
_EVENTS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_events.json")
_reports_lock = threading.Lock()
_rep_seq = [0]
_PERIODS = ("daily", "weekly", "monthly", "yearly")
_PERIOD_LABEL = {
    "daily": ("Daily System Report", "تقرير النظام اليومي"),
    "weekly": ("Weekly Security Report", "تقرير الأمن الأسبوعي"),
    "monthly": ("Monthly Compliance Report", "تقرير الامتثال الشهري"),
    "yearly": ("Annual Security Review", "المراجعة الأمنية السنوية"),
    "status": ("Status Report", "تقرير الحالة"),
    "surprise": ("Incident Alert — New Vulnerability", "تنبيه حادثة — ثغرة جديدة"),
    "scan": ("File Scan Report", "تقرير فحص ملف"),
}
_PERIOD_SECONDS = {"daily": 86400, "weekly": 604800, "monthly": 2592000, "yearly": 31536000}


def _load_json(path, default):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return default




def _load_reports():
    data = _secure_load(_REPORTS_FILE, {})
    od = collections.OrderedDict()
    # keep chronological order by ts
    for rid, rep in sorted(data.items(), key=lambda kv: kv[1].get("ts", 0)):
        od[rid] = rep
    # restore the id sequence so new ids don't collide
    mx = 0
    for rid in od:
        try:
            mx = max(mx, int(str(rid).rsplit("-", 1)[-1]))
        except Exception:
            pass
    _rep_seq[0] = mx
    return od


_REPORTS = _load_reports()          # id -> report  (persisted, real only — no seeding)


def _save_reports():
    try:
        with _reports_lock:
            _secure_save(_REPORTS_FILE, dict(_REPORTS))
    except Exception:
        pass


# ---- lightweight security EVENT LOG (real activity for period reports) ----
_EVENTS = _secure_load(_EVENTS_FILE, [])   # [{ts, kind, sev, cve, asset, text_en, text_ar}]
_events_lock = threading.Lock()


# ==================================================================
#  In-memory SQLite INDEX of events (and audit) for fast queries.
#  PRIVACY: the DB lives in :memory: and never touches disk. The
#  encrypted JSON on disk remains the single source of truth.
# ==================================================================
import sqlite3 as _sqlite

_db = _sqlite.connect(":memory:", check_same_thread=False)
_db_lock = threading.Lock()


def _db_init():
    with _db_lock:
        cur = _db.cursor()
        cur.execute("""CREATE TABLE IF NOT EXISTS events (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            ts INTEGER NOT NULL, kind TEXT, sev TEXT, cve TEXT, asset TEXT,
            text_en TEXT, text_ar TEXT
        )""")
        cur.execute("CREATE INDEX IF NOT EXISTS ix_events_ts ON events(ts)")
        cur.execute("CREATE INDEX IF NOT EXISTS ix_events_kind ON events(kind)")
        cur.execute("CREATE INDEX IF NOT EXISTS ix_events_sev ON events(sev)")
        cur.execute("CREATE INDEX IF NOT EXISTS ix_events_cve ON events(cve)")
        # FTS5 if available for fast text search; fall back to LIKE if not.
        # (use a full-content FTS5 table — contentless FTS5 forbids DELETE,
        #  which would break retention/test cleanup.)
        try:
            cur.execute("""CREATE VIRTUAL TABLE IF NOT EXISTS events_fts
                           USING fts5(text_en, text_ar)""")
            _db_init.has_fts = True
        except _sqlite.OperationalError:
            _db_init.has_fts = False
        _db.commit()


_db_init.has_fts = False
_db_init()


def _db_sync_from_events():
    """Rebuild the SQLite index from the in-memory _EVENTS list. Called at startup
    and any time the source list is replaced (e.g. by retention pruning)."""
    with _db_lock:
        cur = _db.cursor()
        cur.execute("DELETE FROM events")
        if _db_init.has_fts:
            cur.execute("DELETE FROM events_fts")
        rows = [(e.get("ts", 0), e.get("kind", ""), e.get("sev", ""), e.get("cve", ""),
                 e.get("asset", ""), e.get("text_en", ""), e.get("text_ar", "")) for e in _EVENTS]
        cur.executemany("INSERT INTO events(ts,kind,sev,cve,asset,text_en,text_ar) VALUES (?,?,?,?,?,?,?)", rows)
        if _db_init.has_fts and rows:
            cur.executemany("INSERT INTO events_fts(rowid, text_en, text_ar) "
                            "SELECT id, text_en, text_ar FROM events WHERE id=?",
                            [(i + 1,) for i in range(len(rows))])
        _db.commit()


_db_sync_from_events()


def log_event(kind, sev="info", cve="", asset="", text_en="", text_ar="", dedupe_sec=600):
    """Record a real security event. Persists to the encrypted JSON AND indexes
    it in the in-memory SQLite for fast searching.

    dedupe_sec: if an identical event (same kind+cve+text) was logged within this
    many seconds, skip it. This stops the periodic scan from spamming the activity
    feed with the same "WorkloadManager" / FIM line every cycle."""
    now = int(time.time() * 1000)
    with _events_lock:
        if dedupe_sec and _EVENTS:
            cutoff = now - dedupe_sec * 1000
            for prev in reversed(_EVENTS):
                if prev["ts"] < cutoff:
                    break
                if (prev["kind"] == kind and prev.get("cve", "") == cve
                        and prev.get("text_en", "") == text_en
                        and prev.get("text_ar", "") == text_ar):
                    return prev   # duplicate within window -> don't re-log
    ev = {"ts": now, "kind": kind, "sev": sev,
          "cve": cve, "asset": asset, "text_en": text_en, "text_ar": text_ar}
    with _events_lock:
        _EVENTS.append(ev)
        if len(_EVENTS) > 5000:
            del _EVENTS[:len(_EVENTS) - 5000]
        try:
            _secure_save(_EVENTS_FILE, _EVENTS)
        except Exception:
            pass
    # mirror into the search index (best-effort, never blocks the caller)
    try:
        with _db_lock:
            cur = _db.cursor()
            cur.execute("INSERT INTO events(ts,kind,sev,cve,asset,text_en,text_ar) VALUES (?,?,?,?,?,?,?)",
                        (ev["ts"], ev["kind"], ev["sev"], ev["cve"], ev["asset"], ev["text_en"], ev["text_ar"]))
            if _db_init.has_fts:
                cur.execute("INSERT INTO events_fts(rowid, text_en, text_ar) VALUES (?,?,?)",
                            (cur.lastrowid, ev["text_en"], ev["text_ar"]))
            _db.commit()
    except Exception:
        pass
    return ev


def events_in_window(seconds):
    """Indexed range query — O(log N) via the ts index instead of full scan."""
    cutoff = int(time.time() * 1000) - seconds * 1000
    try:
        with _db_lock:
            rows = _db.execute("SELECT ts,kind,sev,cve,asset,text_en,text_ar FROM events "
                               "WHERE ts >= ? ORDER BY ts DESC", (cutoff,)).fetchall()
        return [{"ts": r[0], "kind": r[1], "sev": r[2], "cve": r[3], "asset": r[4],
                 "text_en": r[5], "text_ar": r[6]} for r in rows]
    except Exception:
        # fall back to the in-memory list if SQLite is unhappy
        with _events_lock:
            return [e for e in _EVENTS if e["ts"] >= cutoff]


def events_query(kind=None, sev=None, cve=None, search=None, since_ms=None, limit=500):
    """Fast filtered search using SQLite indexes + optional FTS5 full-text search."""
    where = []; args = []
    if kind:    where.append("kind = ?"); args.append(kind)
    if sev:     where.append("sev = ?"); args.append(sev)
    if cve:     where.append("cve = ?"); args.append(cve)
    if since_ms: where.append("ts >= ?"); args.append(int(since_ms))
    sql = "SELECT ts,kind,sev,cve,asset,text_en,text_ar FROM events"
    if search and _db_init.has_fts:
        sql += " WHERE id IN (SELECT rowid FROM events_fts WHERE events_fts MATCH ?)"
        args.insert(0, search)
        if where:
            sql += " AND " + " AND ".join(where)
    elif search:
        where.append("(text_en LIKE ? OR text_ar LIKE ? OR asset LIKE ? OR cve LIKE ?)")
        like = "%" + search + "%"
        args.extend([like, like, like, like])
        if where: sql += " WHERE " + " AND ".join(where)
    elif where:
        sql += " WHERE " + " AND ".join(where)
    sql += " ORDER BY ts DESC LIMIT ?"
    args.append(int(limit))
    try:
        with _db_lock:
            rows = _db.execute(sql, args).fetchall()
        return [{"ts": r[0], "kind": r[1], "sev": r[2], "cve": r[3], "asset": r[4],
                 "text_en": r[5], "text_ar": r[6]} for r in rows]
    except Exception:
        return []


def _avg(seq):
    seq = [x for x in seq if isinstance(x, (int, float))]
    return round(sum(seq) / len(seq), 1) if seq else 0


def generate_report(period, lang, extra=None, template=""):
    _rep_seq[0] += 1
    rid = f"rep-{period}-{_rep_seq[0]}"
    ar = lang == "ar"
    score = posture_score()
    threats = active_threats()
    counts = {c["k"]: c["n"] for c in _counts()}
    cpu_avg, ram_avg, net_avg = _avg(_series("cpu")), _avg(_series("ram")), _avg(_series("traffic"))
    title = _PERIOD_LABEL.get(period, _PERIOD_LABEL["daily"])[1 if ar else 0]
    tone = "danger" if score < 60 else ("warning" if score < 80 else "success")

    if ar:
        summary = (f"خلال هذه الفترة بلغت درجة الأمان {score}/100 مع {threats} تهديد نشط. "
                   f"متوسط المعالج {cpu_avg}% والذاكرة {ram_avg}GB وحركة الشبكة {net_avg} م.ب/ث. "
                   f"الثغرات المفتوحة: {counts.get('critical',0)} حرجة، {counts.get('high',0)} عالية، "
                   f"{counts.get('medium',0)} متوسطة.")
        sections = [
            {"h": "وضع المخاطر", "b": f"درجة الأمان {score}/100، {threats} تهديد نشط يتطلب المتابعة."},
            {"h": "أداء الموارد", "b": f"المعالج {cpu_avg}% · الذاكرة {ram_avg}GB · الشبكة {net_avg} م.ب/ث (متوسطات الفترة)."},
            {"h": "الثغرات", "b": f"{counts.get('critical',0)} حرجة، {counts.get('high',0)} عالية، {counts.get('medium',0)} متوسطة، {counts.get('low',0)} منخفضة."},
            {"h": "التوصيات", "b": "؛ ".join(r["ar"] for r in _RECS if r["state"] == "open") or "لا توصيات معلّقة."},
        ]
    else:
        summary = (f"Over this period the security score was {score}/100 with {threats} active threats. "
                   f"Average CPU {cpu_avg}%, memory {ram_avg}GB, network {net_avg} Mbps. "
                   f"Open findings: {counts.get('critical',0)} critical, {counts.get('high',0)} high, "
                   f"{counts.get('medium',0)} medium.")
        sections = [
            {"h": "Risk posture", "b": f"Security score {score}/100, {threats} active threats requiring attention."},
            {"h": "Resource performance", "b": f"CPU {cpu_avg}% · memory {ram_avg}GB · network {net_avg} Mbps (period averages)."},
            {"h": "Findings", "b": f"{counts.get('critical',0)} critical, {counts.get('high',0)} high, {counts.get('medium',0)} medium, {counts.get('low',0)} low."},
            {"h": "Recommendations", "b": "; ".join(r["en"] for r in _RECS if r["state"] == "open") or "No pending recommendations."},
        ]

    if extra:
        sections.insert(0, extra)

    # ---- real activity that happened DURING the period (from the event log) ----
    win = _PERIOD_SECONDS.get(period)
    if win:
        evs = events_in_window(win)
        opened = [e for e in evs if e["kind"] == "finding_open"]
        fixed = [e for e in evs if e["kind"] in ("finding_resolved", "finding_patched")]
        scans = [e for e in evs if e["kind"] == "scan"]
        incidents = [e for e in evs if e["kind"] in ("incident", "breach", "leak")]
        if ar:
            sections.append({"h": "نشاط الفترة", "b":
                f"ثغرات جديدة مكتشفة: {len(opened)} · ثغرات تمت معالجتها: {len(fixed)} · "
                f"عمليات فحص: {len(scans)} · حوادث/اختراقات/تسريبات: {len(incidents)}."})
            if incidents:
                sections.append({"h": "الحوادث الأمنية", "b": "؛ ".join(
                    (e.get("text_ar") or e.get("text_en") or e.get("cve") or "حادثة") for e in incidents[:8])})
        else:
            sections.append({"h": "Activity this period", "b":
                f"New findings: {len(opened)} · remediated: {len(fixed)} · "
                f"scans: {len(scans)} · incidents/breaches/leaks: {len(incidents)}."})
            if incidents:
                sections.append({"h": "Security incidents", "b": "; ".join(
                    (e.get("text_en") or e.get("cve") or "incident") for e in incidents[:8])})

    # template-specific framing so each report type is visibly different, not
    # just the same status snapshot. The rich AI version opens separately.
    tpl_titles = {
        "executive": ("إحاطة أمنية تنفيذية", "Executive Security Briefing"),
        "incident":  ("تقرير الاستجابة للحادث", "Incident Response Report"),
        "compliance":("تقرير تدقيق الامتثال", "Compliance Audit Report"),
        "cmar":      ("تقرير تحليل البرمجيات الخبيثة", "Malware Analysis Report"),
    }
    if template in tpl_titles:
        title = tpl_titles[template][0 if ar else 1]
        if template == "executive":
            sections.insert(0, {"h": "الملخّص التنفيذي" if ar else "Executive Summary",
                "b": (f"الوضع الأمني عند {score}/100 ({'مقبول' if score>=80 else 'يحتاج انتباه' if score>=60 else 'حرج'}). "
                      f"أبرز ما يلزم القيادة: معالجة {counts.get('critical',0)+counts.get('high',0)} ثغرة ذات أولوية." if ar else
                      f"Security posture at {score}/100 ({'good' if score>=80 else 'needs attention' if score>=60 else 'critical'}). "
                      f"Leadership priority: address {counts.get('critical',0)+counts.get('high',0)} high-priority findings.")})
        elif template == "incident":
            sections.insert(0, {"h": "حالة الحوادث" if ar else "Incident Status",
                "b": (f"{threats} تهديد نشط قيد المتابعة. راجع قسم نشاط الفترة أدناه للحوادث المسجّلة." if ar else
                      f"{threats} active threats under tracking. See period activity below for logged incidents.")})
        elif template == "compliance":
            try:
                comp = _run_compliance_check()
                passed = comp.get("passed", 0); total = comp.get("total", 0)
                sections.insert(0, {"h": "ملخّص الامتثال (CIS)" if ar else "Compliance Summary (CIS)",
                    "b": (f"اجتاز {passed} من {total} ضابط أمني. النسبة: {round(100*passed/total) if total else 0}%." if ar else
                          f"Passed {passed} of {total} controls. Score: {round(100*passed/total) if total else 0}%.")})
            except Exception:
                pass
        rep_template = template
    else:
        rep_template = ""

    rep = {"id": rid, "period": period, "tone": tone, "score": score,
           "title": title, "summary": summary, "sections": sections,
           "template": rep_template,
           "auto": False, "ts": int(time.time() * 1000)}
    _REPORTS[rid] = rep
    while len(_REPORTS) > 200:
        _REPORTS.popitem(last=False)
    _save_reports()
    return rep


def _reports_list(lang):
    # REAL reports only — no seeding/placeholders. Newest first.
    return list(reversed(list(_REPORTS.values())))


def weekly_digest(lang="en"):
    """Build a smart weekly summary: aggregate the last 7 days of events, top
    noisy rules, isolations, and resumed apps, then have the local model write a
    short human summary. Falls back to a hand-written summary if the model is
    unavailable. This gives the user a periodic 'what happened & what to do' note."""
    now = int(time.time() * 1000)
    week_ago = now - 7 * 86400000
    evs = events_query(since_ms=week_ago, limit=1000)
    # tally
    by_kind = {}
    by_sev = {"critical": 0, "high": 0, "medium": 0, "low": 0, "info": 0}
    rt_events = 0
    for e in evs:
        k = e.get("kind", "other")
        by_kind[k] = by_kind.get(k, 0) + 1
        by_sev[e.get("sev", "info")] = by_sev.get(e.get("sev", "info"), 0) + 1
        if k == "realtime":
            rt_events += 1
    # noisy rules (from the rate limiter)
    noisy = []
    try:
        for key, dq in _rule_hits.items():
            recent = sum(1 for t in dq if now / 1000 - t <= 7 * 86400)
            if recent >= 3:
                noisy.append((key, recent))
        noisy.sort(key=lambda x: -x[1])
    except Exception:
        pass
    # isolations + resumes
    iso_count = 0
    try:
        iso_count = sum(1 for _ in _isolation_log)
    except Exception:
        pass
    resumed = {}
    try:
        for r in _isolation_resume_learn:
            resumed[r["name"]] = resumed.get(r["name"], 0) + 1
    except Exception:
        pass

    with _sec_lock:
        opens = [f for f in _FINDINGS if f["st"] in _OPEN_STATES]

    stats = {
        "total_events": len(evs),
        "realtime_events": rt_events,
        "by_kind": by_kind,
        "by_sev": by_sev,
        "noisy_rules": noisy[:5],
        "isolations": iso_count,
        "resumed": sorted(resumed.items(), key=lambda x: -x[1])[:5],
        "open_findings": len(opens),
        "baseline": _baseline_status(),
    }

    # ---- narrative (model, with hand-written fallback) ----
    is_ar = (lang == "ar")
    fb = (
        (f"خلال الأسبوع الماضي سجّل Sentinel {len(evs)} حدثاً، منها {rt_events} حدث حماية لحظية. "
         f"الثغرات المفتوحة حالياً: {len(opens)}. "
         + (f"أكثر قاعدة أطلقت تنبيهات: «{noisy[0][0]}» ({noisy[0][1]} مرة) — راجِعها فقد تحتاج ضبطاً. "
            if noisy else "لا توجد قواعد مزعجة هذا الأسبوع. ")
         + (f"أضِف البرامج التي استأنفتها ({', '.join(n for n, _ in stats['resumed'])}) لقائمة السماح لتقليل الإزعاج. "
            if stats["resumed"] else "")
         + ("ما زال خط الأساس في وضع التعلّم؛ ستتحسّن الدقّة بعد اكتماله."
            if stats["baseline"]["learning"] else "خط الأساس مكتمل ويعمل على تقليل الإنذارات الكاذبة."))
        if is_ar else
        (f"Over the past week Sentinel recorded {len(evs)} events, including {rt_events} real-time "
         f"protection events. There are currently {len(opens)} open findings. "
         + (f"The noisiest rule was '{noisy[0][0]}' ({noisy[0][1]}x) — review it, it may need tuning. "
            if noisy else "No noisy rules this week. ")
         + (f"Consider allowlisting the apps you resumed ({', '.join(n for n, _ in stats['resumed'])}) to reduce noise. "
            if stats["resumed"] else "")
         + ("The behavioral baseline is still learning; accuracy will improve once it completes."
            if stats["baseline"]["learning"] else "The behavioral baseline is complete and reducing false positives.")))

    narrative = fb
    try:
        prompt = (
            ("أنت محلل أمن. اكتب ملخّصاً أسبوعياً موجزاً (٣-٤ جمل) بالعربية الفصحى فقط بناءً على هذه الأرقام: "
             if is_ar else
             "You are a security analyst. Write a concise weekly summary (3-4 sentences) based on these numbers: ")
            + json.dumps(stats, ensure_ascii=False)
        )
        data, err = _ollama_call("/api/generate", {
            "model": MODELS[_current_model[0]]["tag"], "prompt": prompt, "stream": False,
            "options": {"temperature": 0.3, "num_predict": 220}})
        ans = (data or {}).get("response", "").strip() if data else ""
        if ans and len(ans) >= 30:
            if is_ar:
                letters = [c for c in ans if c.isalpha()]
                latin = sum(1 for c in letters if "a" <= c.lower() <= "z")
                if not (letters and latin > 0.35 * len(letters)):
                    narrative = ans
            else:
                narrative = ans
    except Exception:
        pass

    return {"period": "weekly", "generated_ts": now, "stats": stats, "summary": narrative}


@bp.get("/api/digest/weekly")
@require_auth("analyst")
def digest_weekly():
    return jsonify(weekly_digest(_lang()))



# ---- automatic periodic report generation -------------------------------
_PERIOD_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_periods.json")
_last_period = _load_json(_PERIOD_FILE, {})   # period -> last-generated epoch (s)


def _save_periods():
    try:
        with open(_PERIOD_FILE, "w", encoding="utf-8") as f:
            json.dump(_last_period, f)
    except Exception:
        pass


def _auto_report_loop():
    """Generate each periodic report automatically once its interval elapses."""
    # first run: just start the clock — do NOT flood the list with reports.
    first = not _last_period
    if first:
        now = time.time()
        for p in _PERIOD_SECONDS:
            _last_period.setdefault(p, now)
        _save_periods()
    while True:
        try:
            now = time.time()
            for p, secs in _PERIOD_SECONDS.items():
                last = _last_period.get(p, now)
                if now - last >= secs:
                    rep = generate_report(p, "en")
                    rep["auto"] = True
                    _save_reports()
                    _last_period[p] = now
                    _save_periods()
                    add_notification("report", "info",
                                     f"Auto {p} report ready", f"تقرير {p} تلقائي جاهز",
                                     "", "", route="reports", report_id=rep["id"])
        except Exception:
            pass
        time.sleep(60)   # check every minute


_auto_thread = [None]


def _ensure_auto_reports():
    if _auto_thread[0] is None:
        t = threading.Thread(target=_auto_report_loop, daemon=True)
        _auto_thread[0] = t
        t.start()


# ==================================================================
#  MODEL (local, via Ollama)  + easy q4/q8 switch
# ==================================================================
OLLAMA = os.environ.get("SENTINEL_OLLAMA", "http://127.0.0.1:11434")
# Context window cap. Foundation-Sec-8B defaults to 128K (needs ~18GB RAM for the
# KV cache); 2048 keeps memory modest so it loads on machines with limited free RAM.
NUM_CTX = int(os.environ.get("SENTINEL_NUM_CTX", "2048"))
MODELS = {
    # The exact models you pulled into Ollama (see `ollama list`).
    # Override with env vars if you re-pull under different tags.
    "q4": {"tag": os.environ.get("SENTINEL_MODEL_Q4", "hf.co/gabriellarson/Foundation-Sec-8B-Instruct-GGUF:Q4_K_M"),
           "label": "Foundation-Sec 8B · Q4", "note_en": "Lighter (~4.9GB) · modest hardware",
           "note_ar": "أخف (~4.9GB) · للأجهزة المتواضعة"},
    "q8": {"tag": os.environ.get("SENTINEL_MODEL_Q8", "hf.co/fdtn-ai/Foundation-Sec-8B-Instruct-Q8_0-GGUF:latest"),
           "label": "Foundation-Sec 8B · Q8", "note_en": "Full (~8.5GB) · higher accuracy",
           "note_ar": "كامل (~8.5GB) · أعلى دقة"},
}
_current_model = ["q4"]   # default to the lighter model
_SYS_PROMPT = ("You are Sentinel, a senior cybersecurity operations analyst. "
               "Answer the user's actual question directly, fluently, and with expert depth — "
               "reason from the live context rather than repeating canned phrasing.")
_model_last_error = [None]


_ollama_models_cache = {"ts": 0, "models": []}


def _discover_ollama_models():
    """List models installed in Ollama via GET /api/tags. Cached 60s. Never raises."""
    now = time.time()
    if now - _ollama_models_cache["ts"] < 60 and _ollama_models_cache["models"]:
        return _ollama_models_cache["models"]
    models = []
    try:
        import urllib.request
        with urllib.request.urlopen(f"{OLLAMA}/api/tags", timeout=4) as r:
            data = json.loads(r.read().decode("utf-8"))
            models = data.get("models", []) or []
    except Exception:
        models = []
    _ollama_models_cache.update(ts=now, models=models)
    return models


def _ollama_call(path, payload):
    """POST to Ollama; return (data, error_string). Captures HTTP error bodies."""
    import urllib.request, urllib.error
    body = json.dumps(payload).encode("utf-8")
    req = urllib.request.Request(f"{OLLAMA}{path}", data=body,
                                 headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(req, timeout=300) as r:  # 8B can be slow to load
            return json.loads(r.read().decode("utf-8")), None
    except urllib.error.HTTPError as e:
        try:
            detail = e.read().decode("utf-8", "replace")[:300]
        except Exception:
            detail = ""
        return None, f"HTTP {e.code}: {detail}"
    except Exception as e:
        return None, f"{type(e).__name__}: {e}"


def _rag_context(q, lang, limit=6):
    """LOCAL retrieval-augmented context: pull the user's own data (findings, recent
    events/incidents, compliance, scan) relevant to the question. 100% on-device —
    no external calls, nothing is sent anywhere."""
    ql = (q or "").lower()
    bits = []
    # open findings (always relevant for a security assistant)
    with _sec_lock:
        opens = [f for f in _FINDINGS if f["st"] in _OPEN_STATES]
    if opens:
        bits.append("OPEN FINDINGS: " + "; ".join(
            f"{f['cve']} {f['sev']} CVSS{f['score']} on {f['asset']}" for f in opens[:8]))
    # recent real events / incidents
    with _events_lock:
        evs = list(_EVENTS)[-40:]
    inc = [e for e in evs if e.get("kind") == "incident"]
    if inc:
        bits.append("RECENT INCIDENTS: " + "; ".join(
            (e.get("text_en") or "")[:80] for e in inc[-5:]))
    # keyword hits across events (so "did anything fail to log in" etc. is grounded)
    kw = [w for w in ql.replace("?", " ").split() if len(w) > 3]
    if kw:
        hits = [e for e in evs if any(w in (e.get("text_en", "") + e.get("text_ar", "")).lower() for w in kw)]
        if hits:
            bits.append("RELATED EVENTS: " + "; ".join((e.get("text_en") or e.get("kind"))[:70] for e in hits[-5:]))
    # compliance / scan summaries when the question hints at them
    if any(w in ql for w in ("compliance", "cis", "امتثال", "harden", "baseline")):
        comp = _run_compliance_check()
        if comp.get("total"):
            bits.append(f"COMPLIANCE: {comp['passed']}/{comp['total']} CIS controls pass ({comp['score']}%)")
    if any(w in ql for w in ("software", "installed", "برامج", "مثبت")):
        bits.append(f"INSTALLED SOFTWARE COUNT: {len(_last_scan.get('software', []))}")
    if not bits:
        return ""
    head = ("بيانات النظام الفعلية (استند إليها في إجابتك، محلية بالكامل): "
            if lang == "ar" else
            "The user's actual local system data (ground your answer in this): ")
    return head + " | ".join(bits)


def model_chat(q, lang, max_tokens=None, system_extra="", history=None):
    """Ask the local model. Tries /api/chat, then falls back to /api/generate
    (HF GGUF models without a chat template only work on /api/generate)."""
    key = _current_model[0]
    tag = MODELS[key]["tag"]
    ctx = (f"Live context: security score {posture_score()}/100, {active_threats()} active threats, "
           f"CPU {_latest['cpu']}%, RAM {_latest['ram']}GB, network {_latest['traffic']} Mbps.")
    # give the model the ACTUAL open findings so it can be specific, not generic
    with _sec_lock:
        opens = [f for f in _FINDINGS if f["st"] in _OPEN_STATES]
    if opens:
        items = "; ".join(f"{f['cve']} ({f['sev']}, CVSS {f['score']}) on {f['asset']} — fix: {f['fix_en']}"
                          for f in opens[:6])
        ctx += " Open findings: " + items + "."
    else:
        ctx += " No open vulnerability findings."
    # Higher temperature => more fluent, less canned/robotic. Penalties still curb
    # broken words / foreign-language leakage from this English-centric model.
    opts = {"num_ctx": NUM_CTX, "temperature": 0.5, "top_p": 0.92, "repeat_penalty": 1.1}
    if max_tokens:
        opts["num_predict"] = int(max_tokens)
    lang_rule = ((" Reply ONLY in fluent Modern Standard Arabic — never English, Russian, or any other script. "
                  "Answer the specific question with the depth of an expert; do not pad or repeat boilerplate.")
                 if lang == "ar" else
                 (" Reply in fluent, professional English with expert depth, addressing the specific question asked."))
    sys = _SYS_PROMPT + lang_rule + (" " + system_extra if system_extra else "") + " " + ctx

    # conversation memory: prepend prior turns so the model remembers the dialogue
    msgs = [{"role": "system", "content": sys}]
    if history:
        for turn in history[-6:]:
            role = "assistant" if turn.get("role") == "assistant" else "user"
            content = (turn.get("content") or "").strip()
            if content:
                msgs.append({"role": role, "content": content[:1500]})
    msgs.append({"role": "user", "content": q})

    # 1) chat endpoint (uses the model's instruct template if present)
    chat_payload = {"model": tag, "stream": False, "messages": msgs}
    if opts:
        chat_payload["options"] = opts
    data, err1 = _ollama_call("/api/chat", chat_payload)
    if data:
        ans = ((data.get("message") or {}).get("content") or "").strip()
        if ans:
            _model_last_error[0] = None
            return {"source": "live", "model": key, "answer": ans}

    # 2) generate endpoint (no chat template required) — fold history into the prompt
    hist_txt = ""
    if history:
        for turn in history[-6:]:
            who = "Assistant" if turn.get("role") == "assistant" else "User"
            hist_txt += f"\n{who}: {(turn.get('content') or '').strip()[:1500]}"
    gen_prompt = (hist_txt + "\nUser: " + q).strip() if hist_txt else q
    gen_payload = {"model": tag, "stream": False, "system": sys, "prompt": gen_prompt}
    if opts:
        gen_payload["options"] = opts
    data2, err2 = _ollama_call("/api/generate", gen_payload)
    if data2:
        ans = (data2.get("response") or "").strip()
        if ans:
            _model_last_error[0] = None
            return {"source": "live", "model": key, "answer": ans}

    _model_last_error[0] = f"chat[{err1 or 'empty'}] | generate[{err2 or 'empty'}]"
    # fallback (model offline / not pulled)
    if lang == "ar":
        ans = (f"النموذج ({MODELS[key]['label']}) غير متصل. افتح /api/model/health للسبب. "
               f"الحالة: درجة الأمان {posture_score()}/100، {active_threats()} تهديد نشط.")
    else:
        ans = (f"Model ({MODELS[key]['label']}) is offline. Open /api/model/health for the reason. "
               f"Posture: score {posture_score()}/100, {active_threats()} active threats.")
    return {"source": "mock", "model": key, "answer": ans, "error": _model_last_error[0]}


# ==================================================================
#  PROTECTION ENGINE  (works as a layer over ANY antivirus, read-only)
#  Uses Windows Security Center, which lists every registered AV product
#  (Defender, Kaspersky, Bitdefender, …). No conflict — we only read.
# ==================================================================
import subprocess
import platform

IS_WINDOWS = platform.system() == "Windows"
_protect = {"available": False, "name": None, "enabled": None, "rtp": None,
            "threats": 0, "detections": [], "ts": 0}
_PROTECT_TTL = 30


def _ps(cmd):
    try:
        # CREATE_NO_WINDOW (0x08000000) stops a console window from flashing on
        # screen every time we poll PowerShell — important for the bundled .exe.
        flags = 0x08000000 if os.name == "nt" else 0
        out = subprocess.run(
            ["powershell", "-NoProfile", "-NonInteractive", "-Command", cmd],
            capture_output=True, text=True, timeout=12,
            creationflags=flags,
        )
        return out.stdout.strip() if out.returncode == 0 else None
    except Exception:
        return None


def _refresh_protection():
    """Detect the active AV product (any vendor) + Defender detections if present."""
    if not IS_WINDOWS:
        return
    if time.time() - _protect["ts"] < _PROTECT_TTL:
        return
    _protect["ts"] = time.time()

    # 1) Any registered antivirus via Security Center (vendor-agnostic)
    av = _ps("Get-CimInstance -Namespace root/SecurityCenter2 -ClassName AntiVirusProduct "
             "| Select-Object displayName,productState | ConvertTo-Json")
    if av:
        try:
            j = json.loads(av)
            if isinstance(j, dict):
                j = [j]
            if j:
                # productState bit 0x1000 in the 2nd byte => enabled (heuristic)
                prod = j[0]
                state = int(prod.get("productState", 0))
                _protect["available"] = True
                _protect["name"] = prod.get("displayName")
                _protect["enabled"] = bool((state & 0x1000) or (state & 0x10))
        except Exception:
            pass

    # 2) Real-time protection + detections (Defender-specific bonus, if present)
    status = _ps("Get-MpComputerStatus | Select-Object RealTimeProtectionEnabled | ConvertTo-Json")
    if status:
        try:
            _protect["rtp"] = bool(json.loads(status).get("RealTimeProtectionEnabled"))
        except Exception:
            pass
    dets = _ps("Get-MpThreatDetection | Select-Object ThreatID,InitialDetectionTime | ConvertTo-Json")
    if dets:
        try:
            j = json.loads(dets)
            if isinstance(j, dict):
                j = [j]
            active = [d for d in j if d]
            _protect["detections"] = active[:10]
            _protect["threats"] = len(active)
        except Exception:
            _protect["detections"] = []
            _protect["threats"] = 0


def protection_threats():
    return _protect["threats"] if _protect["available"] else 0


def protection_status():
    return {"available": _protect["available"], "name": _protect["name"],
            "enabled": _protect["enabled"], "rtp": _protect["rtp"]}


# ==================================================================
#  MODEL-DRIVEN OVERVIEW SUMMARY  (refreshed periodically by the model)
# ==================================================================
SUMMARY_SEC = 45
_ai_summary = {"en": None, "ar": None, "ts": 0, "score": None}
_summary_started = [False]


def _trend(key):
    s = _series(key)
    if len(s) < 10:
        return 0
    recent = sum(s[-5:]) / 5.0
    older = sum(s[-10:-5]) / 5.0
    return round(recent - older, 1)


def _summary_loop():
    last_run = 0
    last_score = None
    while True:
        try:
            score_now = posture_score()
            # regenerate on a timer OR immediately when the score changes
            if (time.time() - last_run >= SUMMARY_SEC) or (score_now != last_score):
                last_run = time.time(); last_score = score_now
                with _sec_lock:
                    opens = [f for f in _FINDINGS if f["st"] in _OPEN_STATES]
                # language-specific, natural-language facts (no English tokens / no raw telemetry)
                def _facts(ar):
                    if opens:
                        if ar:
                            items = "؛ ".join(f"{f['cve']} (خطورة {f['sev']}، CVSS {f['score']}) على {f['asset']}، الحل: {f['fix_ar']}" for f in opens[:6])
                            return f"الدرجة {score_now} من 100. عدد التهديدات النشطة {active_threats()}. الثغرات المفتوحة: {items}."
                        items = "; ".join(f"{f['cve']} (severity {f['sev']}, CVSS {f['score']}) on {f['asset']}, fix: {f['fix_en']}" for f in opens[:6])
                        return f"Score {score_now}/100. Active threats {active_threats()}. Open findings: {items}."
                    return (f"الدرجة {score_now} من 100، لا تهديدات نشطة ولا ثغرات مفتوحة." if ar
                            else f"Score {score_now}/100, no active threats and no open findings.")
                for lng in ("en", "ar"):
                    if lng == "ar":
                        prompt = ("اكتب تحليلاً أمنياً موجزاً (جملتان أو ثلاث) بالعربية الفصحى فقط، دون أي كلمات إنجليزية "
                                  "أو رموز أو أرقام مكسورة. اذكر أخطر ثغرة مفتوحة باسمها وأثرها المحتمل والإجراء الموصى به. "
                                  "إن لم توجد ثغرات مفتوحة فاذكر أن الوضع سليم وما الذي يُنصح بمراقبته.\n\n" + _facts(True))
                    else:
                        prompt = ("Write a concise security analysis (2-3 sentences) in clear English. Name the most severe "
                                  "open finding, its likely impact, and the recommended action. If there are no open "
                                  "findings, say the posture is clean and what to keep monitoring.\n\n" + _facts(False))
                    res = model_chat(prompt, lng, max_tokens=220)
                    if res.get("source") == "live" and res.get("answer"):
                        _ai_summary[lng] = res["answer"].strip()
                _ai_summary["ts"] = time.time()
                _ai_summary["score"] = score_now
        except Exception:
            pass
        time.sleep(3)   # check often so score changes reflect quickly


def _ensure_summary():
    if _summary_started[0]:
        return
    _summary_started[0] = True
    threading.Thread(target=_summary_loop, daemon=True).start()


# ==================================================================
#  HELPERS
# ==================================================================
def _lang():
    lang = (request.args.get("lang") or "en").lower()
    return "ar" if lang.startswith("ar") else "en"


VERDICT = {"en": "All systems nominal.", "ar": "جميع الأنظمة طبيعية."}


def _summary(lang):
    # prefer a fresh, model-generated summary; fall back to a STABLE template.
    cached = _ai_summary.get(lang)
    if cached and _ai_summary.get("score") == posture_score() and (time.time() - _ai_summary["ts"] < SUMMARY_SEC * 4):
        return cached
    # template is intentionally free of per-second CPU/RAM numbers so it doesn't
    # re-type every tick; it only changes when score/threats change.
    s, th = posture_score(), active_threats()
    if lang == "ar":
        posture = "قوي" if s >= 80 else ("متوسط" if s >= 60 else "يحتاج إجراءً عاجلاً")
        return (f"وضع المخاطر {posture} (درجة {s}/100). "
                + (f"يوجد {th} تهديد نشط يتطلب المتابعة. " if th else "لا توجد تهديدات نشطة حالياً. ")
                + "جارٍ تحليل النموذج للحالة…")
    posture = "strong" if s >= 80 else ("moderate" if s >= 60 else "needs urgent action")
    return (f"Risk posture is {posture} (score {s}/100). "
            + (f"{th} active threat(s) require attention. " if th else "No active threats right now. ")
            + "Model analysis in progress…")


def _overview_payload(lang):
    s, th = posture_score(), active_threats()
    mbps = "م.ب/ث" if lang == "ar" else "Mbps"
    # real trend direction from the rolling history (no fabricated percentages)
    def _trend(key):
        seq = _series(key)
        if len(seq) < 4:
            return ("flat", "")
        recent = sum(seq[-3:]) / 3.0
        older = sum(seq[-6:-3]) / 3.0 if len(seq) >= 6 else seq[0]
        if older == 0:
            return ("flat", "")
        pct = (recent - older) / max(0.01, abs(older)) * 100
        if pct > 5:   return ("up", f"{abs(pct):.0f}%")
        if pct < -5:  return ("down", f"{abs(pct):.0f}%")
        return ("flat", "")
    cpu_td, cpu_tt = _trend("cpu")
    ram_td, ram_tt = _trend("ram")
    tr_td, tr_tt = _trend("traffic")
    estimated = not HAVE_PSUTIL
    return {
        "source": "estimated" if estimated else "live", "ts": int(time.time() * 1000),
        "estimated": estimated,
        "score": s, "verdict": VERDICT[lang], "summary": _summary(lang), "secured": True,
        "vitals": {"threats": str(th), "uptime": _uptime_str()},
        "chips": {"threats": str(th), "cpu": f"{_latest['cpu']}%",
                  "net": str(_latest["traffic"]), "score": str(s)},
        "cards": {
            "cpu":     {"value": str(_latest["cpu"]),  "unit": "%",  "trendDir": cpu_td, "trendText": cpu_tt},
            "ram":     {"value": str(_latest["ram"]),  "unit": "GB", "trendDir": ram_td, "trendText": ram_tt},
            "threat":  {"value": str(th),                            "trendDir": "flat", "trendText": ""},
            "traffic": {"value": str(_latest["traffic"]), "unit": mbps, "trendDir": tr_td, "trendText": tr_tt},
        },
        "series": {k: _series(k) for k in ("cpu", "ram", "threat", "traffic", "netIn", "netOut")},
        "scoreTrend": list(_score_hist),
        "ramPct": _latest_extra.get("ram_pct", 0),
        "protection": protection_status(),
        "reports": [{"tag": r["title"], "tone": r["tone"], "time": "", "read": 3,
                     "title": r["title"], "ex": r["summary"], "id": r["id"]}
                    for r in _reports_list(lang)[:3]],
        "stream": None,
    }


# ==================================================================
#  REAL SYSTEM METRICS  (uptime, disks, processes, connections, ports…)
# ==================================================================
def _uptime_seconds():
    if HAVE_PSUTIL:
        try:
            return max(0, int(time.time() - psutil.boot_time()))
        except Exception:
            return 0
    return int(time.time() - _PROC_START)


def _uptime_str():
    s = _uptime_seconds()
    d, r = divmod(s, 86400)
    h, r = divmod(r, 3600)
    m = r // 60
    if d:
        return f"{d}d {h}h {m}m"
    if h:
        return f"{h}h {m}m"
    return f"{m}m"


def _disks():
    out = []
    if HAVE_PSUTIL:
        try:
            seen = set()
            for part in psutil.disk_partitions(all=False):
                if part.fstype == "" or part.mountpoint in seen:
                    continue
                seen.add(part.mountpoint)
                try:
                    u = psutil.disk_usage(part.mountpoint)
                except (PermissionError, OSError):
                    continue
                out.append({"mount": part.mountpoint, "pct": round(u.percent, 1),
                            "used_gb": round(u.used / 1024**3, 1),
                            "total_gb": round(u.total / 1024**3, 1)})
        except Exception:
            pass
    return out


def _top_processes(n=8):
    out = []
    if HAVE_PSUTIL:
        try:
            ncpu = psutil.cpu_count(logical=True) or 1
            procs = []
            for p in psutil.process_iter(["pid", "name", "memory_info", "cpu_percent"]):
                try:
                    info = p.info
                    mem = (info.get("memory_info").rss / 1024**2) if info.get("memory_info") else 0
                    # psutil sums CPU across cores (can exceed 100%); Task Manager
                    # shows it divided by core count — normalise to match TM.
                    cpu = (info.get("cpu_percent") or 0) / ncpu
                    procs.append({"pid": info.get("pid"), "name": info.get("name") or "?",
                                  "mem": round(mem, 1), "cpu": round(cpu, 1)})
                except (psutil.NoSuchProcess, psutil.AccessDenied):
                    continue
            procs.sort(key=lambda x: x["mem"], reverse=True)
            out = procs[:n]
        except Exception:
            pass
    return out


def _net_conn_stats():
    established = listening = 0
    ports = []
    if HAVE_PSUTIL:
        try:
            for c in psutil.net_connections(kind="inet"):
                if c.status == "ESTABLISHED":
                    established += 1
                elif c.status == "LISTEN":
                    listening += 1
                    if c.laddr and len(ports) < 24:
                        ports.append(c.laddr.port)
        except (psutil.AccessDenied, PermissionError, Exception):
            pass
    return {"established": established, "listening": listening,
            "ports": sorted(set(ports))}


# ==================================================================
#  COMPLIANCE (CIS-style baseline)  — maps real config checks to controls
# ==================================================================
def _net_accounts():
    """Parse 'net accounts' for password policy (Windows)."""
    out = {}
    if os.name != "nt":
        return out
    raw = _run_cmd(["net", "accounts"], 8)
    import re as _r
    m = _r.search(r"[Mm]inimum password length[^\d]*(\d+)", raw or "")
    if m:
        out["min_len"] = int(m.group(1))
    m = _r.search(r"[Ll]ockout threshold[^\d]*(\d+)", raw or "")
    if m:
        out["lockout"] = int(m.group(1))
    return out


def _run_compliance_check():
    """Return CIS-style controls with pass/fail based on real machine state."""
    win = (os.name == "nt")
    controls = []

    def add(cid, en, ar, status, detail=""):
        controls.append({"id": cid, "title_en": en, "title_ar": ar, "status": status, "detail": detail})

    if not win:
        add("CIS-0", "Compliance checks run on Windows hosts",
            "فحوص الامتثال تعمل على أنظمة ويندوز", "manual", "Non-Windows host")
        return {"controls": controls, "score": 0, "passed": 0, "total": 0}

    # firewall
    fw = _run_cmd(["netsh", "advfirewall", "show", "allprofiles", "state"], 10)
    add("CIS-9.1", "Windows Firewall enabled on all profiles", "تفعيل الجدار الناري لكل الملفات",
        "pass" if (fw and "off" not in fw.lower()) else "fail")
    # Defender RTP
    try:
        rtp = protection_status().get("realtime", None)
        add("CIS-18.9", "Antivirus real-time protection on", "الحماية الفورية لمكافحة الفيروسات",
            "pass" if rtp else ("fail" if rtp is False else "manual"))
    except Exception:
        add("CIS-18.9", "Antivirus real-time protection on", "الحماية الفورية", "manual")
    # SMBv1
    smb1 = _reg_get("HKLM", r"SYSTEM\CurrentControlSet\Services\LanmanServer\Parameters", "SMB1")
    add("CIS-18.3.1", "SMABv1 disabled", "تعطيل SMBv1", "pass" if smb1 != 1 else "fail")
    # UAC
    lua = _reg_get("HKLM", r"SOFTWARE\Microsoft\Windows\CurrentVersion\Policies\System", "EnableLUA")
    add("CIS-2.3.17", "User Account Control enabled", "تفعيل UAC", "pass" if lua != 0 else "fail")
    # RDP NLA
    deny = _reg_get("HKLM", r"SYSTEM\CurrentControlSet\Control\Terminal Server", "fDenyTSConnections")
    nla = _reg_get("HKLM", r"SYSTEM\CurrentControlSet\Control\Terminal Server\WinStations\RDP-Tcp", "UserAuthentication")
    if deny == 0:
        add("CIS-18.10", "RDP requires Network Level Authentication", "RDP يشترط NLA", "pass" if nla == 1 else "fail")
    else:
        add("CIS-18.10", "RDP disabled (or NLA enforced)", "RDP معطّل أو يشترط NLA", "pass")
    # guest account
    g = _run_cmd(["net", "user", "guest"], 8)
    import re as _r2
    guest_active = bool(_r2.search(r"active\s+yes", (g or "").lower()))
    add("CIS-2.3.1.2", "Guest account disabled", "تعطيل حساب الضيف", "fail" if guest_active else "pass")
    # password policy
    na = _net_accounts()
    if "min_len" in na:
        add("CIS-1.1.4", "Minimum password length >= 14", "حد أدنى لطول كلمة السر ≥ ١٤",
            "pass" if na["min_len"] >= 14 else "fail", f"current: {na['min_len']}")
    if "lockout" in na:
        add("CIS-1.2.2", "Account lockout threshold set", "تعيين عتبة قفل الحساب",
            "pass" if 0 < na.get("lockout", 0) <= 10 else "fail", f"current: {na.get('lockout')}")
    # BitLocker (best-effort)
    bl = _run_cmd(["powershell", "-NoProfile", "-Command",
                   "(Get-BitLockerVolume -MountPoint $env:SystemDrive -ErrorAction SilentlyContinue).ProtectionStatus"], 12)
    if bl and bl.strip():
        add("CIS-18.9.11", "System drive encrypted (BitLocker)", "تشفير قرص النظام (BitLocker)",
            "pass" if "1" in bl or "On" in bl else "fail")

    total = len([c for c in controls if c["status"] in ("pass", "fail")])
    passed = len([c for c in controls if c["status"] == "pass"])
    score = round(passed / total * 100) if total else 0
    return {"controls": controls, "score": score, "passed": passed, "total": total}


# ==================================================================
#  NETWORK ANALYSIS — local geo classification (privacy-safe), connection
#  anomaly / scan detection, and a user-managed IP blocklist.
# ==================================================================
# ==================================================================
#  PRIVACY-FIRST GEO-IP  (no MaxMind, no registration, no API calls)
#  Uses a free public CSV (DB-IP Lite, CC-BY-4.0) mirrored on GitHub.
#  Behaviour:
#    - DB is OPTIONAL and OFF by default; until the user downloads it,
#      lookups fall back to private/public classification.
#    - The ONE download is anonymous (no key, no account, generic UA)
#      from raw.githubusercontent.com — a general-purpose CDN serving
#      millions of files; the request is indistinguishable from any
#      other repo download.
#    - After the one-time download, all lookups happen in-memory on
#      this machine. Zero outbound traffic at lookup time.
#    - Honors SENTINEL_OFFLINE: no download when offline.
#    - Attribution: "IP-to-country data © DB-IP.com, CC-BY-4.0".
# ==================================================================
import bisect as _bisect

_GEO_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_geoip.csv")
_GEO_URL = os.environ.get(
    "SENTINEL_GEOIP_URL",
    "https://raw.githubusercontent.com/sapics/ip-location-db/main/dbip-country/dbip-country-ipv4.csv")
_geo_starts = []          # parallel arrays: starts[i], ends[i], cc[i]
_geo_ends = []
_geo_cc = []
_geo_lock = threading.Lock()


def _ip_to_int(ip):
    parts = ip.split(".")
    if len(parts) != 4:
        return None
    try:
        a, b, c, d = (int(p) for p in parts)
    except ValueError:
        return None
    if not (0 <= a < 256 and 0 <= b < 256 and 0 <= c < 256 and 0 <= d < 256):
        return None
    return (a << 24) | (b << 16) | (c << 8) | d


def _load_geo():
    """Load the CSV into sorted parallel arrays for O(log N) lookup."""
    starts, ends, cc = [], [], []
    if not os.path.exists(_GEO_FILE):
        return False
    try:
        with open(_GEO_FILE, "r", encoding="utf-8") as f:
            for line in f:
                line = line.strip()
                if not line or line.startswith("#") or "," not in line:
                    continue
                parts = line.split(",")
                if len(parts) < 3:
                    continue
                a = _ip_to_int(parts[0]); b = _ip_to_int(parts[1])
                if a is None or b is None:
                    continue
                starts.append(a); ends.append(b); cc.append(parts[2][:2])
        # sort by start address
        order = sorted(range(len(starts)), key=lambda i: starts[i])
        with _geo_lock:
            _geo_starts[:] = [starts[i] for i in order]
            _geo_ends[:] = [ends[i] for i in order]
            _geo_cc[:] = [cc[i] for i in order]
        return True
    except Exception:
        return False


def _geo_lookup(ip):
    """Find the country code for an IPv4 address. Local, in-memory, O(log N)."""
    if not _geo_starts:
        return None
    n = _ip_to_int(ip)
    if n is None:
        return None
    idx = _bisect.bisect_right(_geo_starts, n) - 1
    if 0 <= idx < len(_geo_starts) and _geo_starts[idx] <= n <= _geo_ends[idx]:
        return _geo_cc[idx]
    return None


def update_geoip(force=False):
    """Anonymous one-time download of the IP-to-country DB. Privacy: GET only,
    no key, generic UA, no host data is ever transmitted."""
    if _OFFLINE:
        return {"status": "offline"}
    if os.path.exists(_GEO_FILE) and not force:
        return {"status": "ok", "ranges": len(_geo_starts), "skipped": "already present"}
    try:
        req = _ur.Request(_GEO_URL, headers={"User-Agent": "Sentinel-SOC"})
        with _ur.urlopen(req, timeout=120) as r:
            data = r.read()
        with open(_GEO_FILE, "wb") as f:
            f.write(data)
        _load_geo()
        return {"status": "ok", "ranges": len(_geo_starts), "bytes": len(data),
                "attribution": "IP-to-country © DB-IP.com (CC-BY-4.0)"}
    except Exception as e:
        return {"status": "error", "error": str(e)[:200]}


# load the DB at startup if it has already been downloaded
_load_geo()


def _ip_class(ip):
    """Classify an IP. Privacy: 100% local — no outbound lookup ever, no IP leaves
    the machine. Returns 'private', 'loopback', a 2-letter country code, or
    'public' if the optional country DB hasn't been downloaded yet."""
    if not ip:
        return "?"
    if ip.startswith(("127.", "::1")):
        return "loopback"
    if ip.startswith(("10.", "192.168.", "169.254.", "fe80", "fc", "fd")) or \
       any(ip.startswith(f"172.{i}.") for i in range(16, 32)):
        return "private"
    # local country lookup (if the DB has been downloaded)
    cc = _geo_lookup(ip)
    return cc if cc else "public"


def _blocklist():
    return (_SETTINGS.get("blocklist") or [])


def _net_analysis():
    """Established connections grouped by remote IP, scan/anomaly flags, blocklist hits."""
    conns = []
    by_remote = {}
    blocked_hits = []
    bl = set(_blocklist())
    if HAVE_PSUTIL:
        try:
            for c in psutil.net_connections(kind="inet"):
                if c.status != "ESTABLISHED" or not c.raddr:
                    continue
                rip = c.raddr.ip
                by_remote[rip] = by_remote.get(rip, 0) + 1
                if rip in bl:
                    blocked_hits.append(rip)
        except (psutil.AccessDenied, PermissionError, Exception):
            pass
    remotes = sorted(by_remote.items(), key=lambda kv: kv[1], reverse=True)
    top = [{"ip": ip, "count": n, "class": _ip_class(ip)} for ip, n in remotes[:12]]
    # anomaly: a single remote IP holding an unusually high number of connections
    anomalies = [{"ip": ip, "count": n} for ip, n in remotes if n >= 20]
    return {"remotes": top, "distinct_remotes": len(by_remote),
            "anomalies": anomalies, "blocked_hits": sorted(set(blocked_hits)),
            "blocklist": sorted(bl)}


def _firewall_rule_count():
    if os.name != "nt":
        return None
    raw = _run_cmd(["powershell", "-NoProfile", "-Command",
                    "(Get-NetFirewallRule -ErrorAction SilentlyContinue | Measure-Object).Count"], 15)
    try:
        return int((raw or "").strip().split()[0])
    except Exception:
        return None


def _mttr_hours():
    """Mean time to remediate: avg(finding_patched.ts - matching finding_open.ts)."""
    with _events_lock:
        opens = {}
        deltas = []
        for e in _EVENTS:
            if e["kind"] == "finding_open" and e.get("cve"):
                opens[e["cve"]] = e["ts"]
            elif e["kind"] == "finding_patched" and e.get("cve") in opens:
                deltas.append(e["ts"] - opens.pop(e["cve"]))
    if not deltas:
        return None
    return round(sum(deltas) / len(deltas) / 3600000, 1)   # ms -> hours


def _online_users():
    with _auth_lock:
        return 1 if _current_user[0] else 0


# heavy psutil reads (processes/disks/connections) are cached and refreshed in a
# background thread so the /metrics endpoint always returns instantly (no UI timeout).
_metrics_cache = {"disks": [], "processes": [], "connections": {"established": 0, "listening": 0, "ports": []},
                  "ts": 0, "io": [], "battery": None, "gpu": [], "temps": [], "disk_health": [], "services": None}
_metrics_thread = [None]
_io_prev = {"t": None, "per": {}}


def _disk_io_rates():
    """Per-disk read/write throughput in MB/s (delta between samples)."""
    out = []
    if not HAVE_PSUTIL:
        return out
    try:
        now = time.time()
        cur = psutil.disk_io_counters(perdisk=True) or {}
        prev_t = _io_prev["t"]
        prev = _io_prev["per"]
        if prev_t is not None:
            dt = max(0.5, now - prev_t)
            for name, c in cur.items():
                p = prev.get(name)
                if not p:
                    continue
                rd = (c.read_bytes - p.read_bytes) / dt / 1024**2
                wr = (c.write_bytes - p.write_bytes) / dt / 1024**2
                if rd >= 0 and wr >= 0:
                    out.append({"disk": name, "read": round(rd, 2), "write": round(wr, 2)})
        _io_prev["t"] = now
        _io_prev["per"] = cur
        out.sort(key=lambda x: x["read"] + x["write"], reverse=True)
    except Exception:
        pass
    return out[:8]


def _battery():
    if not HAVE_PSUTIL or not hasattr(psutil, "sensors_battery"):
        return None
    try:
        b = psutil.sensors_battery()
        if b is None:
            return None
        secs = b.secsleft
        mins = None if secs in (psutil.POWER_TIME_UNLIMITED, psutil.POWER_TIME_UNKNOWN, None) else max(0, int(secs / 60))
        return {"percent": round(b.percent, 0), "plugged": bool(b.power_plugged), "mins_left": mins}
    except Exception:
        return None


def _temps():
    """CPU/component temperatures. psutil works on Linux; on Windows we try WMI."""
    out = []
    if HAVE_PSUTIL and hasattr(psutil, "sensors_temperatures"):
        try:
            t = psutil.sensors_temperatures() or {}
            for chip, entries in t.items():
                for e in entries:
                    if e.current:
                        out.append({"label": (e.label or chip)[:24], "c": round(e.current, 1)})
        except Exception:
            pass
    if not out and os.name == "nt":
        # best-effort WMI thermal zone (Kelvin*10) — not all machines expose it
        raw = _run_cmd(["powershell", "-NoProfile", "-Command",
                        "(Get-CimInstance -Namespace root/wmi -ClassName MSAcpi_ThermalZoneTemperature "
                        "-ErrorAction SilentlyContinue).CurrentTemperature"], 8)
        import re as _r
        for m in _r.findall(r"\d+", raw or ""):
            try:
                c = int(m) / 10.0 - 273.15
                if 10 < c < 110:
                    out.append({"label": "CPU zone", "c": round(c, 1)})
            except Exception:
                pass
    return out[:6]


def _gpu():
    """NVIDIA GPUs via nvidia-smi (if present). Graceful when absent (AMD/Intel/none)."""
    out = []
    raw = _run_cmd(["nvidia-smi",
                    "--query-gpu=name,utilization.gpu,memory.used,memory.total,temperature.gpu",
                    "--format=csv,noheader,nounits"], 8)
    if raw and "," in raw:
        for line in raw.strip().splitlines():
            parts = [p.strip() for p in line.split(",")]
            if len(parts) >= 5:
                try:
                    out.append({"name": parts[0], "util": float(parts[1]),
                                "mem_used": float(parts[2]), "mem_total": float(parts[3]),
                                "temp": float(parts[4])})
                except Exception:
                    pass
    return out


def _disk_health():
    """Physical-disk health (SMART). Windows: Get-PhysicalDisk; Linux: best-effort."""
    out = []
    if os.name == "nt":
        raw = _run_cmd(["powershell", "-NoProfile", "-Command",
                        "Get-PhysicalDisk | Select-Object FriendlyName,HealthStatus,MediaType | "
                        "ConvertTo-Json -Compress"], 12)
        try:
            data = json.loads(raw) if raw.strip() else []
            data = data if isinstance(data, list) else [data]
            for d in data:
                out.append({"name": (d.get("FriendlyName") or "Disk")[:40],
                            "health": d.get("HealthStatus") or "Unknown",
                            "type": d.get("MediaType") or ""})
        except Exception:
            pass
    return out


def _services_summary():
    """Windows service counts (running/stopped) + a few key security services."""
    if os.name != "nt" or not hasattr(psutil, "win_service_iter"):
        return None
    running = stopped = 0
    try:
        for s in psutil.win_service_iter():
            try:
                st = s.status()
                if st == "running":
                    running += 1
                elif st == "stopped":
                    stopped += 1
            except Exception:
                continue
    except Exception:
        return None
    return {"running": running, "stopped": stopped, "total": running + stopped}


def _refresh_metrics_cache():
    # disks + processes first (fast, safe) so they populate even if the
    # connection scan below is slow / blocked by Windows permissions.
    try:
        _metrics_cache["disks"] = _disks()
    except Exception:
        pass
    try:
        _metrics_cache["processes"] = _top_processes()
    except Exception:
        pass
    _metrics_cache["ts"] = int(time.time() * 1000)
    try:
        _metrics_cache["io"] = _disk_io_rates()       # fast (psutil)
    except Exception:
        pass
    try:
        _metrics_cache["battery"] = _battery()         # fast (psutil)
    except Exception:
        pass
    try:
        _metrics_cache["connections"] = _net_conn_stats()   # may be slow on Windows
    except Exception:
        pass


def _refresh_metrics_heavy():
    """Slower probes (external commands) — refreshed less often to stay light."""
    for key, fn in (("gpu", _gpu), ("temps", _temps), ("disk_health", _disk_health), ("services", _services_summary)):
        try:
            _metrics_cache[key] = fn()
        except Exception:
            pass


def _metrics_loop():
    _refresh_metrics_cache()          # fast first paint (cpu reads accurate after prime)
    if HAVE_PSUTIL:
        try:
            for p in psutil.process_iter():
                try:
                    p.cpu_percent(None)
                except Exception:
                    pass
        except Exception:
            pass
    _refresh_metrics_heavy()          # first GPU/temp/health/services read
    n = 0
    while True:
        _refresh_metrics_cache()
        n += 1
        if n % 6 == 0:                # every ~30s
            _refresh_metrics_heavy()
        time.sleep(5)


def _ensure_metrics():
    # NEVER refresh synchronously here — net_connections can block on Windows and
    # would stall the request. Just start the background thread; the cache fills
    # within a few seconds and the endpoint always returns instantly.
    if _metrics_thread[0] is None:
        t = threading.Thread(target=_metrics_loop, daemon=True)
        _metrics_thread[0] = t
        t.start()


def _system_metrics(lang):
    cores = psutil.cpu_count(logical=True) if HAVE_PSUTIL else 8
    try:
        vm_total = round(psutil.virtual_memory().total / 1024**3, 1) if HAVE_PSUTIL else 16
    except Exception:
        vm_total = 0
    return {
        "uptime": _uptime_str(),
        "uptime_seconds": _uptime_seconds(),
        "cpu": {"pct": _latest.get("cpu", 0), "cores": cores},
        "ram": {"pct": _latest_extra.get("ram_pct", 0), "used_gb": _latest.get("ram", 0), "total_gb": vm_total},
        "disks": _metrics_cache["disks"],
        "processes": _metrics_cache["processes"],
        "connections": _metrics_cache["connections"],
        "online_users": _online_users(),
        "registered_users": len(_USERS),
        "mttr_hours": _mttr_hours(),
        "io": _metrics_cache.get("io", []),
        "battery": _metrics_cache.get("battery"),
        "gpu": _metrics_cache.get("gpu", []),
        "temps": _metrics_cache.get("temps", []),
        "disk_health": _metrics_cache.get("disk_health", []),
        "services": _metrics_cache.get("services"),
        "score": posture_score(),
        "score_trend": list(_score_hist),
        "series": {k: _series(k) for k in ("cpu", "ram", "traffic", "netIn", "netOut")},
    }


@bp.get("/api/system/metrics")
def system_metrics():
    _ensure_sampler()
    _ensure_metrics()
    return jsonify(_system_metrics(_lang()))


# ==================================================================
#  ROUTES
# ==================================================================
@bp.get("/api/dashboard/overview")
def overview():
    _ensure_sampler()
    return jsonify(_overview_payload(_lang()))


@bp.get("/api/stream")
def stream():
    """Server-Sent Events: pushes a live tick (~2s) + events for instant UI."""
    _ensure_sampler()
    lang = _lang()
    q = collections.deque()
    with _sub_lock:
        _subscribers.append(q)

    def gen():
        try:
            # immediate first paint
            yield "data: " + json.dumps({"type": "tick", "data": _overview_payload(lang)}) + "\n\n"
            last = 0
            while True:
                # drain any queued events first (notifications, scan results…)
                while q:
                    yield "data: " + json.dumps(q.popleft()) + "\n\n"
                now = time.time()
                if now - last >= SAMPLE_SEC:
                    last = now
                    yield "data: " + json.dumps({"type": "tick", "data": _overview_payload(lang)}) + "\n\n"
                time.sleep(0.25)
        except GeneratorExit:
            pass
        finally:
            with _sub_lock:
                if q in _subscribers:
                    _subscribers.remove(q)

    return Response(gen(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@bp.post("/api/dashboard/scan-now")
@require_auth("analyst")
def scan_now():
    """Run a REAL vulnerability scan of THIS machine (not a demo). Returns the
    refreshed posture and the number of findings actually discovered."""
    before = {f["cve"] for f in _FINDINGS}
    state = run_scan_now()                 # real OS/baseline/software/KEV scan
    after = [f for f in _FINDINGS if f["cve"] not in before and f.get("st") == "open"]
    discovered = None
    if after:
        # surface the single highest-severity NEW finding for the notification
        rank = {"critical": 4, "high": 3, "medium": 2, "low": 1, "info": 0}
        nf = max(after, key=lambda f: rank.get(f.get("sev"), 0))
        add_notification("alert", "danger" if nf["sev"] in ("critical", "high") else "warning",
                         f"New {nf['sev']} finding: {nf['cve']}",
                         f"ثغرة {nf['sev']} جديدة: {nf['cve']}",
                         f"on {nf.get('asset','')}", f"على {nf.get('asset','')}", route="security")
        discovered = {"cve": nf["cve"], "sev": nf["sev"], "asset": nf.get("asset", ""),
                      "new_count": len(after)}
    _push_event({"type": "tick", "data": _overview_payload(_lang())})
    return jsonify({"ok": True, "source": "live", "discovered": discovered,
                    "findings": state.get("count", len(_FINDINGS)),
                    "score": posture_score(), "threats": active_threats(),
                    "startedAt": int(time.time() * 1000)})


# ==================================================================
#  SCAN AN UPLOADED FILE (log / json / text) -> model analysis -> report
# ==================================================================
import re as _re

_THREAT_WORDS = ["error", "fail", "failed", "denied", "unauthorized", "attack", "malware",
                 "exploit", "injection", "breach", "ransom", "trojan", "brute", "suspicious",
                 "critical", "alert", "blocked", "intrusion", "خطأ", "فشل", "هجوم", "اختراق"]


def _log_heuristic_scan(content):
    """Lightweight keyword/IP triage of an uploaded LOG or TEXT file (not the
    behavioral process engine — that's _heuristic_scan() with no args)."""
    lines = content.splitlines()
    low = content.lower()
    hits = {w: low.count(w) for w in _THREAT_WORDS if low.count(w) > 0}
    ips = sorted(set(_re.findall(r"\b\d{1,3}(?:\.\d{1,3}){3}\b", content)))[:15]
    sev = "danger" if any(k in hits for k in ("attack", "malware", "exploit", "breach", "ransom", "intrusion")) \
        else ("warning" if hits else "success")
    return {"lines": len(lines), "hits": hits, "ips": ips, "tone": sev}


def _file_hashes(raw):
    """Local cryptographic hashes — computed on-device, nothing transmitted."""
    return {"sha256": hashlib.sha256(raw).hexdigest(),
            "sha1": hashlib.sha1(raw).hexdigest(),
            "md5": hashlib.md5(raw).hexdigest()}


def _entropy(raw):
    """Shannon entropy (0–8). High (>7.2) suggests packed/encrypted/compressed content."""
    if not raw:
        return 0.0
    import math
    counts = [0] * 256
    for b in raw[:262144]:           # sample up to 256KB
        counts[b] += 1
    n = sum(counts)
    ent = 0.0
    for c in counts:
        if c:
            p = c / n
            ent -= p * math.log2(p)
    return round(ent, 2)


def _file_type(raw, filename):
    """Identify a file locally by magic bytes / extension (no upload)."""
    head = raw[:8]
    ext = (filename.rsplit(".", 1)[-1].lower() if "." in filename else "")
    if head[:2] == b"MZ":
        return "PE executable (Windows .exe/.dll)"
    if head[:4] == b"\x7fELF":
        return "ELF executable (Linux)"
    if head[:4] == b"%PDF":
        return "PDF document"
    if head[:2] == b"PK":
        if ext in ("docx", "xlsx", "pptx", "docm", "xlsm"):
            return "Office Open XML document"
        return "ZIP archive"
    if head[:4] in (b"\xd0\xcf\x11\xe0",):
        return "Legacy Office document (OLE)"
    if ext in ("ps1", "bat", "cmd", "sh", "vbs", "js", "py"):
        return f"Script ({ext})"
    try:
        raw[:2048].decode("utf-8")
        return "Text / log"
    except Exception:
        return "Binary / unknown"


def _binary_indicators(raw, ftype, lang):
    """Local static indicators — no sandbox, no upload."""
    ar = (lang == "ar")
    flags = []
    low = raw[:200000].lower()
    if "PE executable" in ftype:
        if b"this program cannot be run in dos mode" in low:
            pass
        for s in (b"virtualalloc", b"createremotethread", b"writeprocessmemory",
                  b"loadlibrary", b"getprocaddress", b"wininet", b"urldownloadtofile"):
            if s in low:
                flags.append(("استدعاء API حسّاس: " if ar else "Sensitive API: ") + s.decode())
    if "Office" in ftype and b"vbaproject.bin" in low:
        flags.append("يحتوي ماكرو VBA (vbaProject.bin)" if ar else "Contains VBA macros (vbaProject.bin)")
    for s in (b"powershell", b"-enc", b"frombase64string", b"invoke-expression", b"iex ", b"cmd.exe /c"):
        if s in low:
            flags.append(("نمط تنفيذ مشبوه: " if ar else "Suspicious exec pattern: ") + s.decode(errors="replace"))
    return flags[:12]


def _vt_cfg():
    return _SETTINGS.get("virustotal") or {"enabled": False, "api_key": ""}


# ==================================================================
#  YARA — local rule-based file scanning (no network ever, all on-device)
#  Loads .yar/.yara rule files from sentinel_rules/  (user can drop any
#  rules in there). yara-python is OPTIONAL; if not installed we fall
#  back to a small built-in keyword scanner over the file's bytes.
# ==================================================================
_YARA_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_rules")
_yara_compiled = [None]
_yara_lock = threading.Lock()


def _yara_load():
    """Compile every .yar/.yara file in sentinel_rules/ (and the community/ subdir).
    Returns (count, error)."""
    if not os.path.isdir(_YARA_DIR):
        try:
            os.makedirs(_YARA_DIR, exist_ok=True)
        except Exception:
            pass
        return (0, None)
    try:
        # yara-python may print a noisy "Failed to import libyara.dll" to stderr
        # on broken installs. Temporarily silence the C-level stderr fd during the
        # import only, then restore it (so logging keeps working afterwards).
        import yara             # yara-python (optional dep)
    except Exception:
        return (0, "yara-python not installed")
    files = {}
    # user's own rules (highest priority — never overwritten by updates)
    for fn in os.listdir(_YARA_DIR):
        full = os.path.join(_YARA_DIR, fn)
        if os.path.isfile(full) and fn.lower().endswith((".yar", ".yara")):
            files[fn] = full
    # community rules (auto-updated from public mirrors)
    comm = os.path.join(_YARA_DIR, "community")
    if os.path.isdir(comm):
        for fn in os.listdir(comm):
            full = os.path.join(comm, fn)
            if os.path.isfile(full) and fn.lower().endswith((".yar", ".yara")):
                files["community/" + fn] = full
    if not files:
        return (0, None)
    # Compile each rule file INDIVIDUALLY first, so a single malformed community
    # rule can't take down the whole set (previously one bad file => 0 rules).
    # We keep only the files that compile cleanly, then compile them together.
    good = {}
    skipped = []
    for ns, path in files.items():
        try:
            yara.compile(filepath=path)   # validate this file alone
            good[ns] = path
        except Exception as e:
            skipped.append(f"{os.path.basename(path)}: {str(e)[:60]}")
    if not good:
        return (0, "no valid rules" + ("; " + "; ".join(skipped) if skipped else ""))
    try:
        with _yara_lock:
            _yara_compiled[0] = yara.compile(filepaths=good)
        # remember which were skipped (for the status endpoint / diagnostics)
        _yara_update_state["skipped"] = skipped
        return (len(good), ("skipped " + str(len(skipped)) + " invalid rule(s)") if skipped else None)
    except Exception as e:
        return (0, str(e)[:200])


# ---- periodic community-rule updater (privacy-first) ----
# A small curated list of public YARA rule files mirrored on GitHub.
# All downloads are anonymous one-way GETs from a general-purpose CDN
# (raw.githubusercontent.com). No host info is ever sent. Respects
# SENTINEL_OFFLINE. Files land in sentinel_rules/community/ which never
# overwrites the user's own rules in sentinel_rules/.
_YARA_COMMUNITY_URLS = os.environ.get("SENTINEL_YARA_URLS", ",".join([
    "https://raw.githubusercontent.com/Yara-Rules/rules/master/malware/RANSOM_Cerber.yar",
    "https://raw.githubusercontent.com/Yara-Rules/rules/master/malware/MALW_Emotet.yar",
    "https://raw.githubusercontent.com/Yara-Rules/rules/master/exploit_kits/EK_Angler.yar",
    "https://raw.githubusercontent.com/Yara-Rules/rules/master/webshells/WShell_ASPXSpy.yar",
    "https://raw.githubusercontent.com/Neo23x0/signature-base/master/yara/gen_powershell_susp.yar",
    "https://raw.githubusercontent.com/Neo23x0/signature-base/master/yara/gen_mimikatz.yar",
    "https://raw.githubusercontent.com/Neo23x0/signature-base/master/yara/expl_log4j_cve_2021_44228.yar",
    "https://raw.githubusercontent.com/Neo23x0/signature-base/master/yara/gen_webshells.yar",
    "https://raw.githubusercontent.com/Neo23x0/signature-base/master/yara/gen_cobaltstrike.yar",
])).split(",")
_YARA_UPDATE_HOURS = int(os.environ.get("SENTINEL_YARA_UPDATE_HOURS", "24"))
_yara_update_state = {"ts": 0, "ok": 0, "fail": 0, "files": [], "last_error": ""}
_yara_thread = [None]


def update_yara_rules(force=False):
    """One-shot community-rule refresh. Privacy: anonymous GETs only, generic UA,
    no host data sent. No-op when offline. Returns a status dict."""
    if _OFFLINE:
        return {"status": "offline"}
    # do not refresh more than once per window unless forced
    if not force and time.time() * 1000 - _yara_update_state["ts"] < _YARA_UPDATE_HOURS * 3600 * 1000:
        return {"status": "skipped", "reason": "still fresh"}
    comm_dir = os.path.join(_YARA_DIR, "community")
    try:
        os.makedirs(comm_dir, exist_ok=True)
    except Exception:
        pass
    ok = 0; fail = 0; files = []; last_err = ""
    for url in _YARA_COMMUNITY_URLS:
        url = url.strip()
        if not url:
            continue
        try:
            req = _ur.Request(url, headers={"User-Agent": "Sentinel-SOC"})
            with _ur.urlopen(req, timeout=20) as r:
                body = r.read()
            # only accept text that looks like a yara rule (prevents writing junk)
            if b"rule " not in body[:4096]:
                fail += 1; last_err = f"not a yara rule: {url}"; continue
            fname = os.path.basename(url.split("?")[0]) or "rule.yar"
            with open(os.path.join(comm_dir, fname), "wb") as f:
                f.write(body)
            files.append(fname); ok += 1
        except Exception as e:
            fail += 1; last_err = str(e)[:160]
    _yara_update_state.update(ts=int(time.time() * 1000), ok=ok, fail=fail,
                              files=files, last_error=last_err)
    # recompile so the new community rules are active immediately
    count, err = _yara_load()
    return {"status": "ok" if ok else "error", "downloaded": ok, "failed": fail,
            "loaded": count, "files": files, "error": last_err or err}


def _yara_update_loop():
    """Refresh community rules on startup (if stale) then every N hours."""
    if _OFFLINE:
        return
    # initial refresh after a short delay so startup isn't blocked
    time.sleep(45)
    try:
        update_yara_rules()
    except Exception:
        pass
    while True:
        time.sleep(max(3600, _YARA_UPDATE_HOURS * 3600))
        try:
            update_yara_rules()
        except Exception:
            pass


def _ensure_yara_updater():
    if _yara_thread[0] is None and not _OFFLINE:
        t = threading.Thread(target=_yara_update_loop, daemon=True)
        _yara_thread[0] = t
        t.start()


def _yara_match(raw):
    """Scan a buffer against the loaded rules. Returns a list of rule names."""
    if _yara_compiled[0] is None:
        return []
    try:
        with _yara_lock:
            matches = _yara_compiled[0].match(data=raw)
        return [m.rule for m in matches][:20]
    except Exception:
        return []


def _yara_fallback_available():
    """The built-in keyword fallback engine is always usable (pure Python, no
    native deps), so YARA scanning works even when libyara/yara-python is broken."""
    return True


def _yara_fallback(raw):
    """No yara-python? Fall back to a tiny built-in keyword scan that mimics
    common malware-rule patterns. Privacy: still 100% local."""
    rules = {
        "Suspicious_Base64_PowerShell": (b"-enc ", b"frombase64string", b"powershell -e"),
        "Process_Injection_API_Set": (b"virtualallocex", b"writeprocessmemory", b"createremotethread"),
        "Network_Downloader_API": (b"urldownloadtofile", b"winhttp", b"internetopen"),
        "Office_Macro_Marker": (b"vbaproject.bin", b"autoopen", b"document_open"),
        "Mimikatz_Strings": (b"sekurlsa::logonpasswords", b"mimikatz", b"gentilkiwi"),
        "Cobalt_Strike_Beacon": (b"beacon", b"cobaltstrike", b"smb_beacon"),
        "Ransomware_Hints": (b"your files have been encrypted", b"bitcoin", b"readme.txt"),
        "Shellcode_Hex_Marker": (b"\\x90\\x90\\x90\\x90", b"\\xeb\\xfe"),
    }
    low = raw[:262144].lower()
    hits = []
    for name, needles in rules.items():
        if any(n in low for n in needles):
            hits.append(name)
    return hits


# load whatever rules are present at import time
_yara_load()


# ==================================================================
#  DETECTION ENGINE #1 — LIVE PROCESS SCANNING
#  Scans running processes for: malicious on-disk images (YARA against the
#  executable file), suspicious paths/names, unsigned binaries in user-writable
#  locations, and known living-off-the-land abuse. Privacy: 100% local.
# ==================================================================
_PROC_SUSPICIOUS_DIRS = ("\\temp\\", "\\appdata\\local\\temp", "\\downloads\\",
                         "\\$recycle", "\\programdata\\", "\\users\\public\\",
                         "\\windows\\temp")
# legit system tools frequently abused ("living off the land binaries")
_LOLBINS = {"powershell.exe", "cmd.exe", "wscript.exe", "cscript.exe", "mshta.exe",
            "rundll32.exe", "regsvr32.exe", "certutil.exe", "bitsadmin.exe",
            "msbuild.exe", "installutil.exe", "wmic.exe", "schtasks.exe"}
# names that should ONLY ever run from System32 — elsewhere = masquerading
_SYSTEM32_ONLY = {"svchost.exe", "lsass.exe", "csrss.exe", "winlogon.exe",
                  "services.exe", "smss.exe", "wininit.exe", "spoolsv.exe",
                  "explorer.exe", "taskhostw.exe"}

# TRUSTED publishers/paths — processes here are NOT flagged. These are the
# signed Windows + Microsoft Defender + common vendor locations. This kills the
# false positives where explorer.exe / MsMpEng.exe / NisSrv.exe were marked
# "suspicious" just for existing.
_TRUSTED_PATH_MARKERS = (
    "\\windows\\system32\\", "\\windows\\syswow64\\", "\\windows\\explorer.exe",
    "\\windows\\winsxs\\", "\\windowsapps\\", "\\program files\\",
    "\\program files (x86)\\", "\\programdata\\microsoft\\windows defender\\",
    "\\windows\\microsoft.net\\", "\\windows\\servicing\\",
)
# Microsoft/Defender executables that are legitimate even from ProgramData
_TRUSTED_NAMES = {
    "explorer.exe", "msmpeng.exe", "nissrv.exe", "mpdefendercoreservice.exe",
    "mpcmdrun.exe", "securityhealthservice.exe", "smartscreen.exe",
    "searchindexer.exe", "runtimebroker.exe", "dllhost.exe", "sihost.exe",
    "ctfmon.exe", "fontdrvhost.exe", "dwm.exe", "wmiprvse.exe",
    "msedge.exe", "msedgewebview2.exe", "onedrive.exe", "teams.exe",
}


def _is_trusted_proc(name_l, exe_l):
    """A process is trusted if it's a known Microsoft binary OR it lives in a
    standard signed-system/Program Files location. Trusted = never flagged."""
    if name_l in _TRUSTED_NAMES:
        return True
    if exe_l and any(m in exe_l for m in _TRUSTED_PATH_MARKERS):
        return True
    return False


# Authenticode signature cache: path -> (status, signer). Verifying a signature
# is relatively expensive, so we cache per path for the life of the process.
_sig_cache = {}
# Well-known trusted signers — a valid signature from any of these = legitimate.
_TRUSTED_SIGNERS = (
    "microsoft", "google", "mozilla", "intel", "nvidia", "amd ", "realtek",
    "dell", "hewlett", "hp inc", "lenovo", "asus", "logitech", "adobe",
    "apple", "valve", "discord", "spotify", "dropbox", "oracle", "python",
    "git ", "github", "jetbrains", "cisco", "vmware", "citrix", "zoom",
)


def _verify_signature(exe):
    """Return (is_signed_valid, signer_name). Uses PowerShell Get-AuthenticodeSignature
    on Windows. Cached per path. Never raises. The strongest single anti-false-positive
    signal: a binary with a VALID signature from a known publisher is trustworthy."""
    if os.name != "nt" or not exe:
        return (False, "")
    if exe in _sig_cache:
        return _sig_cache[exe]
    result = (False, "")
    try:
        out = _run_cmd(["powershell", "-NoProfile", "-Command",
                        f"$s=Get-AuthenticodeSignature -LiteralPath '{exe}'; "
                        f"\"$($s.Status)|$($s.SignerCertificate.Subject)\""], 12)
        if out and "|" in out:
            status, _, subject = out.strip().partition("|")
            status = status.strip().lower()
            subject_l = subject.strip().lower()
            if status == "valid":
                signer = subject_l
                # extract CN if present for a cleaner signer label
                import re as _r4
                cn = _r4.search(r"cn=([^,]+)", subject_l)
                signer_name = (cn.group(1).strip() if cn else subject_l)[:80]
                trusted = any(ts in subject_l for ts in _TRUSTED_SIGNERS)
                result = (trusted, signer_name)
    except Exception:
        result = (False, "")
    _sig_cache[exe] = result
    return result


def _proc_is_trustworthy(name_l, exe_l, exe):
    """Layered trust decision used by every detection engine to suppress false
    positives. Order: (1) known MS name, (2) trusted system path,
    (3) VALID Authenticode signature from a recognised publisher."""
    if _is_trusted_proc(name_l, exe_l):
        return True
    signed, _signer = _verify_signature(exe)
    return signed


def scan_processes(limit=600):
    """Inspect every running process for malware indicators. Returns a list of
    {pid, name, path, sev, reasons[], yara[]}. Local-only; no data leaves the box."""
    if not HAVE_PSUTIL:
        return {"available": False, "reason": "psutil not installed", "hits": []}
    hits = []
    seen = 0
    for p in psutil.process_iter(["pid", "name", "exe", "username", "cmdline", "ppid"]):
        if seen >= limit:
            break
        seen += 1
        try:
            info = p.info
            name = (info.get("name") or "").lower()
            exe = info.get("exe") or ""
            exe_l = exe.lower()
            cmd = " ".join(info.get("cmdline") or []).lower()
            reasons = []
            yhits = []
            sev = "info"

            # TRUST GATE: skip signed Microsoft/Windows binaries entirely.
            # This removes the false positives (explorer.exe, MsMpEng.exe, etc.).
            if _is_trusted_proc(name, exe_l):
                continue

            # 1) masquerading: a System32-only name running from a NON-system path
            if name in _SYSTEM32_ONLY and exe_l and not _is_trusted_proc(name, exe_l):
                reasons.append(f"'{name}' running outside System32 ({exe}) — possible masquerading")
                sev = "critical"

            # 2) executable living in a user-writable / temp location
            if exe_l and any(d in exe_l for d in _PROC_SUSPICIOUS_DIRS):
                reasons.append(f"Executable in a suspicious location: {exe}")
                sev = _max_sev(sev, "high")

            # 3) LOLBin with an encoded / download command line
            if name in _LOLBINS:
                if any(k in cmd for k in ("-enc", "-e ", "frombase64string", "downloadstring",
                                          "iex ", "invoke-expression", "-windowstyle hidden",
                                          "bypass", "webclient", "bitstransfer", "certutil -urlcache")):
                    reasons.append(f"Suspicious {name} command line (encoded/download/hidden)")
                    sev = _max_sev(sev, "high")

            # 4) YARA scan of the on-disk image — only flag a REAL rule match,
            #    and only when the file isn't in a trusted location.
            if exe and not _is_trusted_proc(name, exe_l):
                try:
                    if os.path.isfile(exe) and os.path.getsize(exe) <= 12_000_000:
                        with open(exe, "rb") as fh:
                            raw = fh.read()
                        yhits = _yara_match(raw) or (_yara_fallback(raw) if _yara_compiled[0] is None else [])
                        if yhits:
                            reasons.append(f"YARA matched the executable: {', '.join(yhits[:4])}")
                            sev = _max_sev(sev, "high" if len(yhits) < 2 else "critical")
                except Exception:
                    pass

            if reasons:
                # FINAL ANTI-FALSE-POSITIVE GATE: before reporting, check the
                # Authenticode signature. A validly-signed binary from a known
                # publisher is almost never malware. We DROP the finding unless
                # YARA explicitly matched (a signed file matching a malware rule
                # is worth surfacing) or the file sits in a temp/download path.
                signed, signer = _verify_signature(exe)
                in_temp = exe_l and any(d in exe_l for d in _PROC_SUSPICIOUS_DIRS)
                if signed and not yhits and not in_temp:
                    continue                      # trusted publisher -> suppress
                if signed and (yhits or in_temp):
                    # signed but still odd -> keep, but downgrade and note signer
                    sev = _max_sev("medium", sev) if yhits else "medium"
                    reasons.append(f"(signed by {signer})")
                hits.append({"pid": info.get("pid"), "name": info.get("name"),
                             "path": exe, "user": info.get("username"),
                             "sev": sev, "reasons": reasons, "yara": yhits,
                             "signed": signed})
        except (psutil.NoSuchProcess, psutil.AccessDenied, Exception):
            continue
    # log the worst finding as a real event
    if hits:
        worst = max(hits, key=lambda h: _SEV_WEIGHT.get(h["sev"], 0))
        log_event("process_threat", sev=worst["sev"], asset=_plat.node() or "host",
                  text_en=f"Suspicious process: {worst['name']} (pid {worst['pid']}) — {worst['reasons'][0]}",
                  text_ar=f"عملية مشبوهة: {worst['name']} (pid {worst['pid']}) — {worst['reasons'][0]}")
    return {"available": True, "scanned": seen, "hits": hits,
            "ts": int(time.time() * 1000)}


def _max_sev(a, b):
    return a if _SEV_WEIGHT.get(a, 0) >= _SEV_WEIGHT.get(b, 0) else b


# ==================================================================
#  DETECTION ENGINE #2 — PERSISTENCE / AUTORUN AUDIT
#  Enumerates the common autostart locations attackers use to survive reboot:
#  Run keys, Startup folders, Scheduled Tasks, Services. Flags entries that run
#  from suspicious paths or use encoded commands. Privacy: 100% local.
# ==================================================================
def scan_persistence():
    """Audit autostart locations for suspicious persistence. Local-only."""
    if os.name != "nt":
        return {"available": False, "reason": "Windows only", "items": []}
    items = []

    def _flag(source, name, value):
        v = (value or "").lower()
        # skip entries that live in trusted signed locations (Program Files,
        # Windows, Microsoft Defender, etc.) — these are legitimate autostarts.
        if any(m in v for m in _TRUSTED_PATH_MARKERS):
            return
        # extract the executable path from the command line to verify its signature
        exe_path = ""
        try:
            raw = (value or "").strip()
            if raw.startswith('"'):
                exe_path = raw[1:].split('"', 1)[0]
            else:
                exe_path = raw.split()[0] if raw.split() else ""
        except Exception:
            exe_path = ""
        # STRONGEST anti-false-positive: if the target executable carries a VALID
        # signature from a known publisher (Discord, Spotify, NVIDIA, Microsoft…),
        # it is a legitimate autostart even from AppData — don't flag it.
        if exe_path:
            signed, _sg = _verify_signature(exe_path)
            if signed:
                return
        # respect the user's allowlist: apps they marked "Safe" are skipped
        try:
            nm_l = (name or "").lower()
            if nm_l in _rt_allowlist["names"] or (exe_path and any(
                    p and p in exe_path.lower() for p in _rt_allowlist["paths"])):
                return
        except Exception:
            pass
        reasons = []
        # a script host alone isn't suspicious; it must ALSO be encoded/downloading
        bad_combo = any(k in v for k in ("-enc", "-e ", "frombase64string",
                                         "downloadstring", "iex ", "invoke-expression",
                                         "regsvr32 /i:http", "certutil -urlcache",
                                         "bitsadmin /transfer", "-windowstyle hidden"))
        in_temp = any(d in v for d in _PROC_SUSPICIOUS_DIRS)
        if in_temp:
            reasons.append("runs from a user-writable/temp path")
        if bad_combo:
            reasons.append("uses an encoded / download command")
        # only flag when there's a REAL indicator (temp path or bad command),
        # and downgrade to medium unless both are present
        if reasons:
            sev = "high" if (in_temp and bad_combo) else "medium"
            items.append({"source": source, "name": name, "command": value[:300],
                          "reasons": reasons, "sev": sev, "exe": exe_path,
                          "unsigned": True})

    # 1) Run / RunOnce keys (HKLM + HKCU)
    try:
        import winreg
        runkeys = [
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\Run"),
            (winreg.HKEY_LOCAL_MACHINE, r"SOFTWARE\Microsoft\Windows\CurrentVersion\RunOnce"),
            (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows\CurrentVersion\Run"),
            (winreg.HKEY_CURRENT_USER, r"SOFTWARE\Microsoft\Windows\CurrentVersion\RunOnce"),
        ]
        for hive, path in runkeys:
            try:
                with winreg.OpenKey(hive, path) as key:
                    i = 0
                    while True:
                        try:
                            nm, val, _ = winreg.EnumValue(key, i)
                            _flag("Run key", nm, str(val))
                            i += 1
                        except OSError:
                            break
            except Exception:
                continue
    except Exception:
        pass

    # 2) Scheduled tasks with suspicious actions
    try:
        out = _run_cmd(["powershell", "-NoProfile", "-Command",
                        "Get-ScheduledTask | ForEach-Object { $a=$_.Actions | "
                        "Where-Object {$_.Execute}; if($a){ \"$($_.TaskName)|$($a.Execute) $($a.Arguments)\" } }"], 30)
        for line in (out or "").splitlines():
            if "|" in line:
                nm, _, val = line.partition("|")
                _flag("Scheduled Task", nm.strip(), val.strip())
    except Exception:
        pass

    # 3) Services pointing at user-writable binaries
    try:
        out = _run_cmd(["powershell", "-NoProfile", "-Command",
                        "Get-CimInstance Win32_Service | ForEach-Object { \"$($_.Name)|$($_.PathName)\" }"], 30)
        for line in (out or "").splitlines():
            if "|" in line:
                nm, _, val = line.partition("|")
                _flag("Service", nm.strip(), val.strip())
    except Exception:
        pass

    if items:
        worst = items[0]
        log_event("persistence", sev="high", asset=_plat.node() or "host",
                  text_en=f"Suspicious persistence: {worst['source']} '{worst['name']}'",
                  text_ar=f"استمرارية مشبوهة: {worst['source']} '{worst['name']}'")
    return {"available": True, "items": items, "ts": int(time.time() * 1000)}


# ==================================================================
#  DETECTION ENGINE #3 — FILE INTEGRITY MONITORING (FIM)
#  Baselines hashes of critical files; on re-check, reports any that changed,
#  appeared, or vanished. Catches tampering of system/host files. The baseline
#  is stored encrypted locally. Privacy: 100% local.
# ==================================================================
_FIM_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_fim.json")


def _fim_targets():
    """The files we watch by default — high-value OS tamper targets only.
    NOTE: we deliberately do NOT watch the app's own source file — it changes
    every time the user installs an update, which would create a permanent false
    'integrity change' alert."""
    t = []
    if os.name == "nt":
        sysroot = os.environ.get("SystemRoot", "C:\\Windows")
        t += [os.path.join(sysroot, "System32", "drivers", "etc", "hosts"),
              os.path.join(sysroot, "System32", "cmd.exe"),
              os.path.join(sysroot, "System32", "kernel32.dll"),
              os.path.join(sysroot, "explorer.exe")]
    else:
        t += ["/etc/hosts", "/etc/passwd", "/etc/shadow", "/bin/bash"]
    return [p for p in t if os.path.isfile(p)]


def fim_baseline(extra_paths=None):
    """Record current hashes as the trusted baseline."""
    targets = _fim_targets() + list(extra_paths or [])
    base = {}
    for p in targets:
        try:
            with open(p, "rb") as f:
                base[p] = hashlib.sha256(f.read()).hexdigest()
        except Exception:
            base[p] = None
    _secure_save(_FIM_FILE, {"ts": int(time.time() * 1000), "hashes": base})
    return {"ok": True, "count": len([v for v in base.values() if v]), "ts": base and int(time.time()*1000)}


def fim_check():
    """Compare current hashes against the baseline. Reports changed/new/missing."""
    data = _secure_load(_FIM_FILE, None)
    if not data or not data.get("hashes"):
        return {"available": True, "baselined": False, "changes": [],
                "note": "No baseline yet — create one first."}
    changes = []
    for p, old in data["hashes"].items():
        try:
            if not os.path.isfile(p):
                if old is not None:
                    changes.append({"path": p, "kind": "missing", "sev": "high"})
                continue
            with open(p, "rb") as f:
                new = hashlib.sha256(f.read()).hexdigest()
            if old is None:
                changes.append({"path": p, "kind": "appeared", "sev": "medium"})
            elif new != old:
                changes.append({"path": p, "kind": "modified", "sev": "high",
                                "old": old[:12], "new": new[:12]})
        except Exception:
            continue
    if changes:
        log_event("fim", sev="high", asset=_plat.node() or "host",
                  text_en=f"File integrity: {len(changes)} watched file(s) changed",
                  text_ar=f"سلامة الملفات: تغيّر {len(changes)} ملف مراقَب")
    return {"available": True, "baselined": True, "changes": changes,
            "baseline_ts": data.get("ts"), "ts": int(time.time() * 1000)}


# ==================================================================
#  DETECTION ENGINE #4 — NETWORK CONNECTION ANALYSIS (C2 / beaconing)
#  Inspects active TCP connections and maps each to its owning process.
#  Flags: connections to suspicious ports, processes that "shouldn't" talk to
#  the internet, many connections to one remote (beaconing), and listeners on
#  unusual ports. Privacy: 100% local — no IP is ever sent anywhere.
# ==================================================================
# remote ports commonly used by malware C2 / RATs / miners
_SUSPECT_REMOTE_PORTS = {1080, 4444, 4445, 5555, 6666, 6667, 1337, 31337,
                         8443, 9001, 9050, 9051, 12345, 54321, 3333, 14444, 45560}
# processes that normally have NO reason to open internet connections
_NO_NET_PROCS = {"notepad.exe", "calc.exe", "mspaint.exe", "write.exe",
                 "winword.exe", "excel.exe", "powerpnt.exe"}


def scan_network_threats():
    """Analyse live connections for C2/beaconing indicators. Local-only."""
    if not HAVE_PSUTIL:
        return {"available": False, "reason": "psutil not installed", "hits": []}
    hits = []
    remote_counts = {}          # remote ip -> count (beaconing detection)
    try:
        conns = psutil.net_connections(kind="inet")
    except Exception:
        return {"available": False, "reason": "permission denied", "hits": []}

    pid_name = {}
    for c in conns:
        try:
            if not c.raddr or not c.raddr.ip:
                # listener on an unusual port?
                if c.status == "LISTEN" and c.laddr and c.laddr.port in _SUSPECT_REMOTE_PORTS:
                    nm = _pid_name(c.pid, pid_name)
                    hits.append({"sev": "high", "kind": "suspicious_listener",
                                 "detail": f"Process {nm} (pid {c.pid}) listening on suspicious port {c.laddr.port}",
                                 "pid": c.pid, "name": nm})
                continue
            rip = c.raddr.ip
            rport = c.raddr.port
            # skip private/loopback (internal traffic)
            klass = _ip_class(rip)
            if klass in ("private", "loopback"):
                continue
            remote_counts[rip] = remote_counts.get(rip, 0) + 1
            nm = _pid_name(c.pid, pid_name)

            # 1) connection to a known-suspect C2 port
            if rport in _SUSPECT_REMOTE_PORTS:
                hits.append({"sev": "critical", "kind": "c2_port",
                             "detail": f"{nm} (pid {c.pid}) -> {rip}:{rport} (known C2/RAT port)",
                             "pid": c.pid, "name": nm, "remote": f"{rip}:{rport}", "geo": klass})
            # 2) a process that should never reach the internet
            if (nm or "").lower() in _NO_NET_PROCS:
                hits.append({"sev": "high", "kind": "unexpected_network",
                             "detail": f"{nm} (pid {c.pid}) has an outbound connection to {rip} — unusual for this program",
                             "pid": c.pid, "name": nm, "remote": rip, "geo": klass})
        except Exception:
            continue

    # 3) beaconing: many separate connections to one remote host
    for rip, cnt in remote_counts.items():
        if cnt >= 20:
            hits.append({"sev": "medium", "kind": "beaconing",
                         "detail": f"{cnt} connections to a single remote host {rip} (possible beaconing)",
                         "remote": rip, "count": cnt, "geo": _ip_class(rip)})

    if hits:
        worst = max(hits, key=lambda h: _SEV_WEIGHT.get(h["sev"], 0))
        log_event("network_threat", sev=worst["sev"], asset=_plat.node() or "host",
                  text_en=f"Network: {worst['detail']}",
                  text_ar=f"الشبكة: {worst['detail']}")
    return {"available": True, "connections": len(conns), "hits": hits,
            "ts": int(time.time() * 1000)}


def _pid_name(pid, cache):
    if pid in cache:
        return cache[pid]
    nm = "?"
    try:
        if pid and HAVE_PSUTIL:
            nm = psutil.Process(pid).name()
    except Exception:
        nm = "?"
    cache[pid] = nm
    return nm


# ==================================================================
#  DETECTION ENGINE #5 — ACCOUNT & PRIVILEGE AUDIT
#  Looks for the account-level signs of compromise: unexpected admin accounts,
#  recently created users, accounts with non-expiring passwords, and members of
#  the local Administrators group. Privacy: 100% local.
# ==================================================================
def scan_accounts():
    """Audit local accounts for signs of compromise/backdoors. Windows-focused."""
    if os.name != "nt":
        return {"available": False, "reason": "Windows only", "items": []}
    items = []

    # 1) members of the local Administrators group
    try:
        out = _run_cmd(["net", "localgroup", "Administrators"], 12)
        admins = []
        if out:
            started = False
            for line in out.splitlines():
                line = line.strip()
                if line.startswith("---"):
                    started = True; continue
                if started and line and not line.lower().startswith("the command"):
                    admins.append(line)
        # flag if there are many admins or an unexpected one
        if len(admins) > 3:
            items.append({"sev": "medium", "kind": "many_admins",
                          "detail": f"{len(admins)} accounts are local Administrators: {', '.join(admins[:8])}",
                          "accounts": admins})
    except Exception:
        pass

    # 2) enabled accounts + password-never-expires (PowerShell)
    try:
        out = _run_cmd(["powershell", "-NoProfile", "-Command",
                        "Get-LocalUser | Where-Object {$_.Enabled} | "
                        "ForEach-Object { \"$($_.Name)|$($_.PasswordExpires)|$($_.LastLogon)\" }"], 20)
        for line in (out or "").splitlines():
            parts = line.split("|")
            if len(parts) >= 2:
                uname = parts[0].strip()
                pwd_expires = parts[1].strip()
                # built-in accounts that are enabled are worth noting
                if uname.lower() in ("guest", "defaultaccount", "wdagutilityaccount"):
                    items.append({"sev": "high", "kind": "builtin_enabled",
                                  "detail": f"Built-in account '{uname}' is ENABLED — usually should be disabled",
                                  "account": uname})
    except Exception:
        pass

    # 3) hidden admin accounts (name ending with $ is a classic hiding trick)
    try:
        out = _run_cmd(["powershell", "-NoProfile", "-Command",
                        "Get-LocalUser | ForEach-Object { $_.Name }"], 15)
        for line in (out or "").splitlines():
            nm = line.strip()
            if nm.endswith("$") and nm.lower() not in ("",):
                items.append({"sev": "high", "kind": "hidden_account",
                              "detail": f"Account '{nm}' ends with '$' (often used to hide accounts)",
                              "account": nm})
    except Exception:
        pass

    if items:
        worst = max(items, key=lambda h: _SEV_WEIGHT.get(h["sev"], 0))
        log_event("account_audit", sev=worst["sev"], asset=_plat.node() or "host",
                  text_en=f"Account audit: {worst['detail']}",
                  text_ar=f"تدقيق الحسابات: {worst['detail']}")
    return {"available": True, "items": items, "ts": int(time.time() * 1000)}


# ==================================================================
#  DETECTION ENGINE #6 — ROOTKIT DETECTION (no driver, no cloud, pure Python)
#  Uses two well-established, local-only techniques:
#
#  1) CROSS-VIEW DETECTION — the core idea behind RootkitRevealer/GMER.
#     A rootkit hides a process/port by hooking ONE enumeration API but rarely
#     ALL of them. We list processes/ports from independent Windows sources and
#     flag anything that appears in one view but is hidden from another.
#       - processes: psutil (NtQuerySystemInformation) vs `tasklist` (WMI/Win32)
#                    vs PowerShell Get-Process (CLR/WMI)
#       - ports:     psutil vs `netstat`
#
#  2) DRIVER SIGNATURE AUDIT — kernel rootkits load a driver. Legitimate drivers
#     on modern Windows are digitally signed; an UNSIGNED loaded driver, or one
#     in a user-writable path, is a strong rootkit indicator.
#
#  PRIVACY: 100% local. Reads only local OS state. Nothing is ever transmitted.
#  SAFETY:  read-only enumeration — never loads a driver, never patches memory.
# ==================================================================
def _procs_psutil():
    """PID set from psutil (uses NtQuerySystemInformation under the hood)."""
    s = {}
    if not HAVE_PSUTIL:
        return s
    for p in psutil.process_iter(["pid", "name"]):
        try:
            s[int(p.info["pid"])] = (p.info.get("name") or "").lower()
        except Exception:
            continue
    return s


def _procs_tasklist():
    """PID set from tasklist.exe (a different Win32/WMI code path)."""
    s = {}
    if os.name != "nt":
        return s
    out = _run_cmd(["tasklist", "/fo", "csv", "/nh"], 20)
    for line in (out or "").splitlines():
        # "name.exe","1234","Console","1","12,345 K"
        parts = [x.strip().strip('"') for x in line.split('","')]
        if len(parts) >= 2:
            try:
                name = parts[0].strip('"').lower()
                pid = int(parts[1])
                s[pid] = name
            except Exception:
                continue
    return s


def _procs_powershell():
    """PID set from PowerShell Get-Process (CLR/WMI code path)."""
    s = {}
    if os.name != "nt":
        return s
    out = _run_cmd(["powershell", "-NoProfile", "-Command",
                    "Get-Process | ForEach-Object { \"$($_.Id)|$($_.ProcessName)\" }"], 25)
    for line in (out or "").splitlines():
        if "|" in line:
            pid_s, _, nm = line.partition("|")
            try:
                s[int(pid_s.strip())] = nm.strip().lower()
            except Exception:
                continue
    return s


def _ports_psutil():
    if not HAVE_PSUTIL:
        return set()
    out = set()
    try:
        for c in psutil.net_connections(kind="inet"):
            if c.laddr:
                out.add(int(c.laddr.port))
    except Exception:
        pass
    return out


def _ports_netstat():
    if os.name != "nt":
        return set()
    out = set()
    res = _run_cmd(["netstat", "-ano"], 20)
    import re as _r3
    for line in (res or "").splitlines():
        m = _r3.search(r":(\d+)\s", line)
        if m:
            try:
                out.add(int(m.group(1)))
            except Exception:
                continue
    return out


def scan_rootkit():
    """Cross-view + driver-signature rootkit detection. Local & read-only."""
    if os.name != "nt":
        return {"available": False, "reason": "Windows only", "hidden_processes": [],
                "unsigned_drivers": [], "port_discrepancies": []}
    result = {"available": True, "ts": int(time.time() * 1000),
              "hidden_processes": [], "unsigned_drivers": [], "port_discrepancies": [],
              "views": {}}

    # ---- 1) CROSS-VIEW: processes (double-confirmed to avoid false positives) ----
    # A real rootkit hides a PID CONSISTENTLY. Processes that merely start/stop
    # between two enumerations would look "hidden" in a single snapshot — so we
    # take TWO passes a moment apart and only report a PID that is hidden from
    # the SAME tool BOTH times. This removes the explorer/Defender-style noise.
    def _collect():
        return [(n, v) for n, v in (("psutil", _procs_psutil()),
                                    ("tasklist", _procs_tasklist()),
                                    ("powershell", _procs_powershell())) if v]

    views = _collect()
    result["views"] = {n: len(v) for n, v in views}
    if len(views) >= 2:
        # first-pass suspects: present somewhere, missing from some tool
        ignore = {0, 4}
        all_pids = set()
        for _, v in views:
            all_pids |= set(v.keys())
        suspects = {}
        for pid in all_pids - ignore:
            missing = [n for n, v in views if pid not in v]
            present = [n for n, v in views if pid in v]
            if missing and present:
                suspects[pid] = (present, missing,
                                 next((v[pid] for _, v in views if pid in v), "?"))
        if suspects:
            time.sleep(0.8)                       # let transient processes settle
            views2 = _collect()
            v2map = {n: v for n, v in views2}
            for pid, (present, missing, name) in suspects.items():
                # confirm: STILL missing from the same tools AND still alive somewhere
                still_present = any(pid in v2map.get(n, {}) for n in present if n in v2map)
                still_missing = [n for n in missing if n in v2map and pid not in v2map[n]]
                # skip trusted Microsoft names — never a rootkit
                if name in _TRUSTED_NAMES:
                    continue
                # require the hidden PID to still be present in psutil specifically
                # (our most reliable source) — drops tasklist/powershell timing noise
                if "psutil" not in [n for n in present if pid in v2map.get(n, {})]:
                    continue
                if still_present and still_missing:
                    # verify the executable's signature before crying rootkit:
                    # a validly-signed process that's merely slow to enumerate is
                    # not a rootkit. Only an UNSIGNED hidden process is alarming.
                    try:
                        exe = psutil.Process(pid).exe()
                    except Exception:
                        exe = ""
                    signed, _sg = _verify_signature(exe) if exe else (False, "")
                    if signed:
                        continue                  # signed + visible to psutil -> not a rootkit
                    result["hidden_processes"].append({
                        "pid": pid, "name": name,
                        "seen_in": present, "hidden_from": still_missing,
                        "sev": "critical"})

    # ---- 2) CROSS-VIEW: listening ports ----
    p_psutil = _ports_psutil()
    p_netstat = _ports_netstat()
    if p_psutil and p_netstat:
        # a port netstat shows but psutil's API hides (or vice-versa)
        only_netstat = p_netstat - p_psutil
        only_psutil = p_psutil - p_netstat
        # small differences are normal (timing); flag only notable hidden ports
        for port in sorted(only_netstat):
            if port in _SUSPECT_REMOTE_PORTS:
                result["port_discrepancies"].append({
                    "port": port, "seen_in": "netstat", "hidden_from": "psutil API",
                    "sev": "high"})

    # ---- 3) DRIVER SIGNATURE AUDIT ----
    # list loaded drivers with their signing status via PowerShell (signtool-free)
    try:
        out = _run_cmd(["powershell", "-NoProfile", "-Command",
                        "Get-CimInstance Win32_SystemDriver | Where-Object {$_.State -eq 'Running'} | "
                        "ForEach-Object { $p=$_.PathName -replace '\\\\\\?\\?\\\\',''; "
                        "if($p){ try{ $sig=(Get-AuthenticodeSignature $p).Status }catch{ $sig='Unknown' }; "
                        "\"$($_.Name)|$p|$sig\" } }"], 60)
        for line in (out or "").splitlines():
            parts = line.split("|")
            if len(parts) >= 3:
                name, path, status = parts[0].strip(), parts[1].strip(), parts[2].strip()
                pl = path.lower()
                sl = status.lower()
                # drivers in trusted system locations are legitimate even if the
                # signature check returns UnknownError (common for catalog-signed
                # Windows drivers) — only flag those OUTSIDE trusted paths.
                in_trusted = any(m in pl for m in _TRUSTED_PATH_MARKERS) or "\\windows\\system32\\drivers\\" in pl
                # report only a clearly-bad signature (not signed / tampered)
                if sl in ("notsigned", "hashmismatch"):
                    result["unsigned_drivers"].append({
                        "name": name, "path": path, "signature": status,
                        "sev": "critical" if not in_trusted else "high"})
                # driver running from a user-writable/temp path (very abnormal)
                elif any(d in pl for d in _PROC_SUSPICIOUS_DIRS):
                    result["unsigned_drivers"].append({
                        "name": name, "path": path, "signature": "suspicious-path",
                        "sev": "critical"})
    except Exception:
        pass

    total = (len(result["hidden_processes"]) + len(result["unsigned_drivers"])
             + len(result["port_discrepancies"]))
    result["total"] = total
    if total:
        sev = "critical" if (result["hidden_processes"] or
              any(d["sev"] == "critical" for d in result["unsigned_drivers"])) else "high"
        log_event("rootkit", sev=sev, asset=_plat.node() or "host",
                  text_en=f"Rootkit indicators: {len(result['hidden_processes'])} hidden process(es), "
                          f"{len(result['unsigned_drivers'])} suspicious driver(s)",
                  text_ar=f"مؤشّرات Rootkit: {len(result['hidden_processes'])} عملية مخفية، "
                          f"{len(result['unsigned_drivers'])} درايفر مشبوه")
    return result


def _vt_lookup(sha256):
    """OPTIONAL VirusTotal lookup. Privacy: sends ONLY the hash, never the file, and
    only when the user has explicitly enabled it with their own API key. Off by default."""
    cfg = _vt_cfg()
    if _OFFLINE or not cfg.get("enabled") or not cfg.get("api_key"):
        return None
    try:
        req = _ur.Request("https://www.virustotal.com/api/v3/files/" + sha256,
                          headers={"x-apikey": cfg["api_key"], "User-Agent": "Sentinel-SOC"})
        with _ur.urlopen(req, timeout=15) as r:
            data = json.loads(r.read().decode("utf-8", "replace"))
        stats = data.get("data", {}).get("attributes", {}).get("last_analysis_stats", {})
        return {"malicious": stats.get("malicious", 0), "suspicious": stats.get("suspicious", 0),
                "harmless": stats.get("harmless", 0), "undetected": stats.get("undetected", 0)}
    except Exception as e:
        return {"error": str(e)[:120]}


def analyze_file(filename, content, lang, raw=None):
    h = _log_heuristic_scan(content)
    snippet = content[:6000]   # keep the model prompt bounded
    model_text = None
    prompt = (("حلّل ملف السجل/البيانات التالي أمنياً. لخّص أبرز الأحداث، التهديدات المحتملة، "
               "والتوصيات في نقاط موجزة:\n\n" if lang == "ar"
               else "Analyze the following security log/data file. Summarize key events, "
                    "potential threats, and recommendations concisely:\n\n") + snippet)
    res = model_chat(prompt, lang)
    if res.get("source") == "live" and res.get("answer"):
        model_text = res["answer"].strip()

    ar = lang == "ar"
    top_words = ", ".join(f"{k}×{v}" for k, v in sorted(h["hits"].items(), key=lambda x: -x[1])[:6]) or ("لا شيء" if ar else "none")
    if ar:
        summary = (model_text or
                   f"فحصتُ «{filename}»: {h['lines']} سطر. مؤشرات لافتة: {top_words}. "
                   f"عناوين IP مرصودة: {len(h['ips'])}.")
        sections = [
            {"h": "نظرة عامة", "b": f"الملف «{filename}» · {h['lines']} سطر · {len(h['ips'])} عنوان IP."},
            {"h": "مؤشرات أمنية", "b": f"كلمات حرجة: {top_words}."},
            {"h": "عناوين IP", "b": ("، ".join(h["ips"]) or "لا يوجد")},
        ]
    else:
        summary = (model_text or
                   f"Scanned \u201c{filename}\u201d: {h['lines']} lines. Notable indicators: {top_words}. "
                   f"IP addresses seen: {len(h['ips'])}.")
        sections = [
            {"h": "Overview", "b": f"File \u201c{filename}\u201d · {h['lines']} lines · {len(h['ips'])} IPs."},
            {"h": "Security indicators", "b": f"Critical keywords: {top_words}."},
            {"h": "IP addresses", "b": (", ".join(h["ips"]) or "none")},
        ]
    if model_text:
        sections.insert(0, {"h": ("تحليل النموذج" if ar else "Model analysis"), "b": model_text})

    # ---- privacy-first local file forensics (hashing, entropy, type, indicators) ----
    rawb = raw if raw is not None else content.encode("utf-8", "replace")
    hashes = _file_hashes(rawb)
    ent = _entropy(rawb)
    ftype = _file_type(rawb, filename)
    bin_flags = _binary_indicators(rawb, ftype, lang)
    ffacts = (f"النوع: {ftype} · الحجم: {len(rawb)} بايت · الإنتروبيا: {ent}/8"
              if ar else f"Type: {ftype} · Size: {len(rawb)} bytes · Entropy: {ent}/8")
    if ent > 7.2:
        ffacts += (" — إنتروبيا عالية (قد يكون مضغوطاً/مشفّراً/محزَّماً)" if ar
                   else " — high entropy (possibly packed/encrypted/compressed)")
    sections.append({"h": ("الخصائص المحلية" if ar else "Local file properties"), "b": ffacts})
    sections.append({"h": ("البصمات (hash)" if ar else "Hashes"),
                     "b": f"SHA-256: {hashes['sha256']}\nSHA-1: {hashes['sha1']}\nMD5: {hashes['md5']}"})
    if bin_flags:
        sections.append({"h": ("مؤشرات ثابتة" if ar else "Static indicators"), "b": "\n".join("• " + f for f in bin_flags)})

    # ---- YARA rule matches (local, no network) ----
    yara_hits = _yara_match(rawb)
    yara_engine = "yara-python"
    if not yara_hits and _yara_compiled[0] is None:
        yara_hits = _yara_fallback(rawb)
        yara_engine = "built-in"
    if yara_hits:
        sections.append({"h": ("قواعد YARA المُطابقة" if ar else f"YARA matches"),
                         "b": (("(محرّك مبسّط)\n" if yara_engine == "built-in" and not ar
                                else ("(محرّك مبسّط)\n" if yara_engine == "built-in" else ""))
                               + "\n".join("• " + r for r in yara_hits))})
        if len(yara_hits) >= 2:
            h["tone"] = "danger"

    # ---- behavioral/heuristic risk score (combines signals) ----
    risk = 0
    risk_reasons = []
    if ent > 7.2:
        risk += 2; risk_reasons.append("إنتروبيا عالية (محزَّم/مشفّر)" if ar else "high entropy (packed/encrypted)")
    if len(yara_hits) >= 1:
        risk += 3 * len(yara_hits); risk_reasons.append("مطابقة قواعد YARA" if ar else "YARA rule match")
    # suspicious indicators inside binaries (imports/strings) raise the score
    susp_markers = ("createremotethread", "virtualallocex", "writeprocessmemory",
                    "urldownloadtofile", "frombase64string", "-enc ", "powershell -e",
                    "wscript.shell", "schtasks", "reg add", "vssadmin delete")
    low = rawb[:262144].lower()
    hit_markers = [m for m in susp_markers if m.encode() in low]
    if hit_markers:
        risk += len(hit_markers); risk_reasons.append((f"{len(hit_markers)} استدعاء/أمر مشبوه" if ar else f"{len(hit_markers)} suspicious API/command"))
    # double extension or script-as-document tricks
    fl = filename.lower()
    if any(fl.endswith(d) for d in (".pdf.exe", ".doc.exe", ".jpg.exe", ".scr", ".pif", ".vbs.exe")):
        risk += 4; risk_reasons.append("امتداد مزدوج/خادع" if ar else "double/deceptive extension")
    if risk >= 6:
        h["tone"] = "danger"
    elif risk >= 3 and h["tone"] == "success":
        h["tone"] = "warning"
    if risk_reasons:
        verdict_txt = ("درجة الخطورة التقديرية: " + str(risk) + "/10\n" if ar else "Estimated risk score: " + str(risk) + "/10\n")
        sections.append({"h": ("التقييم السلوكي" if ar else "Heuristic assessment"),
                         "b": verdict_txt + "\n".join("• " + r for r in risk_reasons)})

    # ---- OPTIONAL VirusTotal (hash-only, opt-in, user's key) ----
    vt = _vt_lookup(hashes["sha256"])
    if vt and "error" not in vt:
        mal = vt.get("malicious", 0)
        sections.append({"h": ("VirusTotal (بصمة فقط)" if ar else "VirusTotal (hash-only)"),
                         "b": (f"محرّكات صنّفته خبيثاً: {mal} · مشبوه: {vt.get('suspicious',0)} · سليم: {vt.get('harmless',0)}"
                               if ar else
                               f"Engines flagging malicious: {mal} · suspicious: {vt.get('suspicious',0)} · harmless: {vt.get('harmless',0)}")})
        if mal and mal > 0:
            h["tone"] = "danger"

    _rep_seq[0] += 1
    rid = f"rep-scan-{_rep_seq[0]}"
    rep = {"id": rid, "period": "scan", "tone": h["tone"],
           "score": posture_score(),
           "title": (f"تقرير فحص: {filename}" if ar else f"Scan report: {filename}"),
           "summary": summary, "sections": sections, "ts": int(time.time() * 1000),
           "byModel": bool(model_text)}
    _REPORTS[rid] = rep
    _save_reports()
    log_event("scan", sev=h["tone"], asset=filename,
              text_en=f"File scan: {filename} ({h['tone']})",
              text_ar=f"فحص ملف: {filename} ({h['tone']})")
    add_notification("report" if h["tone"] != "danger" else "alert",
                     h["tone"] if h["tone"] != "success" else "success",
                     f"Scan report ready: {filename}", f"جاهز تقرير فحص: {filename}",
                     "", "", route="reports", report_id=rid)
    _push_event({"type": "scan_report", "report_id": rid})
    return rep


# ==================================================================
#  DETECTION ENGINE #7 — BEHAVIORAL HEURISTICS (signature-less)
#  Catches threats by BEHAVIOR, not by matching a known list:
#   - rapid file modification (ransomware-style mass encryption)
#   - process spawning many network connections fast (beaconing/scanning)
#   - a single process opening an unusual number of handles to user docs
#   - suspicious parent->child chains (office app spawning powershell, etc.)
#  Privacy: 100% local, read-only sampling.
# ==================================================================
_HEUR_STATE = {
    "proc_conn_history": {},   # pid -> [(ts, conn_count)]
    "last_doc_mtimes": {},     # path -> mtime  (ransomware canary)
    "alerts": collections.deque(maxlen=100),
}
# parent->child combos that are almost always malicious living-off-the-land
_SUSPICIOUS_CHAINS = {
    ("winword.exe", "powershell.exe"), ("winword.exe", "cmd.exe"),
    ("excel.exe", "powershell.exe"), ("excel.exe", "cmd.exe"),
    ("excel.exe", "wscript.exe"), ("outlook.exe", "powershell.exe"),
    ("powerpnt.exe", "powershell.exe"), ("winword.exe", "wscript.exe"),
    ("mshta.exe", "powershell.exe"), ("wscript.exe", "powershell.exe"),
    ("powershell.exe", "cmd.exe"),  # nested shells via macro droppers
}
# directories whose files we watch as a ransomware canary
_CANARY_DIRS_ENV = "SENTINEL_CANARY_DIRS"


def _heuristic_scan():
    """Behavioral, signature-less detection. Returns a list of finding dicts.
    Each is also logged so it shows in the activity feed and can be auto-isolated."""
    if os.name != "nt" or not HAVE_PSUTIL:
        return {"available": False, "reason": "Windows + psutil required", "alerts": []}
    alerts = []
    now = time.time()

    # ---- 1) parent->child LOLBin chains ----
    try:
        for p in psutil.process_iter(["pid", "name", "ppid"]):
            try:
                name = (p.info.get("name") or "").lower()
                ppid = p.info.get("ppid")
                if not ppid:
                    continue
                parent = psutil.Process(ppid)
                pname = (parent.name() or "").lower()
                if (pname, name) in _SUSPICIOUS_CHAINS:
                    exe = ""
                    try:
                        exe = p.exe()
                    except Exception:
                        pass
                    alerts.append({
                        "type": "lolbin_chain", "sev": "high", "pid": p.info["pid"],
                        "name": name, "parent": pname, "exe": exe,
                        "detail_en": f"{pname} spawned {name} — common macro/dropper behavior",
                        "detail_ar": f"{pname} شغّل {name} — سلوك شائع لبرمجيات الماكرو الخبيثة"})
            except (psutil.NoSuchProcess, psutil.AccessDenied, Exception):
                continue
    except Exception:
        pass

    # ---- 2) per-process connection burst (beaconing / scanning) ----
    try:
        conns_by_pid = {}
        for c in psutil.net_connections(kind="inet"):
            if c.pid:
                conns_by_pid[c.pid] = conns_by_pid.get(c.pid, 0) + 1
        hist = _HEUR_STATE["proc_conn_history"]
        for pid, count in conns_by_pid.items():
            arr = hist.setdefault(pid, [])
            arr.append((now, count))
            # keep only the last 60s
            hist[pid] = [(t, c) for (t, c) in arr if now - t <= 60]
            # a process holding many simultaneous outbound connections is suspect
            if count >= 40:
                try:
                    pr = psutil.Process(pid)
                    nm = (pr.name() or "").lower()
                    exe = pr.exe() if pr else ""
                except Exception:
                    nm, exe = "?", ""
                # trusted browsers/updaters legitimately hold many connections
                if not _proc_is_trustworthy(nm, exe.lower(), exe):
                    alerts.append({
                        "type": "conn_burst", "sev": "medium", "pid": pid,
                        "name": nm, "exe": exe, "count": count,
                        "detail_en": f"{nm} holds {count} live connections (possible beaconing/scan)",
                        "detail_ar": f"{nm} يحمل {count} اتصالاً مباشراً (beaconing/مسح محتمل)"})
    except Exception:
        pass

    # ---- 3) ransomware canary: mass-modification of watched documents ----
    try:
        dirs = os.environ.get(_CANARY_DIRS_ENV, "")
        watch = [d.strip() for d in dirs.split(";") if d.strip()] if dirs else []
        if not watch:
            up = os.path.expanduser("~")
            watch = [os.path.join(up, "Documents"), os.path.join(up, "Desktop"), os.path.join(up, "Pictures")]
        changed = 0
        sampled = 0
        prev = _HEUR_STATE["last_doc_mtimes"]
        cur = {}
        for d in watch:
            if not os.path.isdir(d):
                continue
            for root, _dirs, files in os.walk(d):
                for fn in files[:120]:
                    fp = os.path.join(root, fn)
                    try:
                        mt = os.path.getmtime(fp)
                        cur[fp] = mt
                        sampled += 1
                        if fp in prev and mt != prev[fp]:
                            changed += 1
                    except Exception:
                        continue
                break   # only top level of each watched dir (cheap)
        _HEUR_STATE["last_doc_mtimes"] = cur
        # many docs changing between two close samples = ransomware-like
        if prev and sampled and changed >= max(12, sampled // 3):
            alerts.append({
                "type": "mass_file_change", "sev": "critical", "pid": None,
                "name": "ransomware?", "exe": "",
                "detail_en": f"{changed} watched documents changed at once — possible ransomware encryption",
                "detail_ar": f"{changed} مستنداً مراقباً تغيّر دفعة واحدة — تشفير فدية محتمل"})
    except Exception:
        pass

    # log + remember
    for a in alerts:
        _HEUR_STATE["alerts"].append({**a, "ts": int(now * 1000)})
        log_event("heuristic", sev=a["sev"], asset=_plat.node() or "host",
                  text_en="Behavioral alert: " + a["detail_en"],
                  text_ar="تنبيه سلوكي: " + a["detail_ar"])
    return {"available": True, "ts": int(now * 1000), "alerts": alerts, "count": len(alerts)}


# ==================================================================
#  AUTO-ISOLATION — contain CLEAR threats automatically (opt-in, safe)
#  Only acts on HIGH-CONFIDENCE, unambiguous threats and ALWAYS logs an
#  undo. Controlled by SENTINEL_AUTO_ISOLATE (off by default). The UI lets
#  the user enable "detect + auto-isolate".
# ==================================================================
_isolation_log = collections.deque(maxlen=200)
_auto_isolate_enabled = [os.environ.get("SENTINEL_AUTO_ISOLATE", "").strip() in ("1", "true", "yes")]
# never auto-act on anything whose name is a critical OS process
_NEVER_ISOLATE = {"svchost.exe", "lsass.exe", "csrss.exe", "winlogon.exe", "services.exe",
                  "smss.exe", "wininit.exe", "explorer.exe", "system", "registry",
                  "dwm.exe", "fontdrvhost.exe", "msmpeng.exe", "nissrv.exe"}


def _isolate_process(pid, name, reason):
    """Suspend (not kill) a process so it stops doing damage but can be inspected
    /resumed. Suspending is reversible and far safer than taskkill. Returns a
    result dict with an undo action."""
    nm = (name or "").lower()
    if nm in _NEVER_ISOLATE:
        return {"ok": False, "error": "refused: critical system process", "name": name}
    if os.name != "nt" or not HAVE_PSUTIL:
        return {"ok": False, "error": "not supported"}
    try:
        proc = psutil.Process(pid)
        # verify identity hasn't been recycled to a different process
        if (proc.name() or "").lower() != nm:
            return {"ok": False, "error": "pid no longer matches", "name": name}
        proc.suspend()
        rec = {"ts": int(time.time() * 1000), "action": "suspend", "pid": pid,
               "name": name, "reason": reason,
               "undo_en": f"Resume with: Resume-Process or taskkill /pid {pid} if confirmed malicious",
               "undo_ar": f"الاستئناف: Resume-Process، أو taskkill /pid {pid} إن تأكّد أنه خبيث"}
        _isolation_log.append(rec)
        log_event("isolation", sev="high", asset=_plat.node() or "host",
                  text_en=f"Auto-isolated (suspended) {name} (pid {pid}): {reason}",
                  text_ar=f"عزل تلقائي (تعليق) {name} (pid {pid}): {reason}")
        add_notification("alert", "warning",
                         f"Auto-isolated {name}", f"تم عزل {name} تلقائياً",
                         "", "", route="logs")
        return {"ok": True, "action": "suspend", "pid": pid, "name": name, "undo": rec}
    except (psutil.NoSuchProcess, psutil.AccessDenied) as e:
        return {"ok": False, "error": str(e)[:120], "name": name}
    except Exception as e:
        return {"ok": False, "error": str(e)[:120], "name": name}


def _auto_isolate_from_alerts(alerts):
    """Given heuristic alerts, auto-isolate only the HIGH-CONFIDENCE ones when
    auto-isolation is enabled. Conservative by design: only conn_burst with a
    very high count, mass_file_change's culprit, or a malicious LOLBin chain."""
    if not _auto_isolate_enabled[0]:
        return []
    acted = []
    for a in alerts:
        pid = a.get("pid")
        if not pid:
            continue
        # only contain clear, high/critical behavioral threats
        if a["type"] == "lolbin_chain" or (a["type"] == "conn_burst" and a.get("count", 0) >= 80):
            res = _isolate_process(pid, a.get("name", ""), a.get("detail_en", a["type"]))
            if res.get("ok"):
                acted.append(res)
    return acted


# ==================================================================
#  REAL-TIME PROTECTION ENGINE — per-second monitoring (not periodic)
#  Runs several tight loops in background threads so threats are caught
#  within ~1 second instead of waiting for the hourly scan:
#    A) process watcher  — every new process is inspected the instant it
#       appears; malicious LOLBin chains / bad paths are isolated at once.
#    B) connection watcher — sudden outbound connection bursts (C2/scan).
#    C) file watcher (watchdog) — rapid mass file modification = ransomware;
#       the offending process is found and isolated immediately.
#  Honest scope: this is user-space monitoring. It CANNOT block a file before
#  it executes the way a signed kernel mini-filter (Defender) does; it detects
#  and contains within ~1s of execution. That is the strongest a Python agent
#  can safely do. Toggle with SENTINEL_REALTIME=1 or the UI switch.
# ==================================================================
_realtime_enabled = [os.environ.get("SENTINEL_REALTIME", "1").strip() not in ("0", "false", "no", "off")]
_realtime_threads = {"proc": None, "conn": None, "file": None, "sysmon": None}
_realtime_stats = {"started_ts": 0, "events": 0, "isolations": 0,
                   "last_event_ts": 0, "running": False}
_rt_seen_pids = set()
_rt_lock = threading.Lock()


# ---- Monitor-Only mode: alert but never isolate (recommended for first 1-2 weeks) ----
# Real EDRs run in audit mode first so the admin can tune the allowlist from reality
# before enabling active blocking. Default ON for safety.
_monitor_only = [os.environ.get("SENTINEL_MONITOR_ONLY", "1").strip() not in ("0", "false", "no", "off")]

# ---- Rate limiting / rule-quality tracking (the "golden rule": a rule that fires
# too often is probably a bad rule). We count alerts per reason in a rolling window
# and auto-suppress isolation for any reason that becomes too noisy. ----
_rule_hits = {}            # reason_key -> deque[timestamps]
_rule_suppressed = {}      # reason_key -> until_ts (isolation suppressed)
_RULE_WINDOW = 600         # 10-minute rolling window
_RULE_NOISY = 8            # >8 fires in 10 min => rule is noisy, suppress isolation


def _reason_key(reason):
    # collapse a reason string to a stable rule key (drop pids/paths/numbers)
    import re as _r
    k = _r.sub(r"\d+", "", (reason or "").lower())
    k = _r.sub(r"[^a-z ]", "", k)
    return k.strip()[:60]


def _rule_is_noisy(reason):
    """Track how often this rule fires; suppress isolation if it's too noisy."""
    key = _reason_key(reason)
    now = time.time()
    # already suppressed?
    if _rule_suppressed.get(key, 0) > now:
        return True
    dq = _rule_hits.setdefault(key, collections.deque(maxlen=50))
    dq.append(now)
    recent = sum(1 for t in dq if now - t <= _RULE_WINDOW)
    if recent > _RULE_NOISY:
        # mark this rule noisy for the next 30 min — alert only, no isolation
        _rule_suppressed[key] = now + 1800
        log_event("realtime", sev="info", asset=_plat.node() or "host",
                  text_en=f"Rule auto-tuned: '{key}' fired {recent}x/10min — isolation paused (likely false positives)",
                  text_ar=f"ضبط تلقائي للقاعدة: '{key}' تكرّرت {recent} مرة/١٠د — أُوقف العزل (إيجابيات كاذبة محتملة)",
                  dedupe_sec=300)
        return True
    return False


# ---- Adaptive learning: when the user RESUMES an isolated process, remember the
# culprit so we can suggest allowlisting it (learning from user decisions). ----
_isolation_resume_learn = collections.deque(maxlen=50)


# ─────────────────────────────────────────────────────────────────────
#  BEHAVIORAL BASELINE — learn what's normal on THIS machine, then alert
#  on deviation. This is the single strongest false-positive reducer: a
#  process the user runs every day is "normal" here even if it looks odd
#  generically. We record process names seen during a learning window
#  (default 7 days). After learning, a brand-new never-before-seen process
#  adds a small risk nudge; a well-known one gets a discount.
# ─────────────────────────────────────────────────────────────────────
_BASELINE_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_baseline.json")
_BASELINE_LEARN_DAYS = int(os.environ.get("SENTINEL_BASELINE_DAYS", "7"))
_baseline = {
    "started_ts": 0,            # when learning began (ms)
    "procs": {},                # name -> {count, first_ts, last_ts}
    "hours": [0] * 24,          # activity histogram by hour-of-day
    "learning": True,
}
_baseline_lock = threading.Lock()


def _load_baseline():
    global _baseline
    try:
        data = _secure_load(_BASELINE_FILE, None)
        if data and isinstance(data, dict):
            _baseline.update(data)
    except Exception:
        pass
    # start the clock on first run, and repair any corrupt value (0, negative, or
    # a future timestamp) left over from older builds that caused the ~20000-day bug.
    now_ms = int(time.time() * 1000)
    st = _baseline.get("started_ts") or 0
    if st <= 0 or st > now_ms:
        _baseline["started_ts"] = now_ms
        _save_baseline()


def _save_baseline():
    try:
        _secure_save(_BASELINE_FILE, _baseline)
    except Exception:
        pass


def _baseline_is_learning():
    """True while still inside the learning window."""
    st = _baseline.get("started_ts") or 0
    now_ms = time.time() * 1000
    if st <= 0 or st > now_ms:
        return True   # not started yet / corrupt -> treat as still learning
    age_days = (now_ms - st) / 86400000.0
    learning = age_days < _BASELINE_LEARN_DAYS
    if not learning and _baseline.get("learning"):
        # learning just finished — persist the flip
        _baseline["learning"] = False
        _save_baseline()
    return learning


def _baseline_observe(name):
    """Record a process name as part of normal activity (called during learning)."""
    if not name:
        return
    name = name.lower()
    now = int(time.time() * 1000)
    with _baseline_lock:
        p = _baseline["procs"].get(name)
        if p:
            p["count"] += 1
            p["last_ts"] = now
        else:
            _baseline["procs"][name] = {"count": 1, "first_ts": now, "last_ts": now}
        import datetime as _dt
        _baseline["hours"][_dt.datetime.now().hour] += 1
    # persist more often (every ~10 observations) to limit data loss on a crash;
    # a clean shutdown also saves via the atexit hook registered at startup.
    if sum(v["count"] for v in _baseline["procs"].values()) % 10 == 0:
        _save_baseline()


def _baseline_is_known(name):
    """True if this process name was seen during the learning window."""
    if not name:
        return False
    return name.lower() in _baseline["procs"]


def _baseline_status():
    # guard against a missing/zero/corrupt started_ts (old files saved it as 0,
    # which made age_days explode to ~20000). Repair it to "now" on the fly.
    st = _baseline.get("started_ts") or 0
    now_ms = time.time() * 1000
    if st <= 0 or st > now_ms:
        st = int(now_ms)
        _baseline["started_ts"] = st
        _save_baseline()
    learning = _baseline_is_learning()
    age_days = max(0.0, (now_ms - st) / 86400000.0)
    return {
        "learning": learning,
        "days_elapsed": round(age_days, 1),
        "days_target": _BASELINE_LEARN_DAYS,
        "known_processes": len(_baseline.get("procs", {})),
        "progress_pct": min(100, int(age_days / _BASELINE_LEARN_DAYS * 100)) if _BASELINE_LEARN_DAYS else 100,
    }


def _rt_log(sev, en, ar, isolate_pid=None, isolate_name=None, reason=""):
    """Record a real-time event and optionally isolate. Honors Monitor-Only mode
    and rule-noise suppression so we never spam isolations."""
    with _rt_lock:
        _realtime_stats["events"] += 1
        _realtime_stats["last_event_ts"] = int(time.time() * 1000)
    log_event("realtime", sev=sev, asset=_plat.node() or "host",
              text_en=en, text_ar=ar, dedupe_sec=30)
    # decide whether to isolate: requires a target, auto-isolate ON, NOT monitor-only,
    # and the rule isn't currently flagged as noisy.
    if isolate_pid and _auto_isolate_enabled[0] and not _monitor_only[0]:
        if _rule_is_noisy(reason):
            return  # rule too noisy right now — alert only
        res = _isolate_process(isolate_pid, isolate_name or "", reason or en)
        if res.get("ok"):
            with _rt_lock:
                _realtime_stats["isolations"] += 1


def _rt_inspect_process(pid, name, ppid, exe, cmdline=None):
    """Inspect ONE process using a RISK-SCORING model (not single-signal triggers).
    Aggregates multiple behavioral indicators, subtracts trust signals, then:
        0-39   -> informational (no alert)
        40-69  -> alert only (no isolation)
        70+    -> alert + auto-isolate (if enabled)
    This mirrors how real EDRs decide, and dramatically cuts false positives.
    Shared by the WMI event path and the polling fallback."""
    name = (name or "").lower()
    exe_l = (exe or "").lower()
    score = 0
    reasons_en = []
    reasons_ar = []

    # gather command line once
    cl = (cmdline if cmdline is not None else "")
    if not cl:
        try:
            cl = " ".join(psutil.Process(pid).cmdline())
        except Exception:
            cl = ""
    cl = cl.lower()

    # ---- parent process (for tree analysis + trust) ----
    pname = ""
    parent_trusted = False
    try:
        if ppid:
            par = psutil.Process(ppid)
            pname = (par.name() or "").lower()
            pexe = ""
            try:
                pexe = (par.exe() or "")
            except Exception:
                pass
            if pname in ("services.exe", "svchost.exe", "taskeng.exe", "taskhostw.exe",
                         "msiexec.exe", "wmiprvse.exe", "ccmexec.exe", "gpscript.exe",
                         "sihost.exe", "winlogon.exe", "trustedinstaller.exe"):
                parent_trusted = True
    except Exception:
        pass

    # ---- SCORING: behavioral indicators add points ----
    is_lolbin = name in ("powershell.exe", "cmd.exe", "wscript.exe", "cscript.exe",
                         "mshta.exe", "regsvr32.exe", "rundll32.exe")
    if is_lolbin:
        score += 20; reasons_en.append("script/LOLBin host"); reasons_ar.append("مضيف سكربت/LOLBin")

    # encoded command
    if any(k in cl for k in ("-enc", "-encodedcommand", "frombase64string")):
        score += 30; reasons_en.append("encoded command"); reasons_ar.append("أمر مشفّر")
    # download / remote-exec primitive
    if any(k in cl for k in ("downloadstring", "downloadfile", "webclient", "net.webclient",
                             "invoke-webrequest", "start-bitstransfer", "certutil -urlcache",
                             "bitsadmin /transfer", "iex ", "invoke-expression")):
        score += 40; reasons_en.append("download/remote-exec"); reasons_ar.append("تنزيل/تنفيذ عن بُعد")
    # hidden window
    if any(k in cl for k in ("-w hidden", "-windowstyle hidden")):
        score += 15; reasons_en.append("hidden window"); reasons_ar.append("نافذة مخفية")
    # no-profile / execution-policy bypass (common in attacks)
    if "-nop" in cl or "-noprofile" in cl:
        score += 10; reasons_en.append("-NoProfile"); reasons_ar.append("بلا ملف تعريف")
    if "bypass" in cl and "executionpolicy" in cl:
        score += 15; reasons_en.append("ExecutionPolicy Bypass"); reasons_ar.append("تجاوز سياسة التنفيذ")
    # execution from a temp/writable path
    if exe_l and any(d in exe_l for d in _PROC_SUSPICIOUS_DIRS):
        score += 25; reasons_en.append("runs from temp/writable path"); reasons_ar.append("يعمل من مسار مؤقّت")
    # suspicious parent->child chain (Office spawning a shell, etc.)
    if (pname, name) in _SUSPICIOUS_CHAINS:
        score += 45; reasons_en.append(f"suspicious chain {pname}->{name}"); reasons_ar.append(f"سلسلة مشبوهة {pname}←{name}")

    # nothing notable -> stop early (no alert, no work)
    if score < 40:
        return

    # ---- TRUST signals SUBTRACT points (allowlist + signature) ----
    # valid signature from a known publisher
    try:
        if exe and _proc_is_trustworthy(name, exe_l, exe):
            score -= 30; reasons_en.append("signed/trusted binary"); reasons_ar.append("ملف موقّع/موثوق")
    except Exception:
        pass
    # trusted system launcher as parent
    if parent_trusted:
        score -= 30; reasons_en.append("trusted parent launcher"); reasons_ar.append("مُشغِّل أب موثوق")
    # user allowlist (path / name) — checks the process AND its parent, so adding
    # the launcher (e.g. "WorkloadManager") also clears the powershell it spawns.
    try:
        if _rt_in_allowlist(name, exe_l) or (pname and _rt_in_allowlist(pname, "")):
            score -= 60; reasons_en.append("user allowlist"); reasons_ar.append("قائمة سماح المستخدم")
    except Exception:
        pass

    # ---- BEHAVIORAL BASELINE: learn-then-deviate ----
    # During the learning window we just observe (and trust everything a bit more
    # so we don't nag while learning). After learning, a process we've seen as
    # normal on THIS machine gets a discount; a brand-new one gets a small nudge.
    try:
        if _baseline_is_learning():
            _baseline_observe(name)
            score -= 10; reasons_en.append("baseline learning"); reasons_ar.append("تعلّم خط الأساس")
        elif _baseline_is_known(name):
            score -= 20; reasons_en.append("known-normal on this host"); reasons_ar.append("معتاد على هذا الجهاز")
        else:
            score += 15; reasons_en.append("never seen before"); reasons_ar.append("لم يُرَ من قبل")
    except Exception:
        pass

    # ---- DECISION based on final score ----
    if score < 40:
        return  # trust signals brought it below the alert threshold
    rsn_en = ", ".join(reasons_en)
    rsn_ar = "، ".join(reasons_ar)
    if score >= 70:
        _rt_log("critical",
                f"Real-time: {name} (pid {pid}) risk {score} — {rsn_en}",
                f"لحظي: {name} (pid {pid}) خطورة {score} — {rsn_ar}",
                isolate_pid=pid, isolate_name=name,
                reason=f"risk score {score}: {rsn_en}")
    else:  # 40-69
        _rt_log("high",
                f"Real-time: {name} (pid {pid}) risk {score} (review) — {rsn_en}",
                f"لحظي: {name} (pid {pid}) خطورة {score} (للمراجعة) — {rsn_ar}")


# ---- user allowlist for trusted apps (admin tools, backup, DevOps, internal scripts) ----
_rt_allowlist = {"names": set(), "paths": set()}


def _rt_in_allowlist(name, exe_l):
    name = (name or "").lower()
    exe_l = (exe_l or "").lower()
    # exact name match, or the allowlisted name appears in the process/exe
    for n in _rt_allowlist["names"]:
        if n and (n == name or n in name or (exe_l and n in exe_l)):
            return True
    # path fragment match
    return any(p and exe_l and p in exe_l for p in _rt_allowlist["paths"])


def _rt_process_loop():
    """Watch for NEW processes. Prefers WMI event subscription (event-driven, like
    Defender — the OS wakes us only when a process starts, near-zero idle CPU). If
    WMI/pywin32 is unavailable, falls back to a 1s polling loop with identical
    detection logic so protection strength is never reduced."""
    if not HAVE_PSUTIL:
        return
    # ---- preferred path: WMI event subscription (event-driven, low CPU) ----
    if _rt_try_wmi_process_watch():
        return
    # ---- fallback path: per-second polling (still real-time) ----
    try:
        for p in psutil.process_iter(["pid"]):
            _rt_seen_pids.add(p.info["pid"])
    except Exception:
        pass
    while _realtime_enabled[0]:
        try:
            current = {}
            for p in psutil.process_iter(["pid", "name", "ppid"]):
                current[p.info["pid"]] = p.info
            for pid in (set(current.keys()) - _rt_seen_pids):
                info = current.get(pid) or {}
                exe = ""
                try:
                    exe = psutil.Process(pid).exe() or ""
                except Exception:
                    pass
                _rt_inspect_process(pid, info.get("name"), info.get("ppid"), exe)
            _rt_seen_pids.clear()
            _rt_seen_pids.update(current.keys())
        except Exception:
            pass
        time.sleep(1.0)


def _rt_try_wmi_process_watch():
    """Subscribe to Win32_ProcessStartTrace via WMI. Returns True if the watcher
    ran (event-driven), False if WMI is unavailable so the caller can fall back.
    This is the 'sleep until the OS wakes us' model — minimal idle CPU."""
    if os.name != "nt":
        return False
    try:
        import pythoncom  # from pywin32
        import wmi
    except Exception:
        return False
    try:
        pythoncom.CoInitialize()
        c = wmi.WMI()
        # event-driven: blocks here with no CPU until a process actually starts
        watcher = c.Win32_ProcessStartTrace.watch_for()
        while _realtime_enabled[0]:
            try:
                ev = watcher(timeout_ms=1000)   # wakes on event, else times out to re-check flag
            except wmi.x_wmi_timed_out:
                continue
            except Exception:
                continue
            try:
                pid = int(getattr(ev, "ProcessID", 0) or 0)
                name = (getattr(ev, "ProcessName", "") or "")
                ppid = int(getattr(ev, "ParentProcessID", 0) or 0)
                exe = ""
                cmdline = None
                try:
                    pr = psutil.Process(pid)
                    exe = pr.exe() or ""
                    cmdline = " ".join(pr.cmdline())
                except Exception:
                    pass
                _rt_inspect_process(pid, name, ppid, exe, cmdline)
            except Exception:
                continue
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass
        return True
    except Exception:
        return False


def _rt_conn_loop():
    """Watch outbound connection counts every 2s; flag sudden bursts (C2/scan)."""
    if not HAVE_PSUTIL:
        return
    while _realtime_enabled[0]:
        try:
            counts = {}
            for c in psutil.net_connections(kind="inet"):
                if c.pid and c.status == "ESTABLISHED":
                    counts[c.pid] = counts.get(c.pid, 0) + 1
            for pid, n in counts.items():
                if n >= 50:
                    try:
                        pr = psutil.Process(pid)
                        nm = (pr.name() or "").lower()
                        exe = (pr.exe() or "")
                    except Exception:
                        nm, exe = "?", ""
                    if not _proc_is_trustworthy(nm, exe.lower(), exe):
                        _rt_log("high",
                                f"Real-time: {nm} (pid {pid}) holds {n} live connections (beaconing/scan)",
                                f"لحظي: {nm} (pid {pid}) يحمل {n} اتصالاً (beaconing/مسح)",
                                isolate_pid=(pid if n >= 100 else None), isolate_name=nm,
                                reason="connection flood (real-time)")
        except Exception:
            pass
        time.sleep(2.0)


def _rt_file_loop():
    """Watch document folders for rapid mass modification (ransomware) using
    watchdog if available, else a fast 2s polling fallback."""
    up = os.path.expanduser("~")
    watch = os.environ.get(_CANARY_DIRS_ENV, "")
    dirs = [d.strip() for d in watch.split(";") if d.strip()] if watch else \
           [os.path.join(up, "Documents"), os.path.join(up, "Desktop"), os.path.join(up, "Pictures")]
    dirs = [d for d in dirs if os.path.isdir(d)]
    if not dirs:
        return
    # rolling window of recent modification timestamps
    recent = collections.deque(maxlen=400)

    def _note_change():
        now = time.time()
        recent.append(now)
        # how many changes in the last 5 seconds?
        cutoff = now - 5.0
        burst = sum(1 for t in recent if t >= cutoff)
        if burst >= 25:
            # find the most CPU/IO-active non-trusted process as the likely culprit
            culprit_pid, culprit_name = _rt_guess_file_culprit()
            _rt_log("critical",
                    f"Real-time: {burst} files changed in 5s — possible ransomware encryption",
                    f"لحظي: {burst} ملفاً تغيّر خلال ٥ث — تشفير فدية محتمل",
                    isolate_pid=culprit_pid, isolate_name=culprit_name,
                    reason="mass file modification / ransomware (real-time)")
            recent.clear()

    try:
        from watchdog.observers import Observer
        from watchdog.events import FileSystemEventHandler

        class _H(FileSystemEventHandler):
            def on_modified(self, e):
                if not e.is_directory:
                    _note_change()
            def on_created(self, e):
                if not e.is_directory:
                    _note_change()

        obs = Observer()
        for d in dirs:
            try:
                obs.schedule(_H(), d, recursive=True)
            except Exception:
                pass
        obs.start()
        while _realtime_enabled[0]:
            time.sleep(0.5)
        obs.stop(); obs.join(timeout=2)
    except Exception:
        # polling fallback: snapshot mtimes every 2s
        prev = {}
        while _realtime_enabled[0]:
            try:
                changed = 0
                cur = {}
                for d in dirs:
                    for root, _ds, files in os.walk(d):
                        for fn in files[:200]:
                            fp = os.path.join(root, fn)
                            try:
                                mt = os.path.getmtime(fp); cur[fp] = mt
                                if fp in prev and mt != prev[fp]:
                                    changed += 1
                            except Exception:
                                continue
                        break
                if prev and changed >= 25:
                    culprit_pid, culprit_name = _rt_guess_file_culprit()
                    _rt_log("critical",
                            f"Real-time: {changed} files changed quickly — possible ransomware",
                            f"لحظي: {changed} ملفاً تغيّر بسرعة — فدية محتملة",
                            isolate_pid=culprit_pid, isolate_name=culprit_name,
                            reason="mass file modification / ransomware (real-time)")
                prev = cur
            except Exception:
                pass
            time.sleep(2.0)


def _rt_guess_file_culprit():
    """Best-effort: find the most I/O-active non-trusted process as the likely
    ransomware culprit. Returns (pid, name) or (None, None)."""
    if not HAVE_PSUTIL:
        return (None, None)
    best, best_io = None, -1
    try:
        for p in psutil.process_iter(["pid", "name"]):
            try:
                nm = (p.info.get("name") or "").lower()
                exe = ""
                try:
                    exe = p.exe()
                except Exception:
                    pass
                if _proc_is_trustworthy(nm, exe.lower(), exe):
                    continue
                io = p.io_counters()
                w = getattr(io, "write_bytes", 0)
                if w > best_io:
                    best_io, best = w, (p.info["pid"], nm)
            except Exception:
                continue
    except Exception:
        pass
    return best if best else (None, None)


def _start_realtime():
    """Launch the real-time monitoring threads (idempotent)."""
    if not _realtime_enabled[0] or not HAVE_PSUTIL:
        return
    with _rt_lock:
        if _realtime_stats["running"]:
            return
        _realtime_stats["running"] = True
        _realtime_stats["started_ts"] = int(time.time() * 1000)
    threads = [("proc", _rt_process_loop), ("conn", _rt_conn_loop), ("file", _rt_file_loop)]
    # if Sysmon is installed, add its ETW-grade event reader (injection, tampering,
    # encoded process-create) — much richer telemetry than polling can provide.
    try:
        if _sysmon_available():
            threads.append(("sysmon", _sysmon_loop))
            _realtime_stats["sysmon"] = True
        else:
            _realtime_stats["sysmon"] = False
    except Exception:
        _realtime_stats["sysmon"] = False
    for key, target in threads:
        if _realtime_threads.get(key) is None or not _realtime_threads[key].is_alive():
            t = threading.Thread(target=target, daemon=True, name=f"sentinel-rt-{key}")
            _realtime_threads[key] = t
            t.start()
    log_event("realtime", sev="info", asset=_plat.node() or "host",
              text_en="Real-time protection started (event-driven monitoring active)",
              text_ar="بدأت الحماية اللحظية (مراقبة مدفوعة بالأحداث نشطة)", dedupe_sec=5)


def _stop_realtime():
    _realtime_enabled[0] = False
    with _rt_lock:
        _realtime_stats["running"] = False


@bp.get("/api/realtime/status")
@require_auth("analyst")
def realtime_status():
    return jsonify({"enabled": _realtime_enabled[0],
                    "running": _realtime_stats["running"],
                    "stats": _realtime_stats,
                    "auto_isolate": _auto_isolate_enabled[0],
                    "monitor_only": _monitor_only[0],
                    "baseline": _baseline_status(),
                    "watchdog": _has_watchdog(),
                    "sysmon": _realtime_stats.get("sysmon", False),
                    "engine": ("sysmon+etw" if _realtime_stats.get("sysmon")
                               else ("wmi" if _has_wmi() else "polling"))})


@bp.get("/api/baseline/status")
@require_auth("analyst")
def baseline_status_ep():
    return jsonify(_baseline_status())


@bp.post("/api/baseline/reset")
@require_auth("admin")
def baseline_reset_ep():
    """Restart the learning window from scratch (e.g. after major system changes)."""
    global _baseline
    _baseline = {"started_ts": int(time.time() * 1000), "procs": {},
                 "hours": [0] * 24, "learning": True}
    _save_baseline()
    audit("baseline_reset", "behavioral baseline restarted")
    return jsonify({"ok": True, "baseline": _baseline_status()})


def _has_wmi():
    try:
        import wmi  # noqa
        return True
    except Exception:
        return False


@bp.post("/api/realtime/toggle")
@require_auth("admin")
def realtime_toggle():
    # real-time protection is always-on by design; ensure it's running and report
    # status. (Kept for compatibility; disabling is intentionally not supported.)
    _realtime_enabled[0] = True
    _start_realtime()
    return jsonify({"ok": True, "enabled": True, "running": _realtime_stats["running"],
                    "always_on": True})


def _has_watchdog():
    try:
        import watchdog  # noqa
        return True
    except Exception:
        return False


_ALLOWLIST_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_allowlist.json")


def _load_allowlist():
    """Load the user's trusted-app allowlist from disk into memory."""
    try:
        data = _secure_load(_ALLOWLIST_FILE, {"names": [], "paths": []})
        _rt_allowlist["names"] = set(n.lower() for n in data.get("names", []))
        _rt_allowlist["paths"] = set(p.lower() for p in data.get("paths", []))
    except Exception:
        pass


def _save_allowlist():
    try:
        _secure_save(_ALLOWLIST_FILE, {"names": sorted(_rt_allowlist["names"]),
                                       "paths": sorted(_rt_allowlist["paths"])})
    except Exception:
        pass


@bp.get("/api/allowlist")
@require_auth("analyst")
def allowlist_get():
    return jsonify({"names": sorted(_rt_allowlist["names"]),
                    "paths": sorted(_rt_allowlist["paths"])})


@bp.post("/api/allowlist/add")
@require_auth("analyst")
def allowlist_add():
    """Add a trusted process name or path so it stops triggering alerts/isolation.
    Analyst-level: this is a defensive response action (same tier as triaging a
    finding), and the "Safe" button in live-detection relies on it."""
    data = request.get_json(silent=True) or {}
    nm = (data.get("name") or "").strip().lower()
    pth = (data.get("path") or "").strip().lower()
    if nm:
        _rt_allowlist["names"].add(nm)
    if pth:
        _rt_allowlist["paths"].add(pth)
    _save_allowlist()
    audit("allowlist_add", f"name={nm} path={pth}")
    return jsonify({"ok": True, "names": sorted(_rt_allowlist["names"]),
                    "paths": sorted(_rt_allowlist["paths"])})


@bp.post("/api/allowlist/remove")
@require_auth("admin")
def allowlist_remove():
    data = request.get_json(silent=True) or {}
    nm = (data.get("name") or "").strip().lower()
    pth = (data.get("path") or "").strip().lower()
    _rt_allowlist["names"].discard(nm)
    _rt_allowlist["paths"].discard(pth)
    _save_allowlist()
    return jsonify({"ok": True, "names": sorted(_rt_allowlist["names"]),
                    "paths": sorted(_rt_allowlist["paths"])})


@bp.post("/api/scan/upload")
@require_auth("analyst")
def scan_upload():
    """Accept any file (multipart 'file' OR JSON {filename, content}).
    Binary content is hashed/analysed locally; nothing is uploaded anywhere."""
    filename, content, raw = None, None, None
    if request.files.get("file"):
        f = request.files["file"]
        filename = f.filename or "upload.bin"
        raw = f.read()
        content = raw.decode("utf-8", errors="replace")
    else:
        data = request.get_json(silent=True) or {}
        filename = (data.get("filename") or "upload.txt")[:120]
        content = data.get("content") or ""
        raw = content.encode("utf-8", "replace")
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    if not raw:
        return jsonify({"ok": False, "error": "empty file"}), 400
    raw = raw[:5_000_000]                 # cap raw at ~5MB for hashing/analysis
    content = content[:200000]            # cap text used for the model/heuristics
    rep = analyze_file(filename, content, lang, raw=raw)
    return jsonify({"ok": True, "report": rep})


@bp.get("/api/settings/virustotal")
@require_auth("admin")
def vt_get():
    cfg = dict(_vt_cfg())
    if cfg.get("api_key"):
        cfg["api_key"] = "********"
    return jsonify({"virustotal": cfg, "offline": bool(_OFFLINE)})


@bp.post("/api/settings/virustotal")
@require_auth("admin")
def vt_set():
    data = (request.get_json(silent=True) or {}).get("virustotal") or {}
    with _settings_lock:
        cur = _SETTINGS.setdefault("virustotal", {"enabled": False, "api_key": ""})
        cur["enabled"] = bool(data.get("enabled", cur.get("enabled", False)))
        k = data.get("api_key")
        if isinstance(k, str) and k and "****" not in k:
            cur["api_key"] = k
    _save_settings()
    audit("virustotal_config_changed", f"enabled={_vt_cfg().get('enabled')}")
    return jsonify({"ok": True})


def model_chat_stream(q, lang, max_tokens=None, system_extra="", history=None):
    """Yield answer text in chunks as the model generates (same quality, feels instant)."""
    import urllib.request, urllib.error
    key = _current_model[0]
    tag = MODELS[key]["tag"]
    ctx = (f"Live context: security score {posture_score()}/100, {active_threats()} active threats, "
           f"CPU {_latest['cpu']}%, RAM {_latest['ram']}GB, network {_latest['traffic']} Mbps.")
    with _sec_lock:
        opens = [f for f in _FINDINGS if f["st"] in _OPEN_STATES]
    if opens:
        ctx += " Open findings: " + "; ".join(
            f"{f['cve']} ({f['sev']}, CVSS {f['score']}) on {f['asset']} — fix: {f['fix_en']}" for f in opens[:6]) + "."
    else:
        ctx += " No open vulnerability findings."
    # higher temperature => fluent, natural answers (not canned). penalties still
    # curb broken words / foreign-language leakage from this English-centric model.
    opts = {"num_ctx": NUM_CTX, "temperature": 0.5, "top_p": 0.92, "repeat_penalty": 1.1}
    if max_tokens:
        opts["num_predict"] = int(max_tokens)
    lang_rule = ((" Reply ONLY in fluent Modern Standard Arabic — never English, Russian, or any other script. "
                  "Answer the specific question with the depth of an expert; reason it through, do not repeat boilerplate.")
                 if lang == "ar" else
                 (" Reply in fluent, professional English with expert depth, addressing the specific question asked."))
    sys = _SYS_PROMPT + lang_rule + (" " + system_extra if system_extra else "") + " " + ctx

    # conversation memory: include prior turns
    chat_msgs = [{"role": "system", "content": sys}]
    hist_txt = ""
    if history:
        for turn in history[-6:]:
            role = "assistant" if turn.get("role") == "assistant" else "user"
            content = (turn.get("content") or "").strip()
            if content:
                chat_msgs.append({"role": role, "content": content[:1500]})
                who = "Assistant" if role == "assistant" else "User"
                hist_txt += f"\n{who}: {content[:1500]}"
    chat_msgs.append({"role": "user", "content": q})
    gen_prompt = (hist_txt + "\nUser: " + q).strip() if hist_txt else q

    for path, payload, getter in (
        ("/api/chat", {"model": tag, "stream": True, "options": opts, "messages": chat_msgs},
         lambda o: (o.get("message") or {}).get("content") or ""),
        ("/api/generate", {"model": tag, "stream": True, "options": opts, "system": sys, "prompt": gen_prompt},
         lambda o: o.get("response") or ""),
    ):
        try:
            body = json.dumps(payload).encode("utf-8")
            req = urllib.request.Request(f"{OLLAMA}{path}", data=body,
                                         headers={"Content-Type": "application/json"})
            got = False
            with urllib.request.urlopen(req, timeout=300) as r:
                for raw in r:
                    line = raw.decode("utf-8", "replace").strip()
                    if not line:
                        continue
                    try:
                        o = json.loads(line)
                    except Exception:
                        continue
                    piece = getter(o)
                    if piece:
                        got = True
                        yield piece
                    if o.get("done"):
                        break
            if got:
                _model_last_error[0] = None
                return
        except Exception as e:
            _model_last_error[0] = f"stream {path}: {type(e).__name__}: {e}"
            continue
    # both endpoints failed → single fallback message
    if lang == "ar":
        yield (f"النموذج ({MODELS[key]['label']}) غير متصل. افتح /api/model/health للسبب.")
    else:
        yield (f"Model ({MODELS[key]['label']}) is offline. Open /api/model/health for the reason.")


@bp.post("/api/dashboard/ask")
def ask():
    data = request.get_json(silent=True) or {}
    q = (data.get("q") or "").strip()
    lang = "ar" if str(data.get("lang", "en")).lower().startswith("ar") else "en"
    history = data.get("history") if isinstance(data.get("history"), list) else None
    rag = _rag_context(q, lang)
    # bound the answer length so it returns within the timeout on CPU-only hosts
    res = model_chat(q, lang, max_tokens=800, system_extra=rag, history=history)
    res["q"] = q
    return jsonify(res)


# Plain-language explanation for each finding TYPE, so the user understands what
# a finding means, whether it affects them, and whether to act — without needing
# the LLM. If the model is available we enrich it; otherwise this stands alone.
_FINDING_EXPLAINERS = {
    "PORT-135": ("ar", "منفذ RPC (135) يستخدمه ويندوز داخلياً لخدماته. مفتوح على كل أجهزة ويندوز تقريباً. خطر فقط لو جهازك مكشوف مباشرةً للإنترنت (نادر خلف راوتر منزلي). لو لا تشارك خدمات عبر الشبكة، يمكنك حظره بأمان."),
    "PORT-139": ("ar", "منفذ NetBIOS (139) لمشاركة الملفات القديمة في ويندوز. طبيعي على الشبكات المنزلية. أغلقه فقط إن كنت لا تشارك ملفات/طابعات عبر الشبكة المحلية."),
    "PORT-445": ("ar", "منفذ SMB (445) لمشاركة الملفات الحديثة. طبيعي داخل المنزل، لكنه كان ناقل هجمات مثل WannaCry حين يكون مكشوفاً للإنترنت. خلف راوتر منزلي آمن نسبياً. أغلقه إن لم تشارك ملفات."),
    "KEV": ("ar", "هذا برنامج مُثبَّت عندك ورد في قائمة CISA للثغرات المُستغَلّة فعلياً في هجمات حقيقية. المعنى: إصدارك قد يكون قديماً وفيه ثغرة معروفة. الحل: حدّث البرنامج لآخر إصدار. هذه من أهم الثغرات التي يجب الانتباه لها."),
    "PROC": ("ar", "عملية قيد التشغيل بدت مشبوهة (مكانها أو سلوكها). راجع مسارها: إن كانت برنامجاً تعرفه فهي غالباً سليمة. إن كانت في مجلد مؤقت/تنزيلات ولا تعرفها، افحصها."),
    "PERSIST": ("ar", "مدخل بدء تشغيل (يعمل تلقائياً مع ويندوز). البرامج الخبيثة تستخدم هذا لتبقى بعد إعادة التشغيل. راجعه: إن كان لبرنامج تعرفه فهو سليم، وإلا احذفه."),
    "NET": ("ar", "اتصال شبكي بدا غير طبيعي (منفذ معروف للبرمجيات الخبيثة أو سلوك beaconing). افحص أي برنامج يفتح هذا الاتصال."),
    "ACCT": ("ar", "ملاحظة على حسابات المستخدمين (حساب مخفي أو صلاحيات غير متوقّعة). راجع الحسابات وعطّل ما لا تعرفه."),
    "RKPROC": ("ar", "مؤشّر محتمل لـ rootkit (عملية ظهرت مخفية من بعض أدوات النظام). قد يكون إنذاراً كاذباً من توقيت الفحص. لو تكرّر، شغّل فحص Defender دون اتصال للتأكّد."),
    "RKDRV": ("ar", "درايفر نواة غير موقّع أو في مكان غير معتاد. الدرايفرات الخبيثة تُستخدم في الـ rootkits. تحقّق من مصدره."),
    "FIM": ("ar", "ملف نظام مهم تغيّر منذ آخر فحص. لو كان بسبب تحديث ويندوز فهو طبيعي. وإلا فقد يدل على تلاعب."),
    "FW-DISABLED": ("ar", "جدار حماية ويندوز مُطفأ. يُنصح بشدّة بتشغيله — يحمي جهازك من اتصالات غير مرغوبة. اضغط «معالجة» لتشغيله."),
    "AV-RTP-OFF": ("ar", "الحماية اللحظية في Windows Defender مُطفأة. شغّلها فوراً — هي خط دفاعك الأول ضد البرمجيات الخبيثة."),
    "SMBV1-ENABLED": ("ar", "بروتوكول SMBv1 القديم مُفعّل — وهو ناقل WannaCry الشهير. عطّله إلا إن كنت تتصل بأجهزة قديمة جداً."),
    "UAC-DISABLED": ("ar", "التحكم بحساب المستخدم (UAC) مُطفأ — وهو النوافذ التي تطلب إذنك قبل التغييرات المهمة. أعد تفعيله."),
    "GUEST-ACTIVE": ("ar", "حساب الضيف مُفعّل — موطئ قدم كلاسيكي للمهاجمين. عطّله."),
    "RDP-NO-NLA": ("ar", "سطح المكتب البعيد بلا مصادقة NLA — يسمح بمحاولات اتصال قبل التحقّق من الهوية. فعّل NLA."),
}


def _explain_finding(cve, lang, sev=None, asset=None):
    """Return a plain-language explanation for a finding. Matched by prefix so
    PORT-139-NetBIOS -> PORT-139, KEV-git -> KEV, etc."""
    key = None
    for k in sorted(_FINDING_EXPLAINERS, key=len, reverse=True):
        if cve.startswith(k):
            key = k
            break
    base = _FINDING_EXPLAINERS.get(key, (None, None))[1]
    if not base:
        base = ("هذه ملاحظة أمنية عُثر عليها على جهازك. راجع عمود «الإصلاح» للتفاصيل."
                if lang == "ar" else
                "A security finding on your machine. See the Fix column for details.")
    return base


@bp.post("/api/security/explain")
@require_auth("viewer")
def security_explain():
    """Explain a finding in plain language. Works offline (built-in explainers);
    if the local model is available, it adds a tailored note. 100% on-device."""
    data = request.get_json(silent=True) or {}
    cve = (data.get("cve") or "").strip()
    lang = "ar" if str(data.get("lang", "ar")).lower().startswith("ar") else "en"
    if not cve:
        return jsonify({"ok": False, "error": "missing cve"}), 400
    # find the live finding for context
    f = next((x for x in _FINDINGS if x.get("cve") == cve), None)
    sev = f.get("sev") if f else None
    base = _explain_finding(cve, lang, sev)
    answer = base
    # optional model enrichment (best-effort, never blocks; the built-in
    # explanation above always works even with no model installed).
    if data.get("ai"):
        try:
            desc = (f.get("text_ar") if f and lang == "ar" else f.get("text_en")) if f else cve
            prompt = (f"اشرح باختصار شديد (جملتين) لمستخدم عادي هذه الملاحظة الأمنية: {desc}. "
                      f"هل هي خطيرة فعلاً؟ وهل ينصح بإغلاقها؟"
                      if lang == "ar" else
                      f"Briefly (2 sentences) explain this security finding to a non-expert: {desc}. "
                      f"Is it really dangerous, and should they fix it?")
            res = model_chat(prompt, lang, max_tokens=200)
            if res.get("answer"):
                answer = base + "\n\n" + res["answer"].strip()
        except Exception:
            pass
    return jsonify({"ok": True, "cve": cve, "explanation": answer, "severity": sev})


@bp.post("/api/dashboard/ask-stream")
def ask_stream():
    data = request.get_json(silent=True) or {}
    q = (data.get("q") or "").strip()
    lang = "ar" if str(data.get("lang", "en")).lower().startswith("ar") else "en"
    history = data.get("history") if isinstance(data.get("history"), list) else None
    rag = _rag_context(q, lang)

    def gen():
        try:
            for piece in model_chat_stream(q, lang, max_tokens=800, system_extra=rag, history=history):
                yield "data: " + json.dumps({"delta": piece}) + "\n\n"
        except Exception as e:
            yield "data: " + json.dumps({"delta": f"[error: {e}]"}) + "\n\n"
        yield "data: " + json.dumps({"done": True}) + "\n\n"

    return Response(gen(), mimetype="text/event-stream",
                    headers={"Cache-Control": "no-cache", "X-Accel-Buffering": "no"})


@bp.get("/api/model/ping")
def model_ping():
    """Lightweight liveness check: is Ollama reachable right now AND is the current
    model installed? Cheap (just lists tags, no generation) so the UI can poll it
    to show a truthful connected/offline indicator even when idle."""
    import urllib.request
    res = {"reachable": False, "model_ready": False, "model": _current_model[0]}
    try:
        with urllib.request.urlopen(f"{OLLAMA}/api/tags", timeout=3) as r:
            tags = json.loads(r.read().decode("utf-8"))
            names = [m.get("name") or m.get("model") for m in tags.get("models", [])]
            res["reachable"] = True
            cur_tag = MODELS.get(_current_model[0], {}).get("tag", "")
            res["model_ready"] = cur_tag in names
    except Exception:
        res["reachable"] = False
    res["connected"] = bool(res["reachable"] and res["model_ready"])
    return jsonify(res)


@bp.get("/api/model/health")
def model_health():
    """Diagnostic: tells exactly why the model is or isn't connecting."""
    import urllib.request
    out = {"ollama_url": OLLAMA, "ollama_reachable": False, "installed_models": [],
           "configured": {k: MODELS[k]["tag"] for k in MODELS},
           "current": _current_model[0], "tag_found": {}, "test": {}}
    # 1) is Ollama up? list installed models
    try:
        with urllib.request.urlopen(f"{OLLAMA}/api/tags", timeout=8) as r:
            tags = json.loads(r.read().decode("utf-8"))
            names = [m.get("name") or m.get("model") for m in tags.get("models", [])]
            out["ollama_reachable"] = True
            out["installed_models"] = names
            for k, v in MODELS.items():
                out["tag_found"][k] = v["tag"] in names
    except Exception as e:
        out["error"] = f"cannot reach Ollama at {OLLAMA} — is it running? ({e.__class__.__name__})"
        return jsonify(out)
    # 2) try a tiny real generation with the current model
    t0 = time.time()
    res = model_chat("Reply with the single word: OK", _lang(), max_tokens=8)
    out["test"] = {"source": res.get("source"), "latency_sec": round(time.time() - t0, 1),
                   "answer": res.get("answer", "")[:160]}
    out["working"] = res.get("source") == "live"
    out["lastChatError"] = _model_last_error[0]
    if not out["working"] and not out["tag_found"].get(_current_model[0]):
        out["hint"] = ("Current model tag is not in installed_models — pull it or fix the tag "
                       "(SENTINEL_MODEL_Q4 / SENTINEL_MODEL_Q8).")
    return jsonify(out)


@bp.get("/api/models")
def models():
    lang = _lang()
    # discover what's ACTUALLY installed in Ollama so the user can pick any of
    # their local models, not just the two presets.
    discovered = _discover_ollama_models()
    out = [{"id": k, "label": v["label"], "tag": v["tag"],
            "note": v["note_ar"] if lang == "ar" else v["note_en"],
            "installed": any(d.get("name") == v["tag"] or d.get("name", "").split(":")[0] == v["tag"].split(":")[0] for d in discovered)}
           for k, v in MODELS.items()]
    # add any extra installed models not already in our presets
    known_tags = {v["tag"] for v in MODELS.values()}
    for d in discovered:
        nm = d.get("name", "")
        if nm and nm not in known_tags:
            size_gb = round(d.get("size", 0) / (1024**3), 1) if d.get("size") else None
            out.append({"id": "ollama:" + nm, "label": nm.split("/")[-1][:40],
                        "tag": nm, "installed": True,
                        "note": (f"مثبّت في Ollama" + (f" · {size_gb}GB" if size_gb else "")) if lang == "ar"
                                else (f"Installed in Ollama" + (f" · {size_gb}GB" if size_gb else ""))})
    return jsonify({"current": _current_model[0], "models": out})


@bp.post("/api/model")
def set_model():
    data = request.get_json(silent=True) or {}
    key = (data.get("id") or "").strip()
    if key in MODELS:
        _current_model[0] = key
        return jsonify({"ok": True, "current": key})
    # allow selecting a discovered Ollama model (id is "ollama:<tag>")
    if key.startswith("ollama:"):
        tag = key.split("ollama:", 1)[1]
        MODELS[key] = {"tag": tag, "label": tag.split("/")[-1][:40],
                       "note_en": "Installed in Ollama", "note_ar": "مثبّت في Ollama"}
        _current_model[0] = key
        return jsonify({"ok": True, "current": key})
    return jsonify({"ok": False, "error": "unknown model"}), 400


@bp.get("/api/notifications")
def notifications():
    return jsonify({"notifications": _notifs_out(_lang())})


@bp.post("/api/notifications/clear")
@require_auth("analyst")
def notifications_clear():
    """Clear all current notifications (e.g. stale alerts from before an allowlist
    change). Does not affect the permanent event log."""
    _NOTIFS.clear()
    return jsonify({"ok": True})


@bp.get("/api/reports")
def reports():
    lang = _lang()
    period = request.args.get("period")
    template = request.args.get("template") or ""
    # accept the standard periods plus the on-demand "status" snapshot the UI
    # uses for its "Generate report" button.
    if period and (period in _PERIODS or period in ("status", "surprise", "incident")):
        rep = generate_report(period, lang, template=template)
        return jsonify({"ok": True, "report": rep})
    return jsonify({"reports": _reports_list(lang)})


@bp.get("/api/report/<rid>")
def report_one(rid):
    rep = _REPORTS.get(rid)
    if not rep:
        return jsonify({"error": "not found"}), 404
    return jsonify({"report": rep})


# ----- SECURITY -----
@bp.get("/api/security/overview")
def security_overview():
    _ensure_sampler()
    lang = _lang()
    return jsonify({"source": "live", "counts": _counts(), "riskScore": posture_score(),
                    "findings": _findings_out(lang), "recommendations": _recs_out(lang)})


@bp.get("/api/threats/feed")
def threats_feed():
    """Latest publicly-known actively-exploited vulnerabilities (CISA KEV)."""
    _ensure_threat_intel()
    try:
        limit = int(request.args.get("limit", 20))
    except Exception:
        limit = 20
    kev = sorted(_TI.get("kev", []), key=lambda e: e.get("added", ""), reverse=True)
    return jsonify({
        "source": "CISA Known Exploited Vulnerabilities",
        "ts": _TI.get("ts", 0),
        "total": _TI.get("count", 0),
        "offline": bool(_OFFLINE),
        "error": _TI.get("error", ""),
        "items": kev[:max(1, min(limit, 100))],
        "matches": _match_kev_to_software(),     # exploited vulns on THIS host (local match)
    })


@bp.post("/api/threats/refresh")
@require_auth("analyst")
def threats_refresh():
    if _OFFLINE:
        return jsonify({"ok": False, "offline": True})
    threading.Thread(target=lambda: update_threat_intel(force=True), daemon=True).start()
    return jsonify({"ok": True})


@bp.get("/api/privacy")
def privacy_get():
    return jsonify({"offline": bool(_OFFLINE),
                    "note": "All system analysis is local. Only public threat lists are downloaded; no machine data is ever sent."})


@bp.post("/api/privacy")
@require_auth("admin")
def privacy_set():
    global _OFFLINE
    data = request.get_json(silent=True) or {}
    _OFFLINE = bool(data.get("offline"))
    return jsonify({"ok": True, "offline": _OFFLINE})


@bp.get("/api/threat-intel/latest")
def threat_intel_latest():
    """Latest actively-exploited vulnerabilities (public CISA KEV feed). Privacy:
    fetched via one-way GET; no host data is ever transmitted."""
    kev = sorted(_TI.get("kev", []), key=lambda e: e.get("added", ""), reverse=True)
    return jsonify({
        "status": _TI.get("status", "offline" if _OFFLINE else "pending"),
        "updated": _TI.get("ts", 0), "total": _TI.get("count", 0),
        "offline": bool(_OFFLINE), "error": _TI.get("error"),
        "latest": kev[:20],
    })


@bp.post("/api/threat-intel/update")
@require_auth("analyst")
def threat_intel_update():
    if _OFFLINE:
        return jsonify({"ok": False, "offline": True})
    threading.Thread(target=lambda: update_threat_intel(force=True), daemon=True).start()
    return jsonify({"ok": True, "updating": True})


@bp.get("/api/incidents/scan")
def incidents_scan():
    """Run an on-demand event-log incident sweep (Windows)."""
    _ensure_sampler()
    found = _detect_incidents()
    return jsonify({"ok": True, "found": len(found),
                    "incidents": [{"kind": k, "sev": s, "en": e, "ar": a} for k, s, _, e, a in found]})


@bp.post("/api/security/scan")
@require_auth("analyst")
def security_scan_start():
    """Trigger a real vulnerability/configuration scan of this machine."""
    _ensure_sampler()
    if not _last_scan.get("running"):
        threading.Thread(target=run_scan_now, daemon=True).start()
    return jsonify({"ok": True, "running": True})


@bp.get("/api/security/scan/status")
def security_scan_status():
    return jsonify({"running": bool(_last_scan.get("running")),
                    "ts": _last_scan.get("ts", 0),
                    "count": _last_scan.get("count", 0),
                    "software_count": len(_last_scan.get("software", []))})


@bp.get("/api/security/software")
@require_auth("viewer")
def security_software():
    sw = _last_scan.get("software", [])
    return jsonify({"count": len(sw), "software": sorted(sw, key=lambda a: a["name"].lower())})


@bp.get("/api/compliance")
def compliance_get():
    return jsonify(_run_compliance_check())


@bp.get("/api/network/analysis")
def network_analysis():
    _ensure_sampler()
    data = _net_analysis()
    data["firewall_rules"] = _firewall_rule_count()
    return jsonify(data)


def _valid_ip(ip):
    """Strict validation: only a well-formed IPv4/IPv6 address passes. Prevents any
    argument/parameter injection into the netsh firewall command and stops junk
    from being stored in the blocklist."""
    import ipaddress
    try:
        ipaddress.ip_address(ip)
        return True
    except Exception:
        return False


@bp.post("/api/network/block")
@require_auth("analyst")
def network_block():
    ip = ((request.get_json(silent=True) or {}).get("ip") or "").strip()
    if not ip:
        return jsonify({"ok": False, "error": "ip required"}), 400
    if not _valid_ip(ip):
        return jsonify({"ok": False, "error": "invalid IP address"}), 400
    with _settings_lock:
        bl = _SETTINGS.setdefault("blocklist", [])
        if ip not in bl:
            bl.append(ip)
    _save_settings()
    audit("ip_blocked", ip)
    # best-effort real OS-level block via Windows Firewall (needs admin)
    if os.name == "nt":
        _run_cmd(["netsh", "advfirewall", "firewall", "add", "rule",
                  f"name=Sentinel-Block-{ip}", "dir=out", "action=block", f"remoteip={ip}"], 10)
        _run_cmd(["netsh", "advfirewall", "firewall", "add", "rule",
                  f"name=Sentinel-Block-{ip}", "dir=in", "action=block", f"remoteip={ip}"], 10)
    return jsonify({"ok": True, "blocklist": _SETTINGS.get("blocklist", [])})


@bp.post("/api/network/unblock")
@require_auth("analyst")
def network_unblock():
    ip = ((request.get_json(silent=True) or {}).get("ip") or "").strip()
    if not _valid_ip(ip):
        return jsonify({"ok": False, "error": "invalid IP address"}), 400
    with _settings_lock:
        bl = _SETTINGS.setdefault("blocklist", [])
        if ip in bl:
            bl.remove(ip)
    _save_settings()
    audit("ip_unblocked", ip)
    if os.name == "nt":
        _run_cmd(["netsh", "advfirewall", "firewall", "delete", "rule", f"name=Sentinel-Block-{ip}"], 10)
    return jsonify({"ok": True, "blocklist": _SETTINGS.get("blocklist", [])})


@bp.get("/api/network/geoip/status")
def network_geoip_status():
    """Current Geo-IP database state. Always shows that lookups are local-only."""
    return jsonify({
        "loaded": len(_geo_starts) > 0,
        "ranges": len(_geo_starts),
        "offline": bool(_OFFLINE),
        "file": os.path.basename(_GEO_FILE) if os.path.exists(_GEO_FILE) else None,
        "source": _GEO_URL,
        "attribution": "IP-to-country © DB-IP.com (CC-BY-4.0)",
        "privacy": "Local lookup only. The one-time download is anonymous from a public CDN; no IP is ever sent to any service.",
    })


@bp.post("/api/network/geoip/update")
@require_auth("admin")
def network_geoip_update():
    """Download or refresh the local IP-to-country DB. One-way GET to a public mirror."""
    force = bool((request.get_json(silent=True) or {}).get("force"))
    res = update_geoip(force=force)
    audit("geoip_update", str(res.get("status")))
    return jsonify(res)


@bp.get("/api/yara/status")
def yara_status():
    """YARA rule engine status. All scanning is local; nothing is uploaded."""
    files, comm_files = [], []
    try:
        if os.path.isdir(_YARA_DIR):
            files = sorted(f for f in os.listdir(_YARA_DIR) if f.lower().endswith((".yar", ".yara")))
        comm = os.path.join(_YARA_DIR, "community")
        if os.path.isdir(comm):
            comm_files = sorted(f for f in os.listdir(comm) if f.lower().endswith((".yar", ".yara")))
    except Exception:
        pass
    yara_available = True
    yara_detail = ""
    try:
        import yara  # noqa
        # importing can succeed but the underlying libyara may still be broken;
        # a tiny compile proves the engine actually works.
        try:
            yara.compile(source='rule _t { condition: true }')
        except Exception as ce:
            yara_available = False
            yara_detail = f"installed but libyara failed: {str(ce)[:80]}"
    except Exception as ie:
        yara_available = False
        yara_detail = f"not importable: {str(ie)[:80]}"
    total_rules = len(files) + len(comm_files)
    # "active" means we can actually scan: either yara-python compiled the rules,
    # OR the built-in fallback engine is available (it always is) AND we have rules.
    fallback_active = _yara_fallback_available()
    rules_active = (_yara_compiled[0] is not None) or (fallback_active and total_rules > 0)
    return jsonify({
        "yara_python_installed": yara_available,
        "yara_detail": yara_detail,
        "rules_dir": _YARA_DIR,
        "rule_files": files,
        "community_files": comm_files,
        "total_rules": total_rules,
        "loaded": rules_active,
        "compiled": _yara_compiled[0] is not None,
        "engine": ("yara-python" if (yara_available and _yara_compiled[0] is not None)
                   else ("built-in fallback" if total_rules > 0 else "no rules")),
        "auto_update": {
            "enabled": not _OFFLINE,
            "interval_hours": _YARA_UPDATE_HOURS,
            "last_run_ts": _yara_update_state["ts"],
            "downloaded": _yara_update_state["ok"],
            "failed": _yara_update_state["fail"],
            "last_error": _yara_update_state["last_error"],
            "sources": len([u for u in _YARA_COMMUNITY_URLS if u.strip()]),
        },
        "privacy": "All YARA scanning happens locally on this machine. No file content is ever transmitted. Community rule downloads are anonymous one-way GETs from a public CDN.",
    })


@bp.post("/api/yara/update-community")
@require_auth("admin")
def yara_update_community():
    """Manually refresh the community rule pack from public mirrors."""
    res = update_yara_rules(force=True)
    audit("yara_community_update", str(res.get("status")))
    return jsonify(res)


@bp.post("/api/yara/reload")
@require_auth("admin")
def yara_reload():
    """Recompile rules after the user adds/edits files in sentinel_rules/."""
    count, err = _yara_load()
    audit("yara_reload", f"{count} files, err={err}")
    return jsonify({"ok": err is None, "count": count, "error": err})


# ---- live detection engines ----
@bp.post("/api/detect/processes")
@require_auth("analyst")
def detect_processes():
    """Scan running processes for in-memory malware indicators (local-only)."""
    res = scan_processes()
    audit("scan_processes", f"{len(res.get('hits', []))} hits")
    return jsonify(res)


@bp.post("/api/detect/persistence")
@require_auth("analyst")
def detect_persistence():
    """Audit autostart/persistence locations (local-only)."""
    res = scan_persistence()
    audit("scan_persistence", f"{len(res.get('items', []))} items")
    return jsonify(res)


@bp.get("/api/detect/fim")
@require_auth("analyst")
def detect_fim_check():
    """Check watched files against the integrity baseline (local-only)."""
    return jsonify(fim_check())


@bp.post("/api/detect/fim/baseline")
@require_auth("admin")
def detect_fim_baseline():
    """(Re)create the file-integrity baseline from current file hashes."""
    extra = (request.get_json(silent=True) or {}).get("paths") or []
    res = fim_baseline(extra_paths=[p for p in extra if isinstance(p, str)])
    audit("fim_baseline", str(res.get("count")))
    return jsonify(res)


@bp.post("/api/detect/full")
@require_auth("analyst")
def detect_full():
    """Run all live detection engines at once and return a combined verdict."""
    procs = scan_processes()
    persist = scan_persistence()
    fim = fim_check()
    net = scan_network_threats()
    accounts = scan_accounts()
    rootkit = scan_rootkit()
    heur = _heuristic_scan()
    # auto-isolate clear behavioral threats if the user enabled it
    isolated = _auto_isolate_from_alerts(heur.get("alerts", [])) if heur.get("available") else []
    total = (len(procs.get("hits", [])) + len(persist.get("items", []))
             + len(fim.get("changes", [])) + len(net.get("hits", []))
             + len(accounts.get("items", [])) + rootkit.get("total", 0)
             + heur.get("count", 0))
    return jsonify({"ok": True, "ts": int(time.time() * 1000), "total_findings": total,
                    "processes": procs, "persistence": persist, "fim": fim,
                    "network": net, "accounts": accounts, "rootkit": rootkit,
                    "heuristics": heur, "isolated": isolated})


@bp.post("/api/detect/heuristics")
@require_auth("analyst")
def detect_heuristics():
    """Behavioral, signature-less detection (ransomware/beaconing/LOLBin chains)."""
    res = _heuristic_scan()
    isolated = _auto_isolate_from_alerts(res.get("alerts", [])) if res.get("available") else []
    audit("scan_heuristics", f"alerts={res.get('count', 0)} isolated={len(isolated)}")
    return jsonify({**res, "isolated": isolated})


@bp.get("/api/isolation/log")
@require_auth("analyst")
def isolation_log():
    """Recent auto-isolation actions, with undo instructions."""
    return jsonify({"enabled": _auto_isolate_enabled[0],
                    "items": list(_isolation_log)[::-1]})


@bp.post("/api/isolation/toggle")
@require_auth("admin")
def isolation_toggle():
    """Enable/disable automatic isolation of clear threats (admin only)."""
    data = request.get_json(silent=True) or {}
    _auto_isolate_enabled[0] = bool(data.get("enabled"))
    audit("isolation_toggle", f"enabled={_auto_isolate_enabled[0]}")
    return jsonify({"ok": True, "enabled": _auto_isolate_enabled[0]})


@bp.post("/api/isolation/monitor-only")
@require_auth("admin")
def isolation_monitor_only():
    """Toggle Monitor-Only mode: when ON, the engine ALERTS but never isolates —
    recommended for the first 1-2 weeks so you can tune the allowlist from reality
    before enabling active blocking (exactly how production EDRs roll out)."""
    data = request.get_json(silent=True) or {}
    _monitor_only[0] = bool(data.get("enabled"))
    audit("monitor_only_toggle", f"enabled={_monitor_only[0]}")
    return jsonify({"ok": True, "monitor_only": _monitor_only[0]})


@bp.post("/api/isolation/resume/<int:pid>")
@require_auth("admin")
def isolation_resume(pid):
    """Undo an isolation: resume a previously-suspended process (admin only).
    Also LEARNS: records the process so we can suggest allowlisting it, since the
    user resuming it is a signal it was a false positive."""
    if os.name != "nt" or not HAVE_PSUTIL:
        return jsonify({"ok": False, "error": "not supported"}), 400
    try:
        proc = psutil.Process(pid)
        nm = ""
        try:
            nm = (proc.name() or "").lower()
        except Exception:
            pass
        proc.resume()
        # adaptive learning: this resume suggests a false positive
        suggest = None
        if nm:
            _isolation_resume_learn.append({"name": nm, "ts": int(time.time() * 1000)})
            # if the user has resumed this same process 2+ times, strongly suggest allowlist
            same = sum(1 for r in _isolation_resume_learn if r["name"] == nm)
            suggest = {"name": nm, "times_resumed": same,
                       "recommend_allowlist": same >= 2}
        log_event("isolation", sev="info", asset=_plat.node() or "host",
                  text_en=f"Resumed process pid {pid} (isolation undone)",
                  text_ar=f"استُؤنفت العملية pid {pid} (تراجع عن العزل)")
        audit("isolation_resume", f"pid={pid} name={nm}")
        return jsonify({"ok": True, "pid": pid, "suggest": suggest})
    except Exception as e:
        return jsonify({"ok": False, "error": str(e)[:120]}), 400


@bp.get("/api/isolation/suggestions")
@require_auth("analyst")
def isolation_suggestions():
    """Apps the user has resumed repeatedly → candidates for the allowlist."""
    counts = {}
    for r in _isolation_resume_learn:
        counts[r["name"]] = counts.get(r["name"], 0) + 1
    suggestions = [{"name": n, "times": c} for n, c in counts.items() if c >= 2
                   and n not in _rt_allowlist["names"]]
    return jsonify({"suggestions": sorted(suggestions, key=lambda x: -x["times"])})


@bp.post("/api/detect/rootkit")
@require_auth("analyst")
def detect_rootkit():
    """Cross-view + driver-signature rootkit detection (local & read-only)."""
    res = scan_rootkit()
    audit("scan_rootkit", f"total={res.get('total', 0)}")
    return jsonify(res)


@bp.post("/api/detect/network")
@require_auth("analyst")
def detect_network():
    """Analyse live network connections for C2/beaconing (local-only)."""
    res = scan_network_threats()
    audit("scan_network", f"{len(res.get('hits', []))} hits")
    return jsonify(res)


@bp.post("/api/detect/accounts")
@require_auth("analyst")
def detect_accounts():
    """Audit local accounts for compromise indicators (local-only)."""
    res = scan_accounts()
    audit("scan_accounts", f"{len(res.get('items', []))} items")
    return jsonify(res)


@bp.post("/api/yara/install-starter-pack")
@require_auth("admin")
def yara_starter_pack():
    """Write a small bundled set of starter rules to sentinel_rules/ so the user
    has something to scan with out-of-the-box. No download; rules are inline."""
    try:
        os.makedirs(_YARA_DIR, exist_ok=True)
    except Exception:
        pass
    starter = {
        "sentinel_starter.yar": '''
rule Suspicious_Encoded_PowerShell {
    meta: description = "PowerShell run with -encodedcommand / base64"
    strings:
        $a = "powershell" nocase
        $b = "-enc" nocase
        $c = "FromBase64String" nocase
    condition: $a and ($b or $c)
}

rule Process_Injection_API_Set {
    meta: description = "Classic process-injection API trio"
    strings:
        $a = "VirtualAllocEx" ascii nocase
        $b = "WriteProcessMemory" ascii nocase
        $c = "CreateRemoteThread" ascii nocase
    condition: 2 of them
}

rule Office_With_VBA_Macro {
    meta: description = "Office Open XML container with VBA project"
    strings:
        $pk = { 50 4B 03 04 }
        $vba = "vbaProject.bin" nocase
    condition: $pk at 0 and $vba
}

rule Network_Downloader_API {
    meta: description = "Common one-line downloaders"
    strings:
        $a = "URLDownloadToFile" nocase
        $b = "WinHttp" nocase
        $c = "Invoke-WebRequest" nocase
        $d = "DownloadString" nocase
    condition: 2 of them
}

rule Ransomware_Note_Hints {
    meta: description = "Strings often found in ransom notes"
    strings:
        $a = "your files have been encrypted" nocase
        $b = "bitcoin" nocase
        $c = "decrypt" nocase
        $d = "ransom" nocase
    condition: 3 of them
}
''',
    }
    written = []
    for name, body in starter.items():
        p = os.path.join(_YARA_DIR, name)
        try:
            with open(p, "w", encoding="utf-8") as f:
                f.write(body.lstrip())
            written.append(name)
        except Exception:
            pass
    count, err = _yara_load()
    audit("yara_starter_pack", str(written))
    return jsonify({"ok": True, "files": written, "loaded": count, "error": err})


@bp.get("/api/security/remediation/<cve>")
@require_auth("analyst")
def remediation_options(cve):
    """List the available remediation plans for a finding — commands shown
    IN FULL so the user can review before approving anything."""
    f = next((x for x in _FINDINGS if x["cve"] == cve), None)
    if not f:
        # Synthesize a finding from the id so plans resolve even after a rescan.
        # Infer the category from the cve prefix so software/intel plans appear.
        cat = ""
        up = cve.upper()
        if up.startswith("KEV-"):
            cat = "intel"
        elif up.startswith(("PATCH", "OS-")):
            cat = "config"
        elif up.startswith("PORT-"):
            cat = "exposure"
        f = {"cve": cve, "asset": _plat.node() or "host", "sev": "medium",
             "score": 5.0, "fix_en": "", "fix_ar": "", "cat": cat,
             "title_en": cve, "title_ar": cve}
    plans = _remediation_plans(f)
    return jsonify({"cve": cve, "plans": plans,
                    "note": "Nothing runs until you explicitly approve a plan."})


@bp.post("/api/security/remediate")
@require_auth("admin")
def remediation_execute():
    """Execute ONE approved plan. Requires admin + CSRF + explicit confirm flag."""
    data = request.get_json(silent=True) or {}
    cve = (data.get("cve") or "").strip()
    plan_id = (data.get("plan") or "").strip()
    confirmed = bool(data.get("confirm"))
    if not confirmed:
        return jsonify({"error": "confirmation required — set confirm:true after reviewing the commands"}), 400
    s = _current_session()
    user = (s or {}).get("user", "?")
    res = execute_remediation(cve, plan_id, user)
    code = 200 if res.get("ok") else 400
    return jsonify(res), code


@bp.get("/api/security/remediation-log")
@require_auth("analyst")
def remediation_log():
    return jsonify({"items": list(_remediation_log)[::-1]})


@bp.post("/api/security/finding-action")
@require_auth("analyst")
def security_finding_action():
    data = request.get_json(silent=True) or {}
    cve = (data.get("cve") or "").strip()
    action = (data.get("action") or "").strip()
    new_status = "patched" if action == "apply" else "dismissed"
    asset = ""
    with _sec_lock:
        for f in _FINDINGS:
            if f["cve"] == cve:
                f["st"] = new_status
                asset = f.get("asset", "")
                break
    log_event("finding_patched" if action == "apply" else "finding_dismissed",
              sev="info", cve=cve, asset=asset,
              text_en=f"{cve} on {asset} {'patched' if action=='apply' else 'dismissed'}",
              text_ar=f"{cve} على {asset} {'تمت معالجته' if action=='apply' else 'استُبعد'}")
    _push_event({"type": "tick", "data": _overview_payload(_lang())})
    return jsonify({"source": "live", "cve": cve, "action": action, "status": new_status,
                    "score": posture_score(), "threats": active_threats()})


@bp.post("/api/security/rec-action")
@require_auth("analyst")
def security_rec_action():
    data = request.get_json(silent=True) or {}
    rid = (data.get("id") or "").strip()
    action = (data.get("action") or "").strip()
    state = "applied" if action == "apply" else "dismissed"
    with _sec_lock:
        for r in _RECS:
            if r["id"] == rid:
                r["state"] = state
                break
    return jsonify({"source": "live", "id": rid, "action": action, "state": state})


# ==================================================================
#  PROFESSIONAL PDF REPORTS  — CMAR-style, bilingual (EN + AR), live data
# ==================================================================

def _pdf_fonts():
    """Register an Arabic-capable TTF; return (regular, bold, ar_ok)."""
    if getattr(_pdf_fonts, "_c", None):
        return _pdf_fonts._c
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    reg, bld, ar = "Helvetica", "Helvetica-Bold", False
    for rp, bp, is_ar in [
        (r"C:\Windows\Fonts\segoeui.ttf",  r"C:\Windows\Fonts\segoeuib.ttf",  True),
        (r"C:\Windows\Fonts\tahoma.ttf",   r"C:\Windows\Fonts\tahomabd.ttf",  True),
        (r"C:\Windows\Fonts\arial.ttf",    r"C:\Windows\Fonts\arialbd.ttf",   True),
        ("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
         "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf", False),
    ]:
        if os.path.isfile(rp):
            try:
                pdfmetrics.registerFont(TTFont("SxBody", rp))
                pdfmetrics.registerFont(TTFont("SxBold", bp if os.path.isfile(bp) else rp))
                reg, bld, ar = "SxBody", "SxBold", is_ar
                break
            except Exception:
                continue
    _pdf_fonts._c = (reg, bld, ar)
    return _pdf_fonts._c


def _ar(s):
    try:
        import arabic_reshaper
        from bidi.algorithm import get_display
        return get_display(arabic_reshaper.reshape(s))
    except Exception:
        return s


# ---- Model: write the English Q&A report --------------------------------
_PDF_QUESTIONS_EN = [
    "What is the current overall security posture and risk level?",
    "Which open vulnerabilities or CVEs pose the greatest immediate risk, and why?",
    "What is the potential business and operational impact of the identified threats?",
    "What immediate remediation steps must be taken, in priority order?",
    "What detection, monitoring, and preventive measures are recommended going forward?",
]
_PDF_QUESTIONS_AR = [
    "ما هو الوضع الأمني العام ومستوى المخاطر الحالي؟",
    "أيّ الثغرات المفتوحة أو CVE تشكّل الخطر الفوري الأكبر، ولماذا؟",
    "ما الأثر التشغيلي والتجاري المحتمل للتهديدات المرصودة؟",
    "ما الإجراءات العلاجية الفورية الواجبة، بحسب أولوية التنفيذ؟",
    "ما إجراءات الرصد والمراقبة والوقاية الموصى بها مستقبلاً؟",
]


def _model_qa(live_rep, lang):
    """Ask the model to answer the 5 CMAR questions. Returns list of str answers."""
    score = live_rep["score"]
    threats = live_rep.get("_threats", 0)
    with _sec_lock:
        opens = [f for f in _FINDINGS if f["st"] in _OPEN_STATES]
    findings_ctx = ("; ".join(
        f"{f['cve']} (severity {f['sev']}, CVSS {f['score']}) on {f['asset']} — fix: {f['fix_en']}"
        for f in opens[:8]) or "No open findings — all vulnerabilities resolved.")

    if lang == "ar":
        qs = "\n".join(f"## {q}" for q in _PDF_QUESTIONS_AR)
        prompt = (
            "أنت محلل أمن سيبراني خبير تكتب تقريراً رسمياً. "
            "القاعدة الإلزامية: اكتب بالعربية الفصحى السليمة حصرياً. "
            "ممنوع منعاً باتاً استخدام أي كلمة أو حرف إنجليزي (باستثناء أرقام ومعرّفات الثغرات مثل CVE). "
            "استخدم مصطلحات أمنية عربية صحيحة. أجب على كل سؤال بجملتين واضحتين ومهنيتين. "
            "ابدأ كل إجابة بسطر ## ثم نص الإجابة بالعربية مباشرةً:\n\n"
            f"البيانات الحية: درجة الأمان {score}/100، عدد التهديدات النشطة {threats}، "
            f"الثغرات: {findings_ctx}\n\n{qs}"
        )
    else:
        qs = "\n".join(f"## {q}" for q in _PDF_QUESTIONS_EN)
        prompt = (
            "You are a senior cybersecurity analyst. Answer each question below in 2 concise professional "
            "sentences based on the live data. Start each answer on a new line preceded by ##:\n\n"
            f"Live data: score {score}/100, active threats {threats}, "
            f"findings: {findings_ctx}\n\n{qs}"
        )
    res = model_chat(prompt, lang, max_tokens=900)
    raw = res.get("answer", "") if res.get("source") == "live" else ""
    # parse ## blocks
    answers = []
    cur = ""
    for line in raw.splitlines():
        if line.strip().startswith("##"):
            if cur.strip():
                answers.append(cur.strip())
            cur = ""
        else:
            cur += (" " if cur else "") + line.strip()
    if cur.strip():
        answers.append(cur.strip())
    # pad / fallback
    fallbacks_en = [
        f"Security score is {score}/100 with {threats} active threat(s). "
        + ("All vulnerabilities have been resolved." if not opens else
           f"The most severe open finding is {opens[0]['cve']} ({opens[0]['sev']}) on {opens[0]['asset']}."),
        findings_ctx,
        "Unresolved critical/high vulnerabilities expose systems to potential breach and service disruption.",
        "; ".join(r["en"] for r in _RECS if r["state"] == "open") or "Continue monitoring and patching.",
        "Enforce MFA, review access logs weekly, and schedule automated vulnerability scans.",
    ]
    fallbacks_ar = [
        f"درجة الأمان {score}/100 مع {threats} تهديد نشط. "
        + ("جميع الثغرات تمت معالجتها." if not opens else
           f"أخطر ثغرة مفتوحة: {opens[0]['cve']} (خطورة {opens[0]['sev']}) على {opens[0]['asset']}."),
        findings_ctx,
        "الثغرات الحرجة والعالية غير المعالَجة تُعرّض الأنظمة لاختراق محتمل وتعطّل الخدمات.",
        "؛ ".join(r["ar"] for r in _RECS if r["state"] == "open") or "متابعة المراقبة وتطبيق التحديثات.",
        "تفعيل المصادقة متعددة العوامل ومراجعة سجلات الوصول أسبوعياً وجدولة فحص آلي للثغرات.",
    ]
    fb = fallbacks_en if lang != "ar" else fallbacks_ar
    qs_list = _PDF_QUESTIONS_EN if lang != "ar" else _PDF_QUESTIONS_AR

    # quality guard for Arabic: Foundation-Sec is English-first, so if it produced
    # an answer that is mostly Latin characters (i.e. it ignored the "write in
    # Arabic" instruction), swap that answer for the hand-written Arabic fallback
    # so the report never shows broken/half-English Arabic text.
    if lang == "ar":
        cleaned = []
        for i, a in enumerate(answers):
            letters = [c for c in a if c.isalpha()]
            latin = sum(1 for c in letters if "a" <= c.lower() <= "z")
            arabic = sum(1 for c in letters if "\u0600" <= c <= "\u06FF")
            # too short, or more than ~35% Latin letters => use the clean fallback
            if len(a.strip()) < 15 or (letters and latin > max(3, 0.35 * len(letters)) and latin >= arabic):
                cleaned.append(fb[i] if i < len(fb) else a)
            else:
                cleaned.append(a)
        answers = cleaned

    while len(answers) < len(qs_list):
        answers.append(fb[len(answers)] if len(answers) < len(fb) else "—")
    return answers[:len(qs_list)]


# ---- Canvas helpers -------------------------------------------------------

def _pdf_build(live_rep, lang, page_label):
    """Render ONE language section as a LARGE, enterprise-grade security report:
    full cover page, executive summary, key metrics, risk-by-category with a bar
    chart, severity heat row, event timeline, detailed findings table, prioritized
    recommendations, and a model analysis section. Deep-blue institutional palette,
    minimal colour. Returns PDF bytes."""
    import io
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_LEFT, TA_RIGHT, TA_CENTER
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle, HRFlowable, PageBreak, KeepTogether,
                                    Flowable)
    from reportlab.lib.styles import ParagraphStyle
    import datetime

    reg, bld, ar_ok = _pdf_fonts()
    is_ar = (lang == "ar")           # language choice (controls text + alignment)
    shp = _ar if (is_ar and ar_ok) else (lambda x: x)   # shaping only if font supports it
    AL = TA_RIGHT if is_ar else TA_LEFT

    # ---- restrained corporate palette ----
    NAVY  = HexColor("#0E2438"); STEEL = HexColor("#2E4D6B")
    INK   = HexColor("#1B2733"); MID   = HexColor("#5C6E7E")
    LINE  = HexColor("#D4DCE4"); BG    = HexColor("#F3F6F9")
    WHITE = HexColor("#FFFFFF"); GOLD  = HexColor("#B8893A")
    CRIT  = HexColor("#A8201A"); HIGH  = HexColor("#BC5B00")
    MEDM  = HexColor("#1F5FB0"); LOW   = HexColor("#2E7D46"); GREY = HexColor("#7A8896")
    SEV = {"critical": CRIT, "high": HIGH, "medium": MEDM, "low": LOW}

    def tr(en, ar):
        return ar if is_ar else en

    def P(t, s):
        return Paragraph(shp(str(t)), s)

    # ---- styles ----
    s_cover_t  = ParagraphStyle("ct", fontName=bld, fontSize=30, textColor=WHITE, leading=36, alignment=AL)
    s_cover_s  = ParagraphStyle("cs", fontName=reg, fontSize=13, textColor=HexColor("#B9C8D8"), leading=18, alignment=AL)
    s_cover_m  = ParagraphStyle("cm", fontName=reg, fontSize=10, textColor=HexColor("#8DA0B4"), leading=15, alignment=AL)
    s_h1       = ParagraphStyle("h1", fontName=bld, fontSize=16, textColor=NAVY, leading=20, spaceBefore=4, spaceAfter=8, alignment=AL)
    s_h2       = ParagraphStyle("h2", fontName=bld, fontSize=12.5, textColor=STEEL, leading=16, spaceBefore=10, spaceAfter=4, alignment=AL)
    s_body     = ParagraphStyle("b", fontName=reg, fontSize=10, textColor=INK, leading=15.5, alignment=AL, spaceAfter=4)
    s_muted    = ParagraphStyle("m", fontName=reg, fontSize=8.5, textColor=MID, leading=12, alignment=AL)
    s_kpi_n    = ParagraphStyle("kn", fontName=bld, fontSize=21, textColor=NAVY, leading=23, alignment=TA_CENTER)
    s_kpi_l    = ParagraphStyle("kl", fontName=reg, fontSize=8, textColor=MID, leading=10.5, alignment=TA_CENTER)
    s_th       = ParagraphStyle("th", fontName=bld, fontSize=9, textColor=WHITE, leading=12, alignment=AL)
    s_td       = ParagraphStyle("td", fontName=reg, fontSize=8.5, textColor=INK, leading=11.5, alignment=AL)
    s_badge    = ParagraphStyle("bd", fontName=bld, fontSize=8, textColor=WHITE, leading=10.5, alignment=TA_CENTER)

    # ---- data ----
    score   = int(live_rep.get("score", posture_score()))
    threats = int(live_rep.get("_threats", active_threats()))
    title   = live_rep.get("title", tr("Security Assessment Report", "تقرير تقييم أمني"))
    now     = datetime.datetime.now()
    now_s   = now.strftime("%Y-%m-%d %H:%M")
    host    = _plat.node() or "host"
    with _sec_lock:
        opens = [f for f in _FINDINGS if f["st"] in _OPEN_STATES]
    counts = {"critical": 0, "high": 0, "medium": 0, "low": 0}
    for f in opens:
        counts[f.get("sev", "low")] = counts.get(f.get("sev", "low"), 0) + 1
    verdict = (tr("Weak", "ضعيف") if score < 60 else (tr("Moderate", "متوسط") if score < 80 else tr("Strong", "قوي")))
    vcol = CRIT if score < 60 else (HIGH if score < 80 else LOW)

    # recent security events (timeline)
    try:
        evs = events_query(limit=10)
    except Exception:
        evs = []

    # ===== page furniture =====
    def _chrome(canvas, doc):
        canvas.saveState()
        W, H = A4
        if doc.page == 1:
            # ---- full cover page drawn directly ----
            canvas.setFillColor(NAVY); canvas.rect(0, 0, W, H, fill=1, stroke=0)
            canvas.setFillColor(HexColor("#16314A")); canvas.rect(0, 0, W, 30 * mm, fill=1, stroke=0)
            canvas.setFillColor(GOLD); canvas.rect(0, H * 0.62, 55 * mm, 1.4 * mm, fill=1, stroke=0)
            mxr = W - 22 * mm   # right margin anchor for AR
            lab = shp(tr("CYBERSECURITY ASSESSMENT", "تقييم الأمن السيبراني"))
            ttl = shp(str(title))
            sub = shp(tr(f"Host: {host}", f"المضيف: {host}"))
            dt  = shp(tr(f"Report date: {now_s}", f"تاريخ التقرير: {now_s}"))
            cl  = shp(tr("Classification: CONFIDENTIAL", "التصنيف: سرّي"))
            ty = H * 0.62 - 6 * mm
            canvas.setFillColor(HexColor("#8DA0B4")); canvas.setFont(reg, 11)
            if is_ar: canvas.drawRightString(mxr, ty, lab)
            else:     canvas.drawString(22 * mm, ty, lab)
            canvas.setFillColor(WHITE); canvas.setFont(bld, 27)
            if is_ar: canvas.drawRightString(mxr, ty - 15 * mm, ttl)
            else:     canvas.drawString(22 * mm, ty - 15 * mm, ttl)
            canvas.setFillColor(HexColor("#B9C8D8")); canvas.setFont(reg, 12)
            if is_ar:
                canvas.drawRightString(mxr, ty - 30 * mm, sub)
                canvas.drawRightString(mxr, ty - 37 * mm, dt)
            else:
                canvas.drawString(22 * mm, ty - 30 * mm, sub)
                canvas.drawString(22 * mm, ty - 37 * mm, dt)
            canvas.setFillColor(GOLD); canvas.setFont(bld, 11)
            if is_ar: canvas.drawRightString(mxr, ty - 50 * mm, cl)
            else:     canvas.drawString(22 * mm, ty - 50 * mm, cl)
            # bottom confidentiality strip
            canvas.setFillColor(HexColor("#8DA0B4")); canvas.setFont(reg, 8.5)
            foot = shp(tr("This document contains confidential security information. Handle accordingly.",
                          "تحتوي هذه الوثيقة على معلومات أمنية سرّية. تعامل معها وفقاً لذلك."))
            if is_ar: canvas.drawRightString(mxr, 12 * mm, foot)
            else:     canvas.drawString(22 * mm, 12 * mm, foot)
            canvas.restoreState(); return
        # ---- inner page header/footer ----
        canvas.setFillColor(NAVY); canvas.rect(0, H - 18 * mm, W, 18 * mm, fill=1, stroke=0)
        canvas.setFillColor(GOLD); canvas.rect(0, H - 18 * mm, W, 0.8 * mm, fill=1, stroke=0)
        canvas.setFillColor(WHITE); canvas.setFont(bld, 9.5)
        lbl = shp(tr("SECURITY ASSESSMENT", "تقييم أمني"))
        if is_ar: canvas.drawRightString(W - 18 * mm, H - 11.5 * mm, lbl)
        else:     canvas.drawString(18 * mm, H - 11.5 * mm, lbl)
        canvas.setFont(reg, 8); canvas.setFillColor(HexColor("#9DB0C4"))
        cls = shp(tr("CONFIDENTIAL", "سرّي"))
        if is_ar: canvas.drawString(18 * mm, H - 11.5 * mm, cls)
        else:     canvas.drawRightString(W - 18 * mm, H - 11.5 * mm, cls)
        canvas.setStrokeColor(LINE); canvas.setLineWidth(0.5)
        canvas.line(18 * mm, 13 * mm, W - 18 * mm, 13 * mm)
        canvas.setFillColor(MID); canvas.setFont(reg, 7.5)
        canvas.drawString(18 * mm, 8.5 * mm, shp(tr(f"Generated {now_s}", f"أُنشئ {now_s}")))
        canvas.drawRightString(W - 18 * mm, 8.5 * mm, shp(tr(f"Page {doc.page}", f"صفحة {doc.page}")))
        canvas.restoreState()

    story = []

    # ---- small horizontal bar chart (pure flowable, no PNG) ----
    class BarRow(Flowable):
        def __init__(self, data, w=170 * mm, h=44 * mm):
            Flowable.__init__(self); self.data = data; self.width = w; self.height = h
        def draw(self):
            c = self.canv
            mx = max([v for _, v, _ in self.data] + [1])
            n = len(self.data); gap = 7 * mm
            bw = (self.width - gap * (n - 1)) / n
            for i, (lbl, val, col) in enumerate(self.data):
                x = i * (bw + gap)
                bh = (val / mx) * (self.height - 14 * mm)
                c.setFillColor(col); c.rect(x, 10 * mm, bw, bh, fill=1, stroke=0)
                c.setFillColor(INK); c.setFont(bld, 12)
                c.drawCentredString(x + bw / 2, 10 * mm + bh + 2 * mm, str(val))
                c.setFillColor(MID); c.setFont(reg, 8)
                c.drawCentredString(x + bw / 2, 4 * mm, shp(lbl))

    story = []

    # ===== COVER (blank first page; _chrome draws it) =====
    story.append(Spacer(1, 1))
    story.append(PageBreak())

    # ===== EXECUTIVE SUMMARY =====
    story.append(P(tr("Executive Summary", "الملخّص التنفيذي"), s_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=GOLD, spaceAfter=8))
    exec_txt = live_rep.get("summary") or tr(
        f"This assessment evaluated the security posture of {host}. The overall security score is "
        f"{score}/100 ({verdict}), with {threats} active threat(s) and {len(opens)} open finding(s) "
        f"requiring attention. Critical and high-severity items should be prioritized for remediation.",
        f"قيّم هذا التقرير الوضع الأمني للمضيف {host}. درجة الأمان الإجمالية {score}/١٠٠ ({verdict})، "
        f"مع {threats} تهديد نشط و{len(opens)} نتيجة مفتوحة تتطلّب المعالجة. يجب إعطاء الأولوية "
        f"للعناصر الحرجة والعالية الخطورة.")
    story.append(P(exec_txt, s_body))
    story.append(Spacer(1, 4 * mm))

    # KPI cards
    def kpi_cell(num, lbl, col=NAVY, fs=21):
        return Table([[P(num, ParagraphStyle("x", parent=s_kpi_n, textColor=col, fontSize=fs, leading=fs + 2))],
                      [P(lbl, s_kpi_l)]],
                     style=TableStyle([("TOPPADDING", (0, 0), (-1, -1), 1),
                                       ("BOTTOMPADDING", (0, 0), (-1, -1), 1)]))
    kpis = [[kpi_cell(f"{score}", tr("Security Score", "درجة الأمان")),
             kpi_cell(f"{threats}", tr("Active Threats", "تهديدات نشطة"), CRIT if threats else NAVY),
             kpi_cell(f"{len(opens)}", tr("Open Findings", "ثغرات مفتوحة")),
             kpi_cell(verdict, tr("Verdict", "التقييم"), vcol, fs=15)]]
    kt = Table(kpis, colWidths=[43 * mm] * 4)
    kt.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, -1), BG), ("BOX", (0, 0), (-1, -1), 0.5, LINE),
                            ("INNERGRID", (0, 0), (-1, -1), 0.5, LINE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                            ("TOPPADDING", (0, 0), (-1, -1), 9), ("BOTTOMPADDING", (0, 0), (-1, -1), 9)]))
    story.append(kt)
    story.append(Spacer(1, 6 * mm))

    # ===== RISK BY SEVERITY (chart) =====
    story.append(P(tr("Risk Distribution by Severity", "توزيع المخاطر حسب الخطورة"), s_h2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=LINE, spaceAfter=6))
    story.append(BarRow([
        (tr("Critical", "حرجة"), counts["critical"], CRIT),
        (tr("High", "عالية"), counts["high"], HIGH),
        (tr("Medium", "متوسطة"), counts["medium"], MEDM),
        (tr("Low", "منخفضة"), counts["low"], LOW),
    ]))
    story.append(Spacer(1, 4 * mm))

    # ===== narrative sections from payload =====
    for sec in live_rep.get("sections", []):
        story.append(P(sec.get("h", ""), s_h2))
        story.append(HRFlowable(width="100%", thickness=0.5, color=LINE, spaceAfter=5))
        story.append(P(sec.get("b", ""), s_body))

    story.append(PageBreak())

    # ===== DETAILED FINDINGS =====
    story.append(P(tr("Detailed Findings", "النتائج التفصيلية"), s_h1))
    story.append(HRFlowable(width="100%", thickness=1, color=GOLD, spaceAfter=8))
    if opens:
        hdr = ([tr("Remediation", "الإصلاح"), "CVSS", tr("Severity", "الخطورة"),
                tr("Asset", "الأصل"), tr("ID", "المُعرّف")] if is_ar
               else ["ID", "Asset", "Severity", "CVSS", "Remediation"])
        rows = [[P(h, s_th) for h in hdr]]
        sl = {"critical": tr("Critical", "حرجة"), "high": tr("High", "عالية"),
              "medium": tr("Medium", "متوسطة"), "low": tr("Low", "منخفضة")}
        for f in opens[:18]:
            sev = f.get("sev", "low")
            chip = Table([[P(sl.get(sev, sev), s_badge)]],
                         style=TableStyle([("BACKGROUND", (0, 0), (-1, -1), SEV.get(sev, LOW)),
                                           ("TOPPADDING", (0, 0), (-1, -1), 2), ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                                           ("LEFTPADDING", (0, 0), (-1, -1), 3), ("RIGHTPADDING", (0, 0), (-1, -1), 3)]))
            fix = ((f.get("fix_ar") or f.get("fix_en") or "") if is_ar else (f.get("fix_en") or ""))[:80]
            idv = P(f.get("cve", ""), s_td); asset = P(f.get("asset", ""), s_td)
            cvss = P(f.get("score", ""), s_td); fixp = P(fix, s_td)
            rows.append([fixp, cvss, chip, asset, idv] if is_ar else [idv, asset, chip, cvss, fixp])
        cw = ([58 * mm, 13 * mm, 19 * mm, 28 * mm, 32 * mm] if is_ar
              else [32 * mm, 28 * mm, 19 * mm, 13 * mm, 58 * mm])
        tb = Table(rows, colWidths=cw, repeatRows=1)
        tb.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), NAVY),
                                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, BG]),
                                ("LINEBELOW", (0, 0), (-1, -1), 0.3, LINE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                ("TOPPADDING", (0, 0), (-1, -1), 5), ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                                ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5)]))
        story.append(tb)
        if len(opens) > 18:
            story.append(Spacer(1, 2 * mm))
            story.append(P(tr(f"+{len(opens) - 18} additional findings", f"و{len(opens) - 18} نتيجة إضافية"), s_muted))
    else:
        story.append(P(tr("No open findings. All detected vulnerabilities have been resolved.",
                          "لا توجد نتائج مفتوحة. عولجت جميع الثغرات المكتشفة."), s_body))
    story.append(Spacer(1, 6 * mm))

    # ===== EVENT TIMELINE =====
    story.append(P(tr("Recent Security Events", "أحداث أمنية حديثة"), s_h2))
    story.append(HRFlowable(width="100%", thickness=0.5, color=LINE, spaceAfter=5))
    if evs:
        erows = [[P(tr("Time", "الوقت"), s_th), P(tr("Type", "النوع"), s_th), P(tr("Event", "الحدث"), s_th)]]
        if is_ar:
            erows = [[P(tr("Event", "الحدث"), s_th), P(tr("Type", "النوع"), s_th), P(tr("Time", "الوقت"), s_th)]]
        for e in evs[:10]:
            import datetime as _dt
            t = _dt.datetime.fromtimestamp(e.get("ts", 0) / 1000).strftime("%m-%d %H:%M")
            txt = (e.get("text_ar") or e.get("text_en") or "") if is_ar else (e.get("text_en") or e.get("text_ar") or "")
            kind = e.get("kind", "")
            if is_ar:
                erows.append([P(txt[:60], s_td), P(kind, s_td), P(t, s_td)])
            else:
                erows.append([P(t, s_td), P(kind, s_td), P(txt[:60], s_td)])
        cw2 = ([90 * mm, 26 * mm, 24 * mm] if is_ar else [24 * mm, 26 * mm, 90 * mm])
        et = Table(erows, colWidths=cw2, repeatRows=1)
        et.setStyle(TableStyle([("BACKGROUND", (0, 0), (-1, 0), STEEL),
                                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [WHITE, BG]),
                                ("LINEBELOW", (0, 0), (-1, -1), 0.3, LINE), ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                                ("TOPPADDING", (0, 0), (-1, -1), 4), ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
                                ("LEFTPADDING", (0, 0), (-1, -1), 5), ("RIGHTPADDING", (0, 0), (-1, -1), 5)]))
        story.append(et)
    else:
        story.append(P(tr("No recent events recorded.", "لا أحداث حديثة مسجّلة."), s_muted))

    story.append(PageBreak())

    # ===== MODEL ANALYSIS =====
    qa = _model_qa(live_rep, lang)
    qs = _PDF_QUESTIONS_AR if is_ar else _PDF_QUESTIONS_EN
    if qa:
        story.append(P(tr("Detailed Security Analysis", "التحليل الأمني التفصيلي"), s_h1))
        story.append(HRFlowable(width="100%", thickness=1, color=GOLD, spaceAfter=8))
        for i, (q, a) in enumerate(zip(qs, qa)):
            qs_style = ParagraphStyle(f"q{i}", fontName=bld, fontSize=11, textColor=STEEL,
                                      leading=15, alignment=AL, spaceBefore=10, spaceAfter=3)
            story.append(P(f"{i + 1}. {q}", qs_style))
            story.append(P(a, s_body))

    # ===== closing =====
    story.append(Spacer(1, 8 * mm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=LINE, spaceAfter=4))
    story.append(P(tr("— End of Report —", "— نهاية التقرير —"),
                   ParagraphStyle("end", fontName=reg, fontSize=9, textColor=MID, alignment=TA_CENTER)))

    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=26 * mm, bottomMargin=16 * mm,
                            leftMargin=18 * mm, rightMargin=18 * mm, title=str(title))
    doc.build(story, onFirstPage=_chrome, onLaterPages=_chrome)
    buf.seek(0)
    return buf.read()



def _build_pdf(live_rep, lang):
    """Build a professional bilingual PDF (EN + AR) from live system data."""
    import io
    from reportlab.pdfgen import canvas as _canvas
    from reportlab.lib.pagesizes import A4

    # always refresh with live data
    score = live_rep.get("score", posture_score())
    live_rep = dict(live_rep)
    live_rep["score"]    = posture_score()
    live_rep["_threats"] = active_threats()
    live_rep["tone"]     = ("danger" if live_rep["score"] < 60 else
                            ("warning" if live_rep["score"] < 80 else "success"))

    import copy
    rep_ar = copy.copy(live_rep)
    rep_ar["title"] = (live_rep.get("title_ar") or
                       ("تقرير الحالة" if "Status" in str(live_rep.get("title","")) else live_rep.get("title","")))
    rep_ar["summary"] = generate_report("status", "ar").get("summary", live_rep.get("summary",""))
    pages_en = _pdf_build(live_rep, "en", "Security Report")
    pages_ar = _pdf_build(rep_ar, "ar", "تقرير أمني")

    # merge both PDFs into one
    try:
        from reportlab.lib.pagesizes import A4
        from pypdf import PdfWriter, PdfReader
        import io
        merger = PdfWriter()
        for blob in (pages_en, pages_ar):
            merger.append(PdfReader(io.BytesIO(blob)))
        out = io.BytesIO()
        merger.write(out)
        out.seek(0)
        return out.read()
    except Exception:
        # fallback: just return English if pypdf missing
        return pages_en


def model_full_report(report, lang):
    """Kept for compatibility — PDF now uses _model_qa directly."""
    return report.get("sections") or [], False


# ==================================================================
#  CMAR-STYLE REPORT  — formal, model-written answers, browser-printed
# ==================================================================
# The official CMAR question set (FIRST Malware SIG), adapted so the model
# answers each one as a 25-year senior analyst, grounded in real case data.
_CMAR_Q = [
    ("exec",     "Executive Summary",
                 "الملخّص التنفيذي"),
    ("function", "What capability does the identified threat provide an attacker once active?",
                 "ما القدرات التي يمنحها التهديد المكتشف للمهاجم بمجرد نشاطه؟"),
    ("scope",    "Is this a known threat affecting multiple organizations, or are there indicators of a tailored, targeted attack?",
                 "هل هذا تهديد معروف يصيب عدة مؤسسات، أم توجد مؤشرات على هجوم موجّه مُصمّم خصيصاً؟"),
    ("ioc",      "What indicators of compromise (IOCs) are associated — network artifacts, processes/RAM, registry keys, and files created?",
                 "ما مؤشرات الاختراق المرتبطة — آثار الشبكة، العمليات/الذاكرة، مفاتيح السجلّ، والملفات المُنشأة؟"),
    ("persist",  "Does the threat maintain persistence on the system? If so, by what mechanism?",
                 "هل يحافظ التهديد على استمراريته في النظام؟ وإن كان كذلك فبأي آلية؟"),
    ("vuln",     "Which application, service, or vulnerability is exploited? Is it tied to a known CVE; does a patch exist; and does endpoint protection defend against it?",
                 "أي تطبيق أو خدمة أو ثغرة يجري استغلالها؟ وهل ترتبط بـ CVE معروف، وهل يوجد تصحيح، وهل تحمي حماية النقاط الطرفية منها؟"),
    ("remediate","What remediation options will effectively eliminate the threat and return the system to a fully secure state, in priority order?",
                 "ما خيارات المعالجة التي ستزيل التهديد فعلياً وتُعيد النظام إلى حالة آمنة تماماً، بحسب الأولوية؟"),
]

# ----- additional report templates (executive briefing, incident response, compliance audit) -----
_TEMPLATES = {
    "cmar": {
        "title_en": "Cyber Malware Analysis Report",
        "title_ar": "تقرير تحليل البرمجيات الخبيثة",
        "questions": _CMAR_Q,
        "sys_en": (" You are a principal cyber-threat analyst with 25 years authoring official CMAR incident reports for "
                   "national response teams. Write in a precise, authoritative, in-depth institutional tone. For EACH question "
                   "write TWO to THREE paragraphs (4-7 sentences each): explain the technical mechanism precisely, the likely "
                   "impact on confidentiality/integrity/availability, real-world attack chains tied to each named CVE, the "
                   "evidence/indicators, and the actions. Ground everything in the actual case data, CVE IDs and asset names. "
                   "No filler, no repetition — every sentence carries information."),
        "sys_ar": (" أنت محلّل تهديدات سيبرانية رئيسي بخبرة 25 عاماً في كتابة تقارير الحوادث الرسمية (CMAR) للجهات الوطنية. "
                   "اكتب بالعربية الفصحى فقط بأسلوب مؤسسي رصين وعميق. لكل سؤال اكتب من فقرتين إلى ثلاث فقرات (٤-٧ جمل لكل فقرة): "
                   "اشرح الآلية التقنية، والأثر على السرّية والسلامة والتوافر، وسلاسل الهجوم الواقعية لكل CVE، والأدلّة، ثم الإجراءات. "
                   "استند إلى بيانات الحالة وأسماء الـ CVE والأصول الفعلية. لا حشو ولا تكرار."),
    },
    "executive": {
        "title_en": "Executive Security Briefing",
        "title_ar": "إحاطة أمنية تنفيذية",
        "questions": [
            ("exec",       "Executive Summary",
                           "الملخّص التنفيذي"),
            ("posture",    "What is the current overall security posture, and what is the trend?",
                           "ما الوضع الأمني الإجمالي الحالي، وما اتّجاهه؟"),
            ("topRisks",   "What are the top 3-5 risks today, and why does each one matter in business terms?",
                           "ما أهم ٣-٥ مخاطر اليوم، ولماذا يهم كل واحد منها من منظور الأعمال؟"),
            ("impact",     "What is the realistic business impact (financial, operational, reputational) if these risks materialize?",
                           "ما الأثر الواقعي على الأعمال (مالياً، تشغيلياً، على السمعة) إن تحقّقت هذه المخاطر؟"),
            ("recommend",  "What are the recommended decisions for leadership in priority order, including resource and budget asks?",
                           "ما القرارات الموصى بها للقيادة بحسب الأولوية، بما فيها متطلّبات الموارد والميزانية؟"),
        ],
        "sys_en": (" You are the CISO briefing executive leadership. Each answer: 1-2 short paragraphs, plain business language "
                   "with minimal technical jargon, focused on outcomes and decisions. Ground every claim in concrete numbers and "
                   "names from the case data (score, finding counts, asset names). Be candid, not alarmist."),
        "sys_ar": (" أنت مدير أمن المعلومات (CISO) تخاطب القيادة التنفيذية. كل إجابة: فقرة أو اثنتان قصيرتان بلغة أعمال واضحة، "
                   "بأقل قدر من المصطلحات التقنية، مع تركيز على النتائج والقرارات. استند في كل ادّعاء إلى أرقام وأسماء محدّدة من "
                   "بيانات الحالة (الدرجة، عدد الثغرات، أسماء الأصول). كن صريحاً دون إثارة ذعر."),
    },
    "incident": {
        "title_en": "Incident Response Report",
        "title_ar": "تقرير الاستجابة للحادث",
        "questions": [
            ("exec",         "Executive Summary",
                             "الملخّص التنفيذي"),
            ("detection",    "How was the incident detected — when, by which sensor/event, and what triggered the alert?",
                             "كيف اكتُشف الحادث — متى، وبأي حسّاس/حدث، وما الذي أطلق التنبيه؟"),
            ("timeline",     "What is the timeline of events from initial access to detection (with timestamps)?",
                             "ما الخط الزمني للأحداث من الوصول الأولي إلى الاكتشاف (مع الطوابع الزمنية)؟"),
            ("scope",        "What is the scope of impact — which hosts, accounts, data, and services are affected?",
                             "ما نطاق الأثر — أي الأجهزة والحسابات والبيانات والخدمات متأثّرة؟"),
            ("containment",  "What containment steps were taken (or should be taken) to stop the spread?",
                             "ما خطوات الاحتواء المُتّخذة (أو الواجب اتّخاذها) لوقف الانتشار؟"),
            ("eradication",  "How is the threat being eradicated from affected systems (specific actions per host)?",
                             "كيف يُستأصل التهديد من الأنظمة المتأثّرة (إجراءات محدّدة لكل جهاز)؟"),
            ("recovery",     "What are the recovery steps and the verification that systems are fully restored?",
                             "ما خطوات التعافي والتحقّق من أن الأنظمة عادت إلى حالتها بالكامل؟"),
            ("lessons",      "What are the lessons learned and the specific improvements to detection, prevention, and process?",
                             "ما الدروس المستفادة والتحسينات المحدّدة في الكشف والوقاية والإجراءات؟"),
        ],
        "sys_en": (" You are the incident commander writing a post-incident report. Use precise IR vocabulary (Detection, "
                   "Containment, Eradication, Recovery, Lessons). Tie every answer to the actual events, findings, asset names, "
                   "and timestamps in the case data. 2 paragraphs per answer. Concrete and chronological."),
        "sys_ar": (" أنت قائد فريق الاستجابة للحوادث تكتب تقريراً بعد الحادث. استخدم مصطلحات IR بدقّة (الكشف، الاحتواء، "
                   "الاستئصال، التعافي، الدروس). اربط كل إجابة بالأحداث والثغرات وأسماء الأصول والطوابع الزمنية الفعلية في بيانات "
                   "الحالة. فقرتان لكل إجابة. محدّد وزمني."),
    },
    "compliance": {
        "title_en": "Compliance Audit Report",
        "title_ar": "تقرير تدقيق الامتثال",
        "questions": [
            ("exec",       "Executive Summary",
                           "الملخّص التنفيذي"),
            ("scope",      "What frameworks/controls were assessed (CIS Benchmarks, ISO 27001, NIST CSF) and what version?",
                           "ما الأطر/الضوابط التي قُيّمت (CIS، ISO 27001، NIST CSF) وأي إصدار؟"),
            ("passed",     "Which controls passed, with the supporting evidence (configuration values, registry settings, policies)?",
                           "ما الضوابط التي اجتيزت، مع الأدلة المؤيّدة (قيم التهيئة، إعدادات السجلّ، السياسات)؟"),
            ("failed",     "Which controls failed, with the specific configuration gap for each (cite CIS IDs)?",
                           "ما الضوابط التي فشلت، مع الفجوة التهيئية المحدّدة لكلٍّ منها (اذكر معرّفات CIS)؟"),
            ("remediate",  "What remediation steps close each gap, in priority order with effort estimates?",
                           "ما خطوات المعالجة التي تُغلق كل فجوة، بحسب الأولوية مع تقديرات الجهد؟"),
            ("residual",   "What is the residual risk after remediation, and what compensating controls are required?",
                           "ما الخطر المتبقّي بعد المعالجة، وما الضوابط التعويضية المطلوبة؟"),
        ],
        "sys_en": (" You are a senior IT auditor. Cite control IDs explicitly (e.g. CIS 9.1, CIS 18.9). State findings as factual "
                   "gaps, with the exact configuration evidence and the required remediation. Use a formal audit register tone. "
                   "1-2 paragraphs per answer."),
        "sys_ar": (" أنت مدقّق تقنية معلومات أول. استشهد بمعرّفات الضوابط صراحةً (مثل CIS 9.1 و CIS 18.9). اذكر النتائج كفجوات "
                   "وقائعية، مع دليل التهيئة الدقيق والمعالجة المطلوبة. استخدم نبرة سجلّ تدقيق رسمي. فقرة إلى فقرتان لكل إجابة."),
    },
}


def _template_meta(template):
    return _TEMPLATES.get(template) or _TEMPLATES["cmar"]


def _template_questions(template):
    return _template_meta(template)["questions"]


def _cmar_context(report, lang):
    """Real case data the analyst/model reasons over."""
    score = posture_score()
    threats = active_threats()
    with _sec_lock:
        opens = [dict(f) for f in _FINDINGS if f["st"] in _OPEN_STATES]
    # recent IOC-like events (scans / new findings)
    recent = events_in_window(_PERIOD_SECONDS.get(report.get("period"), 2592000)) if report.get("period") in _PERIOD_SECONDS else _EVENTS[-12:]
    return score, threats, opens, recent


def _cmar_model_answers(report, lang, template="cmar"):
    """ONE model call producing question-specific writeups for the chosen template.
    Falls back to concise real-data answers only if the model is unavailable."""
    score, threats, opens, recent = _cmar_context(report, lang)
    meta = _template_meta(template)
    qs = meta["questions"]
    findings_txt = ("؛ ".join(f"{f['cve']} ({f['sev']}, CVSS {f['score']}) على {f['asset']} — الحل: {f['fix_ar']}" for f in opens) if lang=="ar" else
                    "; ".join(f"{f['cve']} ({f['sev']}, CVSS {f['score']}) on {f['asset']} — fix: {f['fix_en']}" for f in opens)) or ("لا ثغرات مفتوحة" if lang=="ar" else "no open findings")
    # extra local context for the compliance template
    extra_ctx = ""
    if template == "compliance":
        try:
            comp = _run_compliance_check()
            extra_ctx = ("\nCIS controls: " +
                         "; ".join(f"{c['id']} {c['title_en']}: {c['status']}{(' ['+c['detail']+']') if c.get('detail') else ''}"
                                   for c in comp.get("controls", [])))
        except Exception:
            pass
    if lang == "ar":
        q_block = "\n".join(f"## {i}\n{q[2]}" for i, q in enumerate(qs))
        sys_extra = meta["sys_ar"]
        prompt = (f"بيانات الحالة الحقيقية:\nدرجة الأمان: {score}/100\nالتهديدات النشطة: {threats}\nالثغرات المفتوحة: {findings_txt}{extra_ctx}\n\n"
                  f"أجب عن كل سؤال بعمق وتفصيل. ابدأ كل إجابة بسطر يحوي ## ثم رقم السؤال، ثم الإجابة المطوّلة في الأسطر التالية:\n{q_block}")
    else:
        q_block = "\n".join(f"## {i}\n{q[1]}" for i, q in enumerate(qs))
        sys_extra = meta["sys_en"]
        prompt = (f"Real case data:\nSecurity score: {score}/100\nActive threats: {threats}\nOpen findings: {findings_txt}{extra_ctx}\n\n"
                  f"Answer each question in depth. Begin each answer with a line containing ## then the question number, "
                  f"then the detailed multi-paragraph answer on the following lines:\n{q_block}")
    res = model_chat(prompt, lang, max_tokens=2200, system_extra=sys_extra)
    answers = {}
    if res.get("source") == "live" and res.get("answer"):
        cur = None; buf = []
        for line in res["answer"].splitlines():
            st = line.strip()
            m = _re.match(r"^#+\s*(\d+)", st)
            if m:
                if cur is not None:
                    answers[cur] = "\n".join(buf).strip()
                cur = int(m.group(1)); buf = []
                rest = _re.sub(r"^#+\s*\d+[\.\):]?\s*", "", st)
                if rest:
                    buf.append(rest)
            elif cur is not None:
                buf.append(st)          # keep blank lines as paragraph breaks
        if cur is not None:
            answers[cur] = "\n".join(buf).strip()
    fb = _cmar_fallback(score, threats, opens, lang, template=template)
    out = []
    for i, q in enumerate(qs):
        a = answers.get(i, "").strip()
        out.append({"key": q[0], "q": (q[2] if lang == "ar" else q[1]), "a": a or fb[i]})
    return out, (res.get("source") == "live")


def _cmar_fallback(score, threats, opens, lang, template="cmar"):
    """Concise real-data fallback answers, one per question in the chosen template."""
    top = opens[0] if opens else None
    ar = (lang == "ar")
    qs = _template_questions(template)
    # short, factual fallback per question key
    def f(key):
        if key == "exec":
            return (f"بلغت درجة الأمان {score}/100 مع {threats} تهديد نشط. " +
                    ("الوضع سليم ولا ثغرات مفتوحة." if not opens else f"أبرز خطر مفتوح {top['cve']} على {top['asset']}.")) if ar else \
                   (f"Security score is {score}/100 with {threats} active threat(s). " +
                    ("Posture is clean with no open findings." if not opens else f"The leading open risk is {top['cve']} on {top['asset']}."))
        if key == "function":
            return ("لا تهديد فاعل حالياً." if not opens else "قد تمنح الثغرات المفتوحة المهاجم تنفيذ أوامر عن بُعد أو رفع امتيازات.") if ar else \
                   ("No active threat capability at present." if not opens else "Open vulnerabilities could grant RCE, privilege escalation, or data access.")
        if key == "scope":
            return "لا مؤشرات على هجوم موجّه؛ الثغرات المرصودة عامة ومعروفة." if ar else \
                   "No indicators of a tailored attack; the observed findings are common, publicly known issues."
        if key == "ioc":
            return ("لا مؤشرات اختراق نشطة." if not opens else "؛ ".join(f"{f['cve']} على {f['asset']}" for f in opens[:5])) if ar else \
                   ("No active indicators of compromise." if not opens else "; ".join(f"{f['cve']} on {f['asset']}" for f in opens[:5]))
        if key == "persist":
            return ("لا آلية استمرارية مرصودة." if not opens else "تُعالَج الثغرات قبل أن تُتيح استمرارية للمهاجم.") if ar else \
                   ("No persistence mechanism observed." if not opens else "Findings are being remediated before they can establish persistence.")
        if key == "vuln":
            return ("لا ثغرات قيد الاستغلال." if not opens else "؛ ".join(f"{f['cve']} (CVSS {f['score']})" for f in opens[:5])) if ar else \
                   ("No vulnerabilities under active exploitation." if not opens else "; ".join(f"{f['cve']} (CVSS {f['score']})" for f in opens[:5]))
        if key == "remediate" or key == "recommend":
            return ("الاستمرار في المراقبة والتحديث الدوري." if not opens else "؛ ".join(f"{f['fix_ar']} على {f['asset']}" for f in opens[:5])) if ar else \
                   ("Continue monitoring and routine patching." if not opens else "; ".join(f"{f['fix_en']} on {f['asset']}" for f in opens[:5]))
        # executive template
        if key == "posture":
            return (f"الدرجة الإجمالية {score}/100. " + (f"{len(opens)} ثغرة مفتوحة." if opens else "لا ثغرات مفتوحة.")) if ar else \
                   (f"Overall score {score}/100. " + (f"{len(opens)} open finding(s)." if opens else "No open findings."))
        if key == "topRisks":
            return ("لا مخاطر جوهرية حالياً." if not opens else "؛ ".join(f"{f['cve']} ({f['sev']})" for f in opens[:5])) if ar else \
                   ("No material risks at present." if not opens else "; ".join(f"{f['cve']} ({f['sev']})" for f in opens[:5]))
        if key == "impact":
            return ("الأثر المباشر منخفض." if not opens else "الأثر المحتمل: تعطّل خدمات، تسريب بيانات، أو فدية.") if ar else \
                   ("Direct impact is low at present." if not opens else "Potential impact: service disruption, data leakage, or ransomware.")
        # incident template
        if key == "detection":
            return "اكتشاف من قارئ سجلّات ويندوز ومحرّك المسح المحلي." if ar else \
                   "Detected by the Windows event-log reader and local scanning engine."
        if key == "timeline":
            return ("لا حوادث مسجّلة في النافذة." if not opens else f"رُصد {len(opens)} نتيجة مفتوحة في آخر فحص.") if ar else \
                   ("No incidents recorded in the window." if not opens else f"{len(opens)} open finding(s) at last scan.")
        if key == "containment":
            return "عزل الأجهزة المتأثّرة، وحظر العناوين المشبوهة في الجدار الناري." if ar else \
                   "Isolate affected hosts; block suspicious IPs at the firewall."
        if key == "eradication":
            return ("لا تهديد فاعل يستلزم استئصالاً." if not opens else "ترقيع/إزالة المكوّنات المُصابة، فحص جذور (rootkit)، وإعادة بناء عند الشك.") if ar else \
                   ("No active threat requires eradication." if not opens else "Patch/remove affected components, rootkit-scan, and rebuild on suspicion.")
        if key == "recovery":
            return "استعادة الخدمات تدريجياً مع المراقبة المضاعفة لمدة ٧٢ ساعة." if ar else \
                   "Phased service restoration with heightened monitoring for 72 hours."
        if key == "lessons":
            return "تحسين كشف PowerShell المُشفّر، وفرض MFA، وتقليل المنافذ المعرّضة." if ar else \
                   "Improve encoded-PowerShell detection; enforce MFA; reduce exposed ports."
        # compliance template
        if key == "passed":
            try:
                comp = _run_compliance_check()
                p = [c for c in comp.get("controls", []) if c["status"] == "pass"]
                return ("؛ ".join(f"{c['id']} {c['title_ar']}" for c in p[:6]) if p else "لا ضوابط ناجحة مسجّلة.") if ar else \
                       ("; ".join(f"{c['id']} {c['title_en']}" for c in p[:6]) if p else "No controls recorded as passed.")
            except Exception:
                return "—"
        if key == "failed":
            try:
                comp = _run_compliance_check()
                f_ = [c for c in comp.get("controls", []) if c["status"] == "fail"]
                return ("؛ ".join(f"{c['id']} {c['title_ar']} — {c.get('detail','')}" for c in f_[:6]) if f_ else "لا ضوابط فاشلة.") if ar else \
                       ("; ".join(f"{c['id']} {c['title_en']} — {c.get('detail','')}" for c in f_[:6]) if f_ else "No failing controls.")
            except Exception:
                return "—"
        if key == "residual":
            return "الخطر المتبقّي منخفض بعد إغلاق الفجوات؛ يُوصى بمراجعة ربع سنوية." if ar else \
                   "Residual risk is low after closing the gaps; quarterly review recommended."
        return ""
    return [f(q[0]) for q in qs]


# ---- background generation: page polls until the deep report is ready ----
_cmar_jobs = {}            # "rid:lang:template" -> {status, answers, by_model, started}
_cmar_jobs_lock = threading.Lock()


def _cmar_fallback_answers(rid, lang, template="cmar"):
    try:
        rep = _REPORTS.get(rid) or {"period": "status"}
        score, threats, opens, _ = _cmar_context(rep, lang)
    except Exception:
        score, threats, opens = posture_score(), active_threats(), []
    fb = _cmar_fallback(score, threats, opens, lang, template=template)
    qs = _template_questions(template)
    return [{"key": q[0], "q": (q[2] if lang == "ar" else q[1]), "a": fb[i]}
            for i, q in enumerate(qs)]


def _run_cmar_job(rid, lang, template="cmar"):
    key = f"{rid}:{lang}:{template}"
    try:
        rep = _REPORTS.get(rid) or generate_report("status", lang)
        answers, by_model = _cmar_model_answers(rep, lang, template=template)
    except Exception:
        answers, by_model = _cmar_fallback_answers(rid, lang, template=template), False
    if by_model and rid in _REPORTS:                       # persist the model write-up
        _REPORTS[rid]["_cmar_" + lang + "_" + template] = answers
        _save_reports()
    with _cmar_jobs_lock:
        _cmar_jobs[key] = {"status": "ready", "answers": answers,
                           "by_model": by_model, "started": _cmar_jobs.get(key, {}).get("started", time.time())}


@bp.get("/api/report/cmar")
def report_cmar_data():
    """Job-based: returns {status:'generating',elapsed} until the deep model report
    is ready, then {status:'ready',answers,byModel}. ?fast=1 returns instant fallback.
    ?template=cmar|executive|incident|compliance selects the question set."""
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    rid = request.args.get("id", "")
    template = (request.args.get("template") or "cmar").lower()
    if template not in _TEMPLATES:
        template = "cmar"

    if request.args.get("fast"):
        return jsonify({"status": "ready", "answers": _cmar_fallback_answers(rid, lang, template=template), "byModel": False, "template": template})

    rep = _REPORTS.get(rid)
    ck = "_cmar_" + lang + "_" + template
    if rep and rep.get(ck):
        return jsonify({"status": "ready", "answers": rep[ck], "byModel": True, "template": template})

    key = f"{rid}:{lang}:{template}"
    with _cmar_jobs_lock:
        job = _cmar_jobs.get(key)
        if job and job.get("status") == "ready":
            return jsonify({"status": "ready", "answers": job["answers"], "byModel": job["by_model"], "template": template})
        if not job:
            _cmar_jobs[key] = {"status": "generating", "started": time.time()}
            threading.Thread(target=_run_cmar_job, args=(rid, lang, template), daemon=True).start()
        started = _cmar_jobs[key]["started"]
    elapsed = int(time.time() - started)
    # HARD TIMEOUT: the local model can be very slow (or stalled) on CPU. After
    # 90s we stop waiting and return the instant data-driven fallback report so
    # the user is never stuck on the spinner. The model job keeps running and
    # will be used next time if it finishes.
    if elapsed >= 90:
        return jsonify({"status": "ready", "template": template, "byModel": False,
                        "timedout": True,
                        "answers": _cmar_fallback_answers(rid, lang, template=template)})
    return jsonify({"status": "generating", "elapsed": elapsed, "template": template})


def _report_html(report, lang, template="cmar"):
    """Formal report (CMAR / executive / incident / compliance). Structure + real
    data render instantly; the model-written answers load via AJAX. Browser renders
    Arabic/RTL perfectly and prints to PDF (Ctrl+P -> Save as PDF)."""
    import html as _h, datetime
    ar = (lang == "ar")
    meta = _template_meta(template)
    template_title = meta["title_ar"] if ar else meta["title_en"]
    score = posture_score(); threats = active_threats()
    with _sec_lock:
        opens = [dict(f) for f in _FINDINGS if f["st"] in _OPEN_STATES]
    analyst = "—"
    with _auth_lock:
        if _current_user[0] and _current_user[0] in _USERS:
            analyst = _USERS[_current_user[0]]["name"]
    now = datetime.datetime.fromtimestamp((report.get("ts") or int(time.time()*1000))/1000)
    date_s = now.strftime("%Y-%m-%d %H:%M")
    rid = report.get("id", "")
    doc_id = "SOC-" + now.strftime("%Y%m%d") + "-" + str(rid).split("-")[-1].zfill(3)
    title = report.get("title") or template_title
    classification = "سرّي" if ar else "CONFIDENTIAL"

    L = {
        "org": "مركز العمليات الأمنية — Sentinel" if ar else "Sentinel — Security Operations Centre",
        "rtype": "تقرير تحليل حادثة" if ar else "Incident Analysis Report",
        "cls": "التصنيف" if ar else "Classification",
        "docid": "رقم الوثيقة" if ar else "Document ID",
        "exec": "الملخّص التنفيذي" if ar else "Executive Summary",
        "case": "تفاصيل الحالة" if ar else "Case Details",
        "date": "التاريخ" if ar else "Date",
        "analyst": "المحلّل" if ar else "Analyst",
        "score": "درجة الأمان" if ar else "Security Score",
        "period": "النطاق" if ar else "Scope",
        "sir": "متطلّبات المعلومات الأساسية" if ar else "Standing Information Requirements",
        "findings": "النتائج والثغرات المفتوحة" if ar else "Findings & Open Vulnerabilities",
        "cve": "المُعرّف" if ar else "CVE", "asset": "الأصل" if ar else "Asset",
        "sev": "الخطورة" if ar else "Severity", "cvss": "CVSS",
        "fix": "الإجراء التصحيحي" if ar else "Remediation",
        "none": "لا توجد ثغرات مفتوحة — جميعها مُعالَجة." if ar else "No open vulnerabilities — all have been remediated.",
        "notes": "ملاحظات المحلّل الإضافية" if ar else "Additional Examiner Notes",
        "iocs": "مؤشرات الاختراق (IOCs)" if ar else "Indicators of Compromise (IOCs)",
        "conf": "هذه الوثيقة سرّية ومخصّصة للاستخدام الداخلي فقط." if ar else "This document is confidential and intended for internal use only.",
        "print": "طباعة / حفظ PDF" if ar else "Print / Save as PDF",
        "gen": "يقوم المحلّل الآلي بكتابة التحليل…" if ar else "Analyst engine is composing the assessment…",
        "page": "صفحة" if ar else "Page",
    }
    tpl_qs = _template_questions(template)
    qrows = "".join(
        f"<section class='q' data-i='{i}'><h3>{('' if ar else str(i)+'. ')}{_h.escape(q[2] if ar else q[1])}</h3>"
        f"<p class='ans' id='ans{i}'><span class='gen'>{L['gen']}</span></p></section>"
        for i, q in enumerate(tpl_qs) if i != 0)

    if opens:
        frows = "".join(
            f"<tr><td>{_h.escape(f['cve'])}</td><td>{_h.escape(f['asset'])}</td>"
            f"<td>{_h.escape(f['sev'])}</td><td>{f['score']}</td>"
            f"<td>{_h.escape(f['fix_ar'] if ar else f['fix_en'])}</td></tr>" for f in opens)
        findings_tbl = (f"<table class='ftbl'><thead><tr><th>{L['cve']}</th><th>{L['asset']}</th>"
                        f"<th>{L['sev']}</th><th>{L['cvss']}</th><th>{L['fix']}</th></tr></thead><tbody>{frows}</tbody></table>")
    else:
        findings_tbl = f"<p class='clean'>{L['none']}</p>"

    iocs = []
    for f in opens:
        iocs.append(f"{f['cve']} — {f['asset']}")
    iocs_html = ("<br>".join(_h.escape(x) for x in iocs)) if iocs else ("—")

    align = "right" if ar else "left"
    return f"""<!doctype html><html dir="{'rtl' if ar else 'ltr'}" lang="{lang}"><head><meta charset="utf-8">
<title>{_h.escape(title)} — {doc_id}</title><style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:'Segoe UI','Tahoma',Arial,sans-serif;color:#111;background:#e9eaec;line-height:1.65;font-size:13px}}
.sheet{{max-width:800px;margin:22px auto;background:#fff;padding:0 0 40px;box-shadow:0 1px 10px rgba(0,0,0,.15)}}
.masthead{{border-bottom:3px double #111;padding:26px 46px 16px;display:flex;justify-content:space-between;align-items:flex-end}}
.masthead .org{{font-size:12px;letter-spacing:.14em;text-transform:uppercase;color:#333}}
.masthead .ttl{{font-size:22px;font-weight:700;margin-top:6px;font-family:Georgia,'Times New Roman',serif}}
.masthead .meta{{font-size:10.5px;color:#444;text-align:{'left' if ar else 'right'};line-height:1.5}}
.cls-strip{{background:#111;color:#fff;text-align:center;font-size:10.5px;letter-spacing:.32em;padding:4px;text-transform:uppercase}}
.body{{padding:24px 46px}}
h2{{font-family:Georgia,'Times New Roman',serif;font-size:14.5px;text-transform:uppercase;letter-spacing:.04em;border-bottom:1.5px solid #111;padding-bottom:4px;margin:26px 0 12px}}
table.kv{{width:100%;border-collapse:collapse;margin-bottom:8px}}
table.kv td{{border:1px solid #999;padding:7px 11px;font-size:12.5px}}
table.kv td.k{{background:#ededed;font-weight:700;width:30%;white-space:nowrap}}
.exec-box{{border:1px solid #999;border-{'right' if ar else 'left'}:4px solid #111;background:#fafafa;padding:13px 16px;font-size:13px}}
table.ftbl{{width:100%;border-collapse:collapse;font-size:12px;margin-top:4px}}
table.ftbl th{{background:#222;color:#fff;padding:7px 9px;text-align:{align};font-weight:600;font-size:11px;letter-spacing:.03em}}
table.ftbl td{{border:1px solid #bbb;padding:6px 9px}}
table.ftbl tbody tr:nth-child(even){{background:#f4f4f4}}
.clean{{padding:8px 0;font-weight:600}}
section.q{{margin:14px 0}}
section.q h3{{font-size:13px;font-weight:700;margin-bottom:4px;color:#111}}
section.q .ans{{font-size:13px;color:#1c1c1c;text-align:justify}}
.gen{{color:#888;font-style:italic}}
.foot{{border-top:1px solid #999;margin:30px 46px 0;padding-top:8px;font-size:10px;color:#555;display:flex;justify-content:space-between}}
.bar{{position:fixed;top:0;{'left' if ar else 'right'}:0;margin:14px;z-index:9;display:none}}
.bar button{{background:#111;color:#fff;border:0;padding:11px 20px;border-radius:5px;font-size:13px;font-weight:600;cursor:pointer}}
#prep{{position:fixed;top:0;left:0;right:0;bottom:0;background:#e9eaec;z-index:50;display:flex;flex-direction:column;align-items:center;justify-content:center;text-align:center;padding:24px}}
#prep .spin{{width:46px;height:46px;border:4px solid #ccc;border-top-color:#111;border-radius:50%;animation:sx 0.9s linear infinite;margin-bottom:20px}}
@keyframes sx{{to{{transform:rotate(360deg)}}}}
#prep .pt{{font-size:18px;font-weight:700;color:#111;font-family:Georgia,serif}}
#prep .ps{{font-size:13px;color:#555;margin-top:8px;max-width:420px;line-height:1.7}}
#prep .pe{{font-size:12px;color:#888;margin-top:14px;font-variant-numeric:tabular-nums}}
@media print{{body{{background:#fff}}.sheet{{box-shadow:none;margin:0;max-width:100%}}.bar,#prep{{display:none !important}}@page{{margin:14mm}}}}
</style></head><body>
<div id="prep">
  <div class="spin"></div>
  <div class="pt">{'جارٍ تجهيز التقرير العميق' if ar else 'Preparing the in-depth report'}</div>
  <div class="ps">{'يكتب المحلّل الآلي تحليلاً تفصيلياً لكل بند. قد يستغرق هذا عدّة دقائق على النموذج المحلي. سيفتح التقرير تلقائياً عند جهوزه.' if ar else 'The AI analyst is writing a detailed assessment for each item. This can take a few minutes on the local model. The report will open automatically once ready.'}</div>
  <div class="pe" id="prepElapsed">00:00</div>
</div>
<div class="bar"><button onclick="window.print()">{L['print']}</button></div>
<div class="sheet">
  <div class="cls-strip">{classification}</div>
  <div class="masthead">
    <div><div class="org">{L['org']}</div><div class="ttl">{_h.escape(title)}</div></div>
    <div class="meta">{L['docid']}: {doc_id}<br>{date_s}</div>
  </div>
  <div class="body">
    <h2>{L['case']}</h2>
    <table class="kv">
      <tr><td class="k">{L['date']}</td><td>{date_s}</td><td class="k">{L['analyst']}</td><td>{_h.escape(analyst)}</td></tr>
      <tr><td class="k">{L['score']}</td><td>{score}/100</td><td class="k">{L['cls']}</td><td>{classification}</td></tr>
    </table>

    <h2>{L['exec']}</h2>
    <div class="exec-box" id="execbox"><span class="gen">{L['gen']}</span></div>

    <h2>{L['findings']}</h2>
    {findings_tbl}

    <h2>{L['sir']}</h2>
    {qrows}

    <h2>{L['notes']} — {L['iocs']}</h2>
    <p style="font-family:monospace;font-size:12px">{iocs_html}</p>
  </div>
  <div class="foot"><span>{L['conf']}</span><span>{doc_id}</span></div>
</div>
<script>
(function(){{
  var id={json.dumps(rid)}, lang={json.dumps(lang)}, doPrint={json.dumps(bool(request.args.get('print')))};
  var ar={json.dumps(ar)};
  var t0=Date.now();
  function pad(n){{return (n<10?'0':'')+n;}}
  var tick=setInterval(function(){{
    var s=Math.floor((Date.now()-t0)/1000);
    var el=document.getElementById('prepElapsed');
    if(el) el.textContent=pad(Math.floor(s/60))+':'+pad(s%60);
  }},1000);
  function fill(d){{
    (d.answers||[]).forEach(function(item,i){{
      var raw=(item.a||'').replace(/&/g,'&amp;').replace(/</g,'&lt;');
      var html=raw.split(/\\n\\s*\\n/).map(function(p){{return p.replace(/\\n/g,' ').trim();}})
                  .filter(Boolean).map(function(p){{return '<p style="margin-bottom:9px">'+p+'</p>';}}).join('');
      if(i===0){{var e=document.getElementById('execbox'); if(e) e.innerHTML=html;}}
      var a=document.getElementById('ans'+i); if(a) a.innerHTML=html;
    }});
  }}
  function reveal(d){{
    fill(d);
    clearInterval(tick);
    var p=document.getElementById('prep'); if(p) p.style.display='none';
    var bar=document.querySelector('.bar'); if(bar) bar.style.display='block';
    if(doPrint) setTimeout(function(){{window.print();}},500);
  }}
  var base='/api/report/cmar?id='+encodeURIComponent(id)+'&lang='+lang+'&template={template}';
  var polls=0;
  function poll(){{
    polls++;
    fetch(base).then(function(r){{return r.json();}}).then(function(d){{
      if(d&&d.status==='ready'){{ reveal(d); }}
      else if(polls>40){{ fetch(base+'&fast=1').then(function(r){{return r.json();}}).then(reveal).catch(function(){{}}); }}
      else {{ setTimeout(poll,3000); }}   // still generating — keep waiting
    }}).catch(function(){{ setTimeout(poll,4000); }});
  }}
  poll();
  // absolute safety: after 100s, force the instant fallback so we never hang
  setTimeout(function(){{
    var p=document.getElementById('prep');
    if(p && p.style.display!=='none'){{
      fetch(base+'&fast=1').then(function(r){{return r.json();}}).then(reveal).catch(function(){{}});
    }}
  }}, 100000);
}})();
</script>
</body></html>"""


@bp.get("/api/report/<rid>/html")
def report_html(rid):
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    template = (request.args.get("template") or "cmar").lower()
    if template not in _TEMPLATES:
        template = "cmar"
    rep = _REPORTS.get(rid) or generate_report("status", lang)
    return Response(_report_html(rep, lang, template=template), mimetype="text/html")


@bp.get("/api/report/generate-html")
def report_generate_html():
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    period = request.args.get("period", "status")
    template = (request.args.get("template") or "cmar").lower()
    if template not in _TEMPLATES:
        template = "cmar"
    rep = generate_report(period, lang)
    meta = _template_meta(template)
    rep["title"] = meta["title_ar"] if lang == "ar" else meta["title_en"]
    return Response(_report_html(rep, lang, template=template), mimetype="text/html")


# ==================================================================
#  EXPORTS — CSV (tables) and DOCX (reports)
# ==================================================================
import csv as _csv
import io as _io


def _csv_response(rows, headers, fname):
    buf = _io.StringIO()
    buf.write("\ufeff")                       # BOM so Excel reads UTF-8 (Arabic) correctly
    w = _csv.writer(buf)
    w.writerow(headers)
    for r in rows:
        w.writerow(r)
    return Response(buf.getvalue(), mimetype="text/csv; charset=utf-8",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@bp.get("/api/export/findings.csv")
@require_auth("viewer")
def export_findings_csv():
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    rows = [[f["cve"], f["asset"], f["sev"], f["score"], f["st"], f["fix"]] for f in _findings_out(lang)]
    return _csv_response(rows, ["ID", "Asset", "Severity", "CVSS", "Status", "Remediation"], "sentinel-findings.csv")


@bp.get("/api/export/events.csv")
@require_auth("viewer")
def export_events_csv():
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    with _events_lock:
        evs = list(_EVENTS)
    rows = [[dt.datetime.fromtimestamp(e["ts"] / 1000).strftime("%Y-%m-%d %H:%M:%S"),
             e.get("kind", ""), e.get("sev", ""), e.get("cve", ""), e.get("asset", ""),
             e.get("text_ar" if lang == "ar" else "text_en", "")] for e in evs]
    return _csv_response(rows, ["Time", "Kind", "Severity", "ID", "Asset", "Detail"], "sentinel-events.csv")


@bp.get("/api/export/software.csv")
@require_auth("viewer")
def export_software_csv():
    sw = sorted(_last_scan.get("software", []), key=lambda a: a["name"].lower())
    rows = [[a["name"], a.get("version", "")] for a in sw]
    return _csv_response(rows, ["Software", "Version"], "sentinel-software.csv")


@bp.get("/api/export/audit.csv")
@require_auth("admin")
def export_audit_csv():
    rows = [[dt.datetime.fromtimestamp(e["ts"] / 1000).strftime("%Y-%m-%d %H:%M:%S"),
             e.get("user", ""), e.get("action", ""), e.get("detail", ""), e.get("ip", "")]
            for e in list(_AUDIT)[::-1]]
    return _csv_response(rows, ["Time", "User", "Action", "Detail", "IP"], "sentinel-audit.csv")


def _report_docx(report, lang, template="cmar"):
    """Build the selected report template as a real Word document (python-docx)."""
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    ar = (lang == "ar")
    rid = report.get("id", "")
    meta = _template_meta(template)
    answers = _REPORTS.get(rid, {}).get("_cmar_" + lang + "_" + template)
    if not answers:
        answers, _ = _cmar_model_answers(report, lang, template=template)
    amap = {a["key"]: a["a"] for a in answers}
    score, threats, opens, _ = _cmar_context(report, lang)

    doc = Document()
    al = WD_ALIGN_PARAGRAPH.RIGHT if ar else WD_ALIGN_PARAGRAPH.LEFT
    def para(text, size=11, bold=False, color=None, align=None):
        p = doc.add_paragraph()
        p.alignment = align if align is not None else al
        r = p.add_run(text)
        r.bold = bold; r.font.size = Pt(size)
        if color:
            r.font.color.rgb = RGBColor(*color)
        return p

    para("CONFIDENTIAL" if not ar else "سرّي", 9, True, (120, 120, 120),
         WD_ALIGN_PARAGRAPH.CENTER)
    para(("Sentinel — Security Operations Centre" if not ar else "سنتينل — مركز العمليات الأمنية"), 10, False, (90, 90, 90))
    # title from the SELECTED template — overrides any title cached on the stored report
    # (so a single underlying report can be rendered under any template)
    title = meta["title_ar"] if ar else meta["title_en"]
    h = para(title, 20, True)
    para((f"Document ID: SOC-{rid}    Date: " if not ar else f"رقم الوثيقة: SOC-{rid}    التاريخ: ") +
         dt.datetime.now().strftime("%Y-%m-%d %H:%M"), 9, False, (110, 110, 110))

    para(("Case Details" if not ar else "تفاصيل الحالة"), 14, True)
    tbl = doc.add_table(rows=2, cols=2)
    tbl.style = "Light Grid Accent 1"
    cells = tbl.rows[0].cells
    cells[0].text = ("Security Score" if not ar else "درجة الأمان"); cells[1].text = f"{score}/100"
    cells = tbl.rows[1].cells
    cells[0].text = ("Active Threats" if not ar else "التهديدات النشطة"); cells[1].text = str(threats)

    para(("Executive Summary" if not ar else "الملخّص التنفيذي"), 14, True)
    para(amap.get("exec", ""), 11)

    para(("Findings & Open Vulnerabilities" if not ar else "النتائج والثغرات المفتوحة"), 14, True)
    if opens:
        ft = doc.add_table(rows=1, cols=4)
        ft.style = "Light Grid Accent 1"
        hdr = ft.rows[0].cells
        for i, t in enumerate((["ID", "Asset", "Severity", "CVSS"] if not ar else ["المعرّف", "الأصل", "الخطورة", "CVSS"])):
            hdr[i].text = t
        for f in opens:
            c = ft.add_row().cells
            c[0].text = str(f.get("cve", "")); c[1].text = str(f.get("asset", ""))
            c[2].text = str(f.get("sev", "")); c[3].text = str(f.get("score", ""))
    else:
        para(("No open findings." if not ar else "لا ثغرات مفتوحة."), 11)

    para(("Standing Information Requirements" if not ar else "متطلّبات المعلومات الأساسية"), 14, True)
    n = 1
    for q in _template_questions(template):
        if q[0] == "exec":
            continue
        para((f"{n}. " if not ar else "") + (q[2] if ar else q[1]), 12, True)
        para(amap.get(q[0], ""), 11)
        n += 1

    para(("This document is confidential and intended for internal use only." if not ar
          else "هذه الوثيقة سرّية ومخصّصة للاستخدام الداخلي فقط."), 8, False, (130, 130, 130))

    bio = _io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


@bp.get("/api/report/<rid>/docx")
def report_docx(rid):
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    template = (request.args.get("template") or "cmar").lower()
    if template not in _TEMPLATES:
        template = "cmar"
    rep = _REPORTS.get(rid) or generate_report("status", lang)
    try:
        data = _report_docx(rep, lang, template=template)
    except ImportError:
        return jsonify({"ok": False, "error": "python-docx not installed. Run: pip install python-docx"}), 500
    except Exception as e:
        return jsonify({"ok": False, "error": f"DOCX generation failed: {e}"}), 500
    fname = "sentinel-" + template + "-" + str(rep.get("id", "report")) + ".docx"
    return Response(data, mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@bp.get("/api/report/templates")
def report_templates():
    """List available report templates (for the UI dropdown)."""
    out = [{"id": k, "title_en": v["title_en"], "title_ar": v["title_ar"]} for k, v in _TEMPLATES.items()]
    return jsonify({"templates": out})


def _pdf_response(report, lang):
    try:
        data = _build_pdf(report, lang)   # always uses live data internally
    except ImportError as e:
        return jsonify({"ok": False, "error": f"PDF library missing: {e}. Run: pip install reportlab arabic-reshaper 'python-bidi==0.4.2'"}), 500
    except Exception as e:
        return jsonify({"ok": False, "error": f"PDF generation failed: {e}"}), 500
    fname = "sentinel-report-" + str(report.get("id", "report")) + ".pdf"
    return Response(data, mimetype="application/pdf",
                    headers={"Content-Disposition": f'attachment; filename="{fname}"'})


@bp.get("/api/report/<rid>/pdf")
def report_pdf(rid):
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    rep = _REPORTS.get(rid) or generate_report("status", lang)
    return _pdf_response(rep, lang)


@bp.get("/api/report/generate-pdf")
def report_generate_pdf():
    lang = "ar" if str(request.args.get("lang", "en")).lower().startswith("ar") else "en"
    period = request.args.get("period", "status")
    rep = generate_report(period, lang)
    rep["title"] = "تقرير الحالة" if lang == "ar" else "Status Report"
    return _pdf_response(rep, lang)


# ==================================================================
#  EMPLOYEE AUTH  (multi-user, self-registration, persisted to disk)
# ==================================================================
import hashlib
import struct

_USERS_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_users.json")
_AUDIT_FILE = os.path.join(os.path.dirname(os.path.abspath(__file__)), "sentinel_audit.json")
_users_lock = threading.Lock()
import hmac as _hmac
import base64 as _b64
import secrets as _secrets
from functools import wraps as _wraps

_ROLES = ("viewer", "analyst", "admin")
_LOCK_THRESHOLD = 5            # failed logins before lockout
_LOCK_MINUTES = 15
_SESSION_HOURS = 12


# ---- password hashing (PBKDF2, with legacy sha256 migration) ----
def _hash_pw(p, salt=None):
    salt = salt or _secrets.token_hex(16)
    h = hashlib.pbkdf2_hmac("sha256", p.encode("utf-8"), salt.encode("utf-8"), 120000)
    return {"algo": "pbkdf2", "salt": salt, "hash": h.hex(), "iter": 120000}


def _verify_pw(stored, p):
    if isinstance(stored, str):                     # legacy sha256+salt
        return stored == hashlib.sha256(("sentinel-salt:" + p).encode("utf-8")).hexdigest()
    try:
        h = hashlib.pbkdf2_hmac("sha256", p.encode("utf-8"), stored["salt"].encode("utf-8"), stored.get("iter", 120000))
        return _hmac.compare_digest(h.hex(), stored["hash"])
    except Exception:
        return False


def _pwd(p):                                         # kept for compatibility
    return _hash_pw(p)


def _password_problem(pw):
    if len(pw) < 8:
        return "Password must be at least 8 characters"
    if not any(c.isalpha() for c in pw) or not any(c.isdigit() for c in pw):
        return "Password must contain both letters and numbers"
    return None


# ---- TOTP (RFC 6238) — works with Google Authenticator / Authy, stdlib only ----
def _totp_secret():
    return _b64.b32encode(_secrets.token_bytes(20)).decode("ascii").rstrip("=")


def _totp_at(secret, t):
    pad = "=" * ((8 - len(secret) % 8) % 8)
    key = _b64.b32decode(secret + pad)
    counter = struct.pack(">Q", int(t // 30))
    h = _hmac.new(key, counter, hashlib.sha1).digest()
    o = h[-1] & 0x0F
    code = (struct.unpack(">I", h[o:o + 4])[0] & 0x7FFFFFFF) % 1000000
    return f"{code:06d}"


def _totp_verify(secret, code):
    code = (code or "").strip().replace(" ", "")
    if len(code) != 6 or not secret:
        return False
    now = time.time()
    return any(_totp_at(secret, now + drift * 30) == code for drift in (-1, 0, 1))


def _otpauth_uri(uname, secret):
    return f"otpauth://totp/Sentinel:{uname}?secret={secret}&issuer=Sentinel&digits=6&period=30"


# ---- audit log (who did what, when) ----
_AUDIT = collections.deque(maxlen=2000)


def _load_audit():
    try:
        for e in _secure_load(_AUDIT_FILE, []):
            _AUDIT.append(e)
    except Exception:
        pass


def _save_audit():
    try:
        _secure_save(_AUDIT_FILE, list(_AUDIT))
    except Exception:
        pass


def audit(action, detail="", user=None):
    try:
        u = user or (_current_session() or {}).get("user") or "-"
    except Exception:
        u = user or "-"
    _AUDIT.append({"ts": int(time.time() * 1000), "user": u, "action": action,
                   "detail": str(detail)[:300], "ip": request.remote_addr if request else ""})
    _save_audit()


def _pwd_legacy_check():
    pass


def _load_users():
    try:
        with open(_USERS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _save_users():
    try:
        with open(_USERS_FILE, "w", encoding="utf-8") as f:
            json.dump(_USERS, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


_USERS = _load_users()          # starts EMPTY — first registered user becomes admin
_current_user = [None]          # most-recent login (used to stamp report author)
_load_audit()


def _initials(name):
    parts = [w for w in name.split() if w]
    return ("".join(p[0] for p in parts[:2]) or "?").upper()


def _user_public(uname):
    u = _USERS.get(uname)
    if not u:
        return None
    locked = u.get("locked_until", 0) > int(time.time() * 1000)
    return {"username": uname, "name": u["name"], "role": u["role"],
            "initials": _initials(u["name"]),
            "email": u.get("email", ""),
            "mfa": bool((u.get("mfa") or {}).get("enabled")),
            "locked": locked,
            "last_login": u.get("last_login"), "last_logout": u.get("last_logout"),
            "online": any(s["user"] == uname for s in _SESSIONS.values())}


@bp.post("/api/auth/register")
def auth_register():
    data = request.get_json(silent=True) or {}
    uname = (data.get("username") or "").strip().lower()
    name = (data.get("name") or "").strip()
    pw = data.get("password") or ""
    first = (len(_USERS) == 0)
    # only an admin can create further accounts once the first one exists
    if not first:
        s = _current_session()
        if not s or not _role_ok(_USERS.get(s["user"], {}).get("role", "viewer"), "admin"):
            return jsonify({"ok": False, "error": "only an admin can create accounts"}), 403
        if not _hmac.compare_digest(request.headers.get("X-CSRF-Token") or "", s["csrf"]):
            return jsonify({"ok": False, "error": "invalid CSRF token"}), 403
    role = "admin" if first else (data.get("role") or "analyst").strip().lower()
    if role not in _ROLES:
        role = "analyst"
    if not uname or not name:
        return jsonify({"ok": False, "error": "username and name are required"}), 400
    prob = _password_problem(pw)
    if prob:
        return jsonify({"ok": False, "error": prob}), 400
    with _users_lock:
        if uname in _USERS:
            return jsonify({"ok": False, "error": "username already exists"}), 409
        _USERS[uname] = {"name": name, "role": role, "pwd": _hash_pw(pw),
                         "email": (data.get("email") or "").strip(),
                         "mfa": None, "failed": 0, "locked_until": 0,
                         "created": int(time.time() * 1000),
                         "last_login": None, "last_logout": None}
        _save_users()
    audit("user_created", f"{uname} ({role})", user=(_current_session() or {}).get("user", uname))
    return jsonify({"ok": True, "user": _user_public(uname), "firstAdmin": first})


@bp.post("/api/auth/login")
def auth_login():
    data = request.get_json(silent=True) or {}
    uname = (data.get("username") or "").strip().lower()
    pw = data.get("password") or ""
    code = data.get("code") or ""
    now = int(time.time() * 1000)
    with _auth_lock:
        u = _USERS.get(uname)
        if not u:
            audit("login_fail", f"unknown user {uname}", user="-")
            return jsonify({"ok": False, "error": "invalid credentials"}), 401
        if u.get("locked_until", 0) > now:
            mins = int((u["locked_until"] - now) / 60000) + 1
            return jsonify({"ok": False, "error": f"account locked, try again in {mins} min"}), 423
        if not _verify_pw(u["pwd"], pw):
            u["failed"] = u.get("failed", 0) + 1
            if u["failed"] >= _LOCK_THRESHOLD:
                u["locked_until"] = now + _LOCK_MINUTES * 60000
                u["failed"] = 0
                audit("account_locked", uname, user="-")
            _save_users()
            audit("login_fail", f"bad password {uname}", user="-")
            return jsonify({"ok": False, "error": "invalid credentials"}), 401
        # password OK — check MFA if enabled
        mfa = u.get("mfa") or {}
        if mfa.get("enabled"):
            if not code:
                return jsonify({"ok": False, "mfaRequired": True}), 200
            if not _totp_verify(mfa.get("secret", ""), code):
                audit("mfa_fail", uname, user=uname)
                return jsonify({"ok": False, "error": "invalid authentication code", "mfaRequired": True}), 401
        # upgrade legacy hash
        if isinstance(u["pwd"], str):
            u["pwd"] = _hash_pw(pw)
        u["failed"] = 0
        u["locked_until"] = 0
        u["last_login"] = now
        u["last_logout"] = None
        _current_user[0] = uname
        tok = _new_session(uname)
        _save_users()
    audit("login", uname, user=uname)
    resp = make_response(jsonify({"ok": True, "user": _user_public(uname),
                                  "csrf": _SESSIONS[tok]["csrf"]}))
    resp.set_cookie("sx_session", tok, httponly=True, samesite="Lax", max_age=_SESSION_HOURS * 3600)
    return resp


@bp.post("/api/auth/logout")
def auth_logout():
    tok = request.cookies.get("sx_session")
    with _auth_lock:
        s = _SESSIONS.pop(tok, None) if tok else None
        uname = s["user"] if s else _current_user[0]
        if uname and uname in _USERS:
            _USERS[uname]["last_logout"] = int(time.time() * 1000)
            _save_users()
        if _current_user[0] == uname:
            _current_user[0] = None
    if uname:
        audit("logout", uname, user=uname)
    resp = make_response(jsonify({"ok": True}))
    resp.delete_cookie("sx_session")
    return resp


@bp.get("/api/auth/me")
@require_auth("viewer")
def auth_me():
    s = _current_session()
    if not s:
        return jsonify({"user": None, "needsSetup": (len(_USERS) == 0)})
    return jsonify({"user": _user_public(s["user"]), "csrf": s["csrf"]})


@bp.get("/api/auth/users")
@require_auth("admin")
def auth_users():
    return jsonify({"users": [_user_public(u) for u in _USERS]})


# ---- MFA enrolment ----
@bp.post("/api/auth/mfa/setup")
def auth_mfa_setup():
    s = _current_session()
    if not s:
        return jsonify({"ok": False, "error": "authentication required"}), 401
    if not _hmac.compare_digest(request.headers.get("X-CSRF-Token") or "", s["csrf"]):
        return jsonify({"ok": False, "error": "invalid CSRF token"}), 403
    secret = _totp_secret()
    with _users_lock:
        _USERS[s["user"]]["mfa"] = {"enabled": False, "secret": secret}
        _save_users()
    return jsonify({"ok": True, "secret": secret, "otpauth": _otpauth_uri(s["user"], secret)})


@bp.post("/api/auth/mfa/enable")
def auth_mfa_enable():
    s = _current_session()
    if not s:
        return jsonify({"ok": False, "error": "authentication required"}), 401
    if not _hmac.compare_digest(request.headers.get("X-CSRF-Token") or "", s["csrf"]):
        return jsonify({"ok": False, "error": "invalid CSRF token"}), 403
    code = (request.get_json(silent=True) or {}).get("code") or ""
    u = _USERS.get(s["user"], {})
    mfa = u.get("mfa") or {}
    if not mfa.get("secret") or not _totp_verify(mfa["secret"], code):
        return jsonify({"ok": False, "error": "invalid code — try again"}), 400
    with _users_lock:
        u["mfa"]["enabled"] = True
        _save_users()
    audit("mfa_enabled", s["user"], user=s["user"])
    return jsonify({"ok": True})


@bp.post("/api/auth/mfa/disable")
def auth_mfa_disable():
    s = _current_session()
    if not s:
        return jsonify({"ok": False, "error": "authentication required"}), 401
    if not _hmac.compare_digest(request.headers.get("X-CSRF-Token") or "", s["csrf"]):
        return jsonify({"ok": False, "error": "invalid CSRF token"}), 403
    with _users_lock:
        _USERS[s["user"]]["mfa"] = None
        _save_users()
    audit("mfa_disabled", s["user"], user=s["user"])
    return jsonify({"ok": True})


@bp.post("/api/auth/user/role")
@require_auth("admin")
def auth_set_role():
    data = request.get_json(silent=True) or {}
    uname = (data.get("username") or "").strip().lower()
    role = (data.get("role") or "").strip().lower()
    if uname not in _USERS or role not in _ROLES:
        return jsonify({"ok": False, "error": "invalid user or role"}), 400
    with _users_lock:
        _USERS[uname]["role"] = role
        _save_users()
    audit("role_changed", f"{uname} -> {role}")
    return jsonify({"ok": True, "user": _user_public(uname)})


@bp.post("/api/auth/user/delete")
@require_auth("admin")
def auth_delete_user():
    data = request.get_json(silent=True) or {}
    uname = (data.get("username") or "").strip().lower()
    s = _current_session()
    if uname == (s or {}).get("user"):
        return jsonify({"ok": False, "error": "cannot delete your own account"}), 400
    with _users_lock:
        if uname in _USERS:
            del _USERS[uname]
            _save_users()
    audit("user_deleted", uname)
    return jsonify({"ok": True})


@bp.get("/api/audit")
@require_auth("admin")
def audit_log():
    items = list(_AUDIT)[-200:][::-1]
    return jsonify({"count": len(_AUDIT), "items": items})


def _alerts_public():
    """Config with secrets masked for display."""
    import copy
    c = copy.deepcopy(_alerts_cfg())
    if c.get("email", {}).get("password"):
        c["email"]["password"] = "********"
    if c.get("telegram", {}).get("token"):
        t = c["telegram"]["token"]
        c["telegram"]["token"] = (t[:6] + "…") if len(t) > 6 else "********"
    return c


@bp.get("/api/settings/alerts")
@require_auth("admin")
def settings_alerts_get():
    return jsonify({"alerts": _alerts_public()})


@bp.post("/api/settings/alerts")
@require_auth("admin")
def settings_alerts_set():
    data = (request.get_json(silent=True) or {}).get("alerts") or {}
    with _settings_lock:
        cur = _SETTINGS.setdefault("alerts", {})
        cur["enabled"] = bool(data.get("enabled", cur.get("enabled", False)))
        cur["min_sev"] = data.get("min_sev", cur.get("min_sev", "high"))
        for ch in ("email", "telegram", "webhook"):
            incoming = data.get(ch) or {}
            slot = cur.setdefault(ch, {})
            for k, v in incoming.items():
                if k in ("password", "token") and isinstance(v, str) and ("****" in v or v.endswith("…")):
                    continue       # keep stored secret when the masked placeholder comes back
                slot[k] = v
    _save_settings()
    audit("alerts_config_changed", "")
    return jsonify({"ok": True, "alerts": _alerts_public()})


@bp.post("/api/settings/alerts/test")
@require_auth("admin")
def settings_alerts_test():
    res = _dispatch_alert_sync("Test alert",
                               "This is a test alert from your Sentinel SOC console.", "critical")
    audit("alert_test", str(res))
    return jsonify({"ok": True, "results": res})


# ----- static (same-origin) -----
PROJECT_ROOT = os.path.dirname(DASHBOARD_DIR)   # parent (holds colors_and_type.css + fonts/)


@bp.get("/")
def index():
    resp = make_response(send_from_directory(DASHBOARD_DIR, "index.html"))
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    return resp


def _no_cache(resp):
    """Force the browser to always re-fetch code files so updates take effect
    immediately (no stale cached .jsx/.js/.css after an upgrade)."""
    resp.headers["Cache-Control"] = "no-store, no-cache, must-revalidate, max-age=0"
    resp.headers["Pragma"] = "no-cache"
    resp.headers["Expires"] = "0"
    return resp


@bp.get("/<path:path>")
def static_files(path):
    # serve from dashboard/ first; fall back to the project root so the
    # self-hosted Thmanyah fonts (../colors_and_type.css, fonts/*.woff2) load.
    # SECURITY: reject any path that escapes the allowed roots (path traversal).
    def _safe(root, rel):
        root_abs = os.path.abspath(root)
        target = os.path.abspath(os.path.join(root_abs, rel))
        # the resolved path MUST stay inside the root
        if target != root_abs and not target.startswith(root_abs + os.sep):
            return None
        return target if os.path.isfile(target) else None

    # code/markup files must never be cached, so upgrades always take effect
    nocache_ext = (".jsx", ".js", ".css", ".html")
    t1 = _safe(DASHBOARD_DIR, path)
    if t1:
        resp = make_response(send_from_directory(DASHBOARD_DIR, path))
        return _no_cache(resp) if path.lower().endswith(nocache_ext) else resp
    t2 = _safe(PROJECT_ROOT, path)
    if t2:
        # never serve our own secrets/data even if they sit under the root
        base = os.path.basename(t2).lower()
        if base.endswith((".key",)) or base.startswith("sentinel_") or base == "sentinel.log":
            return ("forbidden", 403)
        resp = make_response(send_from_directory(PROJECT_ROOT, path))
        return _no_cache(resp) if path.lower().endswith(nocache_ext) else resp
    return ("not found", 404)


# ==================================================================
#  OPERATIONS: file logging, self-health metrics, rate limiting,
#  automatic backups, and data-retention policy.
# ==================================================================
import logging as _logging
from logging.handlers import RotatingFileHandler as _RFH
import collections as _col2

_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_LOG_FILE = os.path.join(_BASE_DIR, "sentinel.log")
_BACKUP_DIR = os.path.join(_BASE_DIR, "backups")
_DATA_FILES = ["sentinel_users.json", "sentinel_reports.json", "sentinel_events.json",
               "sentinel_audit.json", "sentinel_settings.json", "sentinel_threatintel.json"]

_log = _logging.getLogger("sentinel")
if not _log.handlers:
    _log.setLevel(_logging.INFO)
    try:
        h = _RFH(_LOG_FILE, maxBytes=2_000_000, backupCount=5, encoding="utf-8")
        h.setFormatter(_logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
        _log.addHandler(h)
    except Exception:
        pass
    sh = _logging.StreamHandler()
    sh.setFormatter(_logging.Formatter("%(asctime)s %(levelname)s %(message)s"))
    _log.addHandler(sh)

# ---- self-health counters ----
_HEALTH = {"requests": 0, "errors": 0, "latency_ms": _col2.deque(maxlen=200),
           "started": time.time(), "last_alert": 0}
_health_lock = threading.Lock()

# ---- simple per-IP rate limiter for mutating endpoints ----
_RATE = {}                       # ip -> deque[timestamps]
_RATE_LIMIT = int(os.environ.get("SENTINEL_RATE_LIMIT", "120"))   # writes per minute/IP
_rate_lock = threading.Lock()


def _rate_ok(ip):
    now = time.time()
    with _rate_lock:
        dq = _RATE.setdefault(ip, _col2.deque())
        while dq and now - dq[0] > 60:
            dq.popleft()
        if len(dq) >= _RATE_LIMIT:
            return False
        dq.append(now)
        return True


@bp.before_request
def _before_req():
    request.environ["_t0"] = time.time()
    if request.method in ("POST", "PUT", "DELETE", "PATCH"):
        ip = request.remote_addr or "?"
        if not _rate_ok(ip):
            return jsonify({"ok": False, "error": "rate limit exceeded, slow down"}), 429


@bp.after_request
def _after_req(resp):
    try:
        t0 = request.environ.get("_t0")
        if t0:
            ms = (time.time() - t0) * 1000
            with _health_lock:
                _HEALTH["requests"] += 1
                _HEALTH["latency_ms"].append(ms)
                if resp.status_code >= 500:
                    _HEALTH["errors"] += 1
                    _log.error("%s %s -> %s", request.method, request.path, resp.status_code)
    except Exception:
        pass
    return resp


@bp.get("/api/events/search")
def events_search():
    """Indexed event search. Backed by an in-memory SQLite (+ FTS5 when available)
    for sub-millisecond filtering even at 50k+ events. Privacy: the index lives
    purely in RAM; nothing leaves the machine."""
    args = request.args
    since = None
    win = args.get("window")
    if win:
        try:
            mult = {"h": 3600, "d": 86400, "w": 604800}.get(win[-1].lower(), 1)
            since = int(time.time() * 1000) - int(win[:-1] or 1) * mult * 1000
        except Exception:
            since = None
    try:
        limit = max(1, min(int(args.get("limit", 500)), 5000))
    except Exception:
        limit = 500
    items = events_query(kind=args.get("kind") or None,
                         sev=args.get("sev") or None,
                         cve=args.get("cve") or None,
                         search=args.get("q") or None,
                         since_ms=since, limit=limit)
    return jsonify({"count": len(items), "items": items,
                    "fts": _db_init.has_fts,
                    "engine": "sqlite-memory"})


@bp.get("/api/health")
def health():
    """Self-monitoring: process resource use, request stats, latency (Prometheus-ish)."""
    lat = list(_HEALTH["latency_ms"])
    proc_cpu = proc_mem = None
    if HAVE_PSUTIL:
        try:
            p = psutil.Process()
            proc_cpu = round(p.cpu_percent(interval=0.0) / (psutil.cpu_count() or 1), 1)
            proc_mem = round(p.memory_info().rss / 1024**2, 1)
        except Exception:
            pass
    return jsonify({
        "status": "ok",
        "uptime_seconds": int(time.time() - _HEALTH["started"]),
        "requests": _HEALTH["requests"],
        "errors": _HEALTH["errors"],
        "latency_ms_avg": round(sum(lat) / len(lat), 1) if lat else 0,
        "latency_ms_max": round(max(lat), 1) if lat else 0,
        "process_cpu_pct": proc_cpu,
        "process_mem_mb": proc_mem,
        "threads": threading.active_count(),
        "sessions": len(_SESSIONS),
        "findings": len(_FINDINGS),
        "events": len(_EVENTS),
    })


# ---- self resource-use watchdog (logs/alerts if the app itself is heavy) ----
def _self_watch_loop():
    if not HAVE_PSUTIL:
        return
    try:
        p = psutil.Process()
        p.cpu_percent(None)
    except Exception:
        return
    ncpu = psutil.cpu_count() or 1
    while True:
        time.sleep(30)
        try:
            cpu = p.cpu_percent(None) / ncpu
            mem = p.memory_info().rss / 1024**2
            if (cpu > 85 or mem > 1500) and time.time() - _HEALTH["last_alert"] > 600:
                _HEALTH["last_alert"] = time.time()
                msg_en = f"Sentinel itself is using high resources: CPU {cpu:.0f}%, RAM {mem:.0f} MB"
                msg_ar = f"التطبيق نفسه يستهلك موارد عالية: المعالج {cpu:.0f}%، الذاكرة {mem:.0f} ميغابايت"
                _log.warning(msg_en)
                try:
                    add_notification("alert", "warning", msg_en, msg_ar, route="alerts")
                except Exception:
                    pass
        except Exception:
            pass


# ---- automatic backups + retention ----
def _make_backup():
    import zipfile
    try:
        os.makedirs(_BACKUP_DIR, exist_ok=True)
        stamp = dt.datetime.now().strftime("%Y%m%d-%H%M%S")
        path = os.path.join(_BACKUP_DIR, f"sentinel-backup-{stamp}.zip")
        with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as z:
            for fn in _DATA_FILES:
                fp = os.path.join(_BASE_DIR, fn)
                if os.path.exists(fp):
                    z.write(fp, fn)
        # keep only the latest 14 backups
        backups = sorted(f for f in os.listdir(_BACKUP_DIR) if f.endswith(".zip"))
        for old in backups[:-14]:
            try:
                os.remove(os.path.join(_BACKUP_DIR, old))
            except Exception:
                pass
        _log.info("backup written: %s", os.path.basename(path))
        return path
    except Exception as e:
        _log.error("backup failed: %s", e)
        return None


_RETENTION_DAYS = int(os.environ.get("SENTINEL_RETENTION_DAYS", "90"))


def _apply_retention():
    cutoff = int((time.time() - _RETENTION_DAYS * 86400) * 1000)
    try:
        with _events_lock:
            n = len(_EVENTS)
            _EVENTS[:] = [e for e in _EVENTS if e.get("ts", 0) >= cutoff]
            if len(_EVENTS) != n:
                _secure_save(_EVENTS_FILE, _EVENTS)
                _log.info("retention: trimmed %d old events", n - len(_EVENTS))
        # keep the search index in sync (indexed DELETE is much faster than rewriting JSON)
        try:
            with _db_lock:
                _db.execute("DELETE FROM events WHERE ts < ?", (cutoff,))
                _db.commit()
        except Exception:
            pass
    except Exception:
        pass
    try:
        before = len(_AUDIT)
        kept = [e for e in _AUDIT if e.get("ts", 0) >= cutoff]
        if len(kept) != before:
            _AUDIT.clear()
            _AUDIT.extend(kept)
            _save_audit()
    except Exception:
        pass


def _maintenance_loop():
    last_backup = 0
    while True:
        try:
            if time.time() - last_backup > 86400:     # daily backup
                _make_backup()
                _apply_retention()
                last_backup = time.time()
        except Exception:
            pass
        time.sleep(3600)


_ops_started = [False]


def _ensure_ops():
    if _ops_started[0]:
        return
    _ops_started[0] = True
    threading.Thread(target=_self_watch_loop, daemon=True).start()
    threading.Thread(target=_maintenance_loop, daemon=True).start()
    _log.info("Sentinel operations engine started")


@bp.post("/api/backup")
@require_auth("admin")
def backup_now():
    path = _make_backup()
    if not path:
        return jsonify({"ok": False, "error": "backup failed"}), 500
    audit("backup_created", os.path.basename(path))
    return jsonify({"ok": True, "file": os.path.basename(path)})


@bp.get("/api/backup/download")
@require_auth("admin")
def backup_download():
    path = _make_backup()
    if not path or not os.path.exists(path):
        return jsonify({"ok": False, "error": "backup failed"}), 500
    with open(path, "rb") as f:
        data = f.read()
    return Response(data, mimetype="application/zip",
                    headers={"Content-Disposition": f'attachment; filename="{os.path.basename(path)}"'})


def create_app():
    app = Flask(__name__)
    # DoS guard: reject any request body larger than 8 MB before it is buffered
    # into memory (uploads are capped again at 5 MB inside the handler).
    app.config["MAX_CONTENT_LENGTH"] = 8 * 1024 * 1024
    app.register_blueprint(bp)

    # turn any uncaught exception into a clean JSON error instead of an opaque
    # HTTP 500 page, so the UI can show something actionable.
    @app.errorhandler(500)
    def _on_500(e):
        return jsonify({"ok": False, "error": "internal error — check the server console"}), 500

    @app.errorhandler(Exception)
    def _on_exc(e):
        from werkzeug.exceptions import HTTPException
        if isinstance(e, HTTPException):
            return e
        _log.exception("unhandled error")
        return jsonify({"ok": False, "error": f"internal error: {type(e).__name__}"}), 500

    # baseline security response headers (defence-in-depth for a local web UI)
    @app.after_request
    def _security_headers(resp):
        resp.headers.setdefault("X-Content-Type-Options", "nosniff")
        resp.headers.setdefault("X-Frame-Options", "DENY")
        resp.headers.setdefault("Referrer-Policy", "no-referrer")
        resp.headers.setdefault("Cross-Origin-Opener-Policy", "same-origin")
        # The UI is FULLY self-hosted. This CSP forbids loading scripts, styles,
        # fonts, images, or media from anywhere except this machine ('self'),
        # and forbids the page from sending data anywhere off-box (connect-src 'self').
        # 'unsafe-inline'/'unsafe-eval' are required only because the UI compiles
        # JSX in-browser with Babel; no external origin is ever allowed.
        resp.headers.setdefault(
            "Content-Security-Policy",
            "default-src 'self'; "
            "img-src 'self' data: blob:; "
            "media-src 'self' blob:; "
            "style-src 'self' 'unsafe-inline'; "
            "script-src 'self' 'unsafe-inline' 'unsafe-eval' blob:; "
            "worker-src 'self' blob:; "
            "connect-src 'self'; font-src 'self' data:; "
            "object-src 'none'; base-uri 'self'; frame-ancestors 'none'")
        return resp

    _ensure_ops()
    _log.info("Sentinel app created")
    return app


if __name__ == "__main__":
    port = int(os.environ.get("SENTINEL_PORT", "8000"))
    _log.info("starting Sentinel on 127.0.0.1:%d", port)
    create_app().run(host="127.0.0.1", port=port, debug=False, threaded=True)
